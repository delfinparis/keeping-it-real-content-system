#!/usr/bin/env python
# Builds the KIRP interview prep .docx for Chris Wolfe, eXp Realty (El Dorado Hills, CA)
# Built to prompts/06_interview_prep.md v5: research -> draft -> stress test + council -> EP polish
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)

for name, size, color in (('Heading 1', 18, 0x1F3864), ('Heading 2', 14, 0x2E5496), ('Heading 3', 11.5, 0x2E5496)):
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor((color >> 16) & 255, (color >> 8) & 255, color & 255)
    st.font.bold = True

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.8)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def p(text='', bold=False, italic=False, size=None, space=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space)
    r = par.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    return par


def rich(segments, style=None, space=6):
    par = doc.add_paragraph(style=style)
    par.paragraph_format.space_after = Pt(space)
    for seg in segments:
        text = seg[0]
        b = seg[1] if len(seg) > 1 else False
        i = seg[2] if len(seg) > 2 else False
        r = par.add_run(text)
        r.bold, r.italic = b, i
    return par


def bullet(text_bold, text_rest=''):
    par = doc.add_paragraph(style='List Bullet')
    par.paragraph_format.space_after = Pt(2)
    if text_bold:
        par.add_run(text_bold).bold = True
    if text_rest:
        par.add_run(text_rest)
    return par


def quoted(label, text):
    par = doc.add_paragraph(style='Intense Quote')
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.space_before = Pt(3)
    r = par.add_run(label + ' ')
    r.bold = True
    r.italic = True
    r2 = par.add_run(text)
    r2.italic = True
    return par


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            par = cells[i].paragraphs[0]
            run = par.add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def q(num, question, if_vague, reveals, serves, note=None, short=None, permission=None):
    rich([(str(num) + '. ', True), (question, True)], space=2)
    if short:
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(3)
        par.paragraph_format.left_indent = Inches(0.25)
        r = par.add_run('SAY THIS: ')
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        r2 = par.add_run('"' + short + '"')
        r2.bold = True
        r2.font.size = Pt(10.5)
    if permission:
        quoted('First, say this:', '"' + permission + '"')
    quoted('If vague, ask:', if_vague)
    quoted('Ideal answer reveals:', reveals)
    quoted('Serves:', serves)
    if note:
        rich([('PRODUCER NOTE: ', True), (note, False, True)], space=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def bridge(text):
    rich([('BRIDGE TO NEXT BLOCK (read as written): ', True), ('"' + text + '"', False, True)])


LIVE_PRIMARY = "Chris Wolfe Did 33 Deals in His First Year as an Agent. Here Is His Daily Schedule"
LIVE_BACKUP = "He Lost a 12-Year Flipping Business in Austin. Then He Started Over as a Rookie Agent"

# =====================================================================
# TITLE
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('KEEPING IT REAL PODCAST')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Interview Prep: Chris Wolfe')
r.bold = True
r.font.size = Pt(20)
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Chris Wolfe Real Estate at eXp Realty of Northern California  |  El Dorado Hills, CA  |  Target runtime 43 minutes')
r.italic = True
r.font.size = Pt(10)

# =====================================================================
# 1. QUICK REFERENCE CARD
# =====================================================================
h1('1. Quick Reference Card')
p('One page. Glance at this during the interview.', italic=True, space=8)

h3('Who he is')
bullet('Name: ', 'Chris Wolfe. Solo agent with one executive assistant, Chris Wolfe Real Estate at eXp Realty of Northern California.')
bullet('Based: ', 'El Dorado Hills, California, in the Sacramento foothills. Office at 4364 Town Center Boulevard, Suite 114. Lists as far as South Lake Tahoe and Lodi.')
bullet('Markets: ', 'El Dorado Hills luxury, with Serrano as the home turf. Also The Promontory, Ridgeview, The Summit, Kalithea. Active listings on the MLS roster run from 575 thousand to 5.95 million.')
bullet('Guest type: ', 'A, producing agent, with a small C overlay because he built a software product. Product rule applies: one question on Activate Daily, plus whatever Rapid Fire gives you.')

h3('Verified numbers (safe to say on air)')
bullet('Licensed September 20, 2024. ', 'California DRE salesperson license 01894853, issued 9/20/24, expires 9/19/28, no disciplinary action. This is the single most important fact in the packet, because every other public number about him is in conflict with it. See Watch Out For.')
bullet('Two brokerages in his first year. ', 'Lyon Real Estate from 9/25/24 to 3/23/25, then eXp Realty of Northern California from March 2025 to now. Both from the DRE record.')
bullet('Twelve years of flipping is real. ', 'He is CEO of The Wolfe Firm, an Austin and Tampa value add investment firm, on LinkedIn since April 2011. He is also managing member of Wolfe Capital Investments LLC, incorporated in Texas July 28, 2020, with 11 recorded Austin area acquisitions totaling about 5.84 million and the last one observed in August 2023.')
bullet('Team of two. ', 'Chris plus Lexi Mercado, executive assistant, with a property management and retail banking background. That is the whole team page.')
bullet('Social: ', 'Instagram @chriswolfe_realestate, roughly 8,600 followers as of today. The bio says "$42M+ SOLD THIS YEAR." YouTube @chriswolferealestate with neighborhood and market update videos.')
bullet('Reviews: ', '5.0 from 8 reviews on Zillow. Four named client testimonials on his site.')
bullet('His market right now: ', 'El Dorado Hills, August 2026. Median list price about 980 thousand, roughly 72 to 74 days on market, down about 1 to 2 percent year over year. A slow, expensive market.')
bullet('The Austin crash he lived through: ', 'Austin prices fell 10.2 percent from July 2022 to April 2023 against a national decline of about 1 percent. ATTOM reported Austin flippers took a median loss of 18,640 dollars per flip in 2023, the worst of any large metro.')

h3('Guest supplied, do not state as verified')
bullet('', '33 transactions and 510 thousand in net commission in year one. 40 transactions and 1,038,095 in net commission year two to date. From his intake only. The Instagram bio ("$42M+ sold this year") is directionally consistent, but nothing public confirms the counts. Say "you told me" in front of every one of these.')
bullet('', 'Activate Daily, activatedaily.com. His own daily operating software, now sold to other agents, "very inexpensive." The site would not load from here and a web filter flagged the domain, which is probably a new domain false positive, but open it yourself before you say the URL on air. Do not state a price.')
bullet('', 'The mentor worth 500 million plus who gave him the ladder line. Unnamed. Do not ask who it is.')

h3('Personal, publicly shared, use for warmth only')
bullet('', 'Nothing personal is public. No family, no hobbies, no hometown. His phone has a Fresno area code and his firm listed Central California as a buying market, so there may be a California connection before El Dorado Hills. Ask, do not assert.')
bullet('', 'He sent two stories: the open house client who ate a sandwich over the vacant floor, drank his iced tea and asked the assistant for a ride home, and the flip so infested that the walls were painted brown with roach feces and a white square appeared when you took a picture down. Both are levity beats, not questions.')

h3('Contact and social')
bullet('', 'Website eldoradohillsliving.com. Email chris@chriswolferealestate.com. Phone 559.289.8218. Instagram @chriswolfe_realestate. YouTube @chriswolferealestate. Facebook page Chris Wolfe Real Estate.')

h3('Connection to KIR')
bullet('', 'No prior KIR appearance found. No prior podcast appearance of any kind found. He has never told this story on a microphone, which is rare for a guest with these numbers.')
bullet('', 'He is at eXp, so there is no Kale overlap and nothing to manage there.')
bullet('', 'Real Geeks is a sponsor and sells agent websites. Block 3 is about his website. Do not steer toward the sponsor angle.')

h3('THE CORE TOPIC')
p('A man who lost a 12 year flipping business in the Austin crash moved to a market where he knew nobody, got licensed in September 2024, and closed 33 deals in year one with a daily operating system he built before he had a single client. One topic, four angles: the two years of losing, the actual day, marketing the market instead of himself, and the rung he is refusing to reach for.', bold=True)

h3('Overasked questions to avoid')
bullet('', 'There are no prior interviews, so nothing is technically overasked. What is exhausted is his own website copy: "strategic marketing, data driven pricing, strong negotiation." Any question that lets him answer in that language is wasted. Ask for yesterday, a number, or a name.')
bullet('', '"Tell me about your flipping business." He has, in his words, a book of those stories. They are fun and they are not the episode. One levity beat, then back.')
bullet('', '"How do you get so much business so fast." That is the question he gets from everyone, and it is why he built the software. Do not ask it in that form. Every block asks a narrower version of it.')

h3('The "I have interviewed hundreds" moment')
p("Q13, and nowhere else: \"I've interviewed hundreds of agents on this show, and the ones who blow up in year one very often blow up in year three, because they reached before the grip was solid.\"", italic=True)

h3('Live stream title (paste into Restream before you hit record)')
p(LIVE_PRIMARY + ' (' + str(len(LIVE_PRIMARY)) + ' characters)', bold=True)
p('Backup: ' + LIVE_BACKUP + ' (' + str(len(LIVE_BACKUP)) + ' characters)', italic=True)
p('The primary uses his intake number. Confirm the 33 with him in the green room before you paste it. The backup uses only verified facts. Neither is the published episode title. Pick that from 2A after you hear the interview.', italic=True)

h3('WATCH OUT FOR')
bullet('His public bio does not match his license. ', 'His website, Zillow, the Luxury Home Magazine profile and the MLS roster say "over 13 years of experience," "$600 million bought and sold," "over 700 transactions," "Top 250 agent nationwide," and his Serrano guide says "20+ years of living here as a Realtor and a neighbor." The DRE record says licensed September 20, 2024, and his intake says he moved here after Austin. The 13 years and the 600 million are almost certainly the flipping and development career. The 20 years is not explainable. Do not say any of these numbers on air as fact. Ask him in the green room how he wants the prior career described, then use his words.')
bullet('Rapid Fire will spend your best setup. ', 'He sent both his best and worst advice. The best advice, the ladder rung line, is the spine of Block 4. When he gives it, say "Love it, and we\'re coming back to that," then run the callback at Q13 as written. His worst advice, "don\'t go into business for yourself," is safe to spend, and it is a decent one liner to hand back to him at the top of Block 1 if he needs a runway.')
bullet('Rapid Fire will also spend the software. ', 'His tool answer will be Activate Daily. Say "Love it, we\'ll get to that," and do not let him describe it there. Q14 is where it lives, and it is the only product question in the packet.')
bullet('The business loss is personal and possibly legal. ', 'A 12 year business dying over two years may involve partners, lenders or investors who lost money. Ask what it felt like and what he did. Do not ask about lawsuits, bankruptcy, partners or who got hurt. He sets the depth.')
bullet('"Net commission" needs a definition. ', 'His intake says 510 thousand and 1.04 million in net commissions. Ask in the green room whether that is after the brokerage split or gross. Say the number the way he defines it.')
bullet('This episode wants to become flipping war stories. ', 'He has a decade of them and they are good. The roach house is a 20 second laugh, not a block. Use the Drift Guardrail line in Section 4 and come back to the day.')

h3('THE TWO STANDING REMINDERS')
p('Ask the short version.', bold=True, space=2)
p('Count to three before you respond.', bold=True)

doc.add_page_break()

# =====================================================================
# 2. EPISODE FRAMEWORK
# =====================================================================
h1('2. Episode Framework')

h2('2A. Title Options')
table(['#', 'Title', 'Why it works'], [
    ['1', 'He Lost His Flipping Business in the Austin Crash. Then He Did 33 Deals as a Rookie Agent. (Chris Wolfe)',
     'Specific claim, number, name. The loss opens the gap and the rookie number closes it in a way nobody expects. Works for every segment.'],
    ['2', '33 Deals in Year One, 40 by September of Year Two: Chris Wolfe on the Daily Schedule That Built It',
     'Two numbers and the promise. This is the title for the largest audience segment, because it tells them exactly what they are going to get.'],
    ['3', 'Chris Wolfe Started Over at Zero in a Market Where He Knew Nobody. Here Is What He Did Every Day.',
     'Relatability first. Every agent who moved, restarted or is thinking about it sees themselves. Weaker as a click, stronger as a share.'],
], widths=[0.35, 3.6, 3.1])
p('Recommended: #2 for the published episode, if he confirms the numbers on air. The goal is individual producing agents and new agents, and #2 is the only option that carries the numbers and the mechanism in one line. Run #1 as the live title if you want the story hook cold. Both production numbers are his, not verified, so the title lives or dies on him saying them out loud in the interview.', italic=True)

h2('2B. Cold Open Hook')
p('"Chris Wolfe spent 12 years flipping houses in Austin. Then the market turned in 2022 and he spent the next two years watching that business die. He moved to El Dorado Hills, California, where he knew nobody, got his license in September of 2024, and closed 33 transactions in his first year. We\'re going to talk about exactly what his day looks like, and why. Stay tuned."', bold=True)

h2('2C. Episode Arc')
p('Core topic: how a first year agent with no sphere produced 33 closings, and the daily operating system, market marketing and restraint that made it repeatable in year two.', bold=True)
p('Why this topic: his numbers are not the interesting part. The interesting part is that he designed the agent he wanted to be before he had a client, built a day around it, and then marketed the town instead of himself. That is the part the largest slice of this audience cannot get anywhere else, because nobody has interviewed him, and it works whether the listener is in a million dollar market or a 300 thousand dollar one. The flipping career and the software are context and a callback, not the episode.')
p('Four angles, one topic. Block 1 is the two years of losing and the decision that came out of it. Block 2 is the actual day, hour by hour. Block 3 is why his website is named after the town. Block 4 is the rung he is refusing to reach for, with the software handled in one question. Every block ends on something a listener can do this week.', italic=True)

doc.add_page_break()

# =====================================================================
# 3. INTERVIEW QUESTIONS
# =====================================================================
h1('3. Interview Questions')

h2('Rapid Fire (0:00 to 2:00, standardized, read as written)')
p('1. Best real estate advice you have ever received?', space=2)
p('2. Worst real estate advice you have ever received?', space=2)
p('3. One tool or app you cannot run your business without?', space=2)
p('4. What would surprise people most about your day to day?', space=2)
p('No follow ups. "Love it" and move on. Two exceptions, below.', italic=True, space=4)
rich([('EXCEPTION ONE: ', True), ("He already sent his best advice, the ladder rung line from a mentor worth 500 million plus, and it is the spine of Block 4. When he says it, say \"Love it, and we're coming back to that.\" Then run the callback at Q13. Do not discuss it here.", False, True)])
rich([('EXCEPTION TWO: ', True), ("His tool answer will be Activate Daily, his own software. Say \"Love it, we'll get to that.\" Do not let him describe it in Rapid Fire. It gets one question, Q14, and that is the whole product segment.", False, True)])

h2('BLOCK 1: The Two Years of Losing (2:00 to 12:00)')
p('Audience note: new agents and anyone starting over, which in this market is a lot of people. The stakes of the whole episode get set here. If this block is a highlight reel, nothing after it lands.', italic=True, space=2)
p('Arc: the loss, the choice, what the old career actually bought him, the first move for anyone at zero.')

q(1,
  "You ran a flipping business in Austin for 12 years. In 2022 the Austin market turned harder than almost anywhere in the country, and you told me you spent the next two years losing that business. I don't want the montage version. What did losing it actually look like, month to month? What was the first sign, and what was the last property you sold?",
  'What was the month you knew it was over, and what was the last thing you sold?',
  'The real stakes, not the comeback framing. The specific month and the specific property are what make the rest of the episode believable.',
  'All segments.',
  permission="Tell me if this is too personal, but I want to start with the part that didn't work.",
  note="Ask it, then count to three and say nothing. This is the most likely clip in the episode. If he needs a runway, the receipts are in 4F: Austin prices fell 10.2 percent from July 2022 to April 2023, and ATTOM says Austin flippers lost a median of 18,640 dollars per flip in 2023, worst in the country. Do not ask about partners, lenders, lawsuits or bankruptcy. He sets the depth.",
  short='What did two years of losing the business actually look like, month to month?')

q(2,
  "You could have gone back to flipping somewhere cheaper. Instead you got a sales license in a state where you had no database, in a town where you knew nobody. You told me you had to define what kind of agent you wanted to be and put a plan in place. So what did you write down about that agent before you'd closed a single deal?",
  'Give me two things that were on the page. Price point, activity, hours, anything.',
  'That the plan came before the production. This is the thesis of the episode: the agent was designed first, then the day was built to produce him.',
  'New agents, individual agents.',
  short="What did you write down about the agent you wanted to be, before you'd closed one deal?")

q(3,
  "Here's the scene. You're eight months licensed and you're sitting at a listing appointment in Serrano for a two million dollar house. The seller asks how long you've been doing this. What exactly do you say, and what from the 12 years actually shows up in that room?",
  'Say the sentence you say to the seller. Then name one thing from flipping that a normal rookie could not have done in that appointment.',
  'How he frames a prior career honestly, which is the borrowed credibility play every second career agent needs, and the real advantage list: construction, pricing, negotiation, comfort with risk.',
  'New agents, individual agents.',
  note="This is where you let him address the 13 years and 600 million on his website without it being a gotcha. If he says the 600 million is the flipping career, that is the answer. If the seller sentence sounds like his website, ask for the words he actually says out loud.",
  short="Eight months licensed, a two million dollar listing appointment. The seller asks how long you've done this. What do you say?")

q(4,
  "There's an agent listening who is starting over. New market, no sphere, or coming back after a bad couple of years. Before they make a single call, what's the first thing they define, and what's the first thing they do tomorrow morning?",
  'The one decision, and the one action. Not a list.',
  'Portable implementation that needs no capital, no market and no prior career.',
  'New agents, individual agents.',
  short='Agent starting from zero. What do they define first, and what do they do tomorrow morning?')

bridge("So the plan was written before the first deal. Which brings us to the thing everybody asks you about, which is what the actual day looks like.")

h2('BLOCK 2: The Actual Day (12:00 to 22:00)')
p('Audience note: the largest segment, individual producing agents. This is the block they came for and the block they will send to a colleague. Every answer here needs a clock time or a count.', italic=True, space=2)
p('Arc: yesterday, the ranked sources, the block he kept breaking, the first block a listener protects.')

q(5,
  "Take me through yesterday. Not the ideal day, the actual one. What time did the first block start, what was in it, when did prospecting end, and what did you do after that?",
  'What time did the phone come out, and what time did it go away?',
  'The real time blocks, with the real times. This is the artifact the audience wants to steal.',
  'Individual agents, new agents.',
  note="The moment he says \"it depends\" or \"every day is different,\" say \"I believe you, and I want yesterday,\" and ask again. Do not accept a template day.",
  short='Take me through yesterday. Hour by hour. Not the ideal day, the real one.')

q(6,
  "NAR's newest member profile says the median agent with two years or less in the business closed about two transaction sides last year. You closed 33. Those did not come from one source. Rank your prospecting activities by how many of the 33 they actually produced, and tell me roughly how many hours a week each one got.",
  'Of the 33, how many came from a conversation you started versus one that came to you?',
  'Which activities actually produce, and the ratio of outbound to inbound. This is the current moment question with a real number in it, and it is also the most useful 90 seconds in the episode.',
  'Individual agents, new agents, team leaders.',
  note="If he answers with categories, ask for the count per category. The NAR figure is from the 2026 member profile as reported by HousingWire, so say \"about two\" rather than reading it as a precise statistic.",
  short='Rank your prospecting activities by how many of the 33 closings each one actually produced.')

q(7,
  "Every agent who time blocks has one block they keep putting on the calendar and keep breaking. Which one was it for you, and what did you cut from the year one schedule when you built year two?",
  'Name the block. Then name the thing you deleted.',
  'The honest version of time blocking. Nobody keeps every block, and the block he chose to protect tells you what actually drives the business.',
  'Individual agents.',
  short='Which block did you keep breaking, and what did you cut from the schedule in year two?')

q(8,
  "The agent listening has a calendar with nothing on it before 10 a.m. What is the one block they put in first, what goes in it, and what number makes it done for the day?",
  'The block, the activity, and the count.',
  'A single implementable block with a definition of done, which is what turns this from inspiration into homework.',
  'Individual agents, new agents.',
  short="First block an agent puts on the calendar. What's in it, and what number makes it done?")

bridge("The day is how you got the conversations. Your website is how the conversations found you, and it doesn't have your name on it.")

h2('BLOCK 3: Marketing the Market, Not Yourself (22:00 to 32:00)')
p('Audience note: individual agents and team leaders, and every new agent who was told to post their face and a sold sign. The counterintuitive move is that the brand is the town.', italic=True, space=2)
p('Arc: the town, who makes it, the expensive bet that produced nothing, the first piece of content for someone with no budget.')

q(9,
  "Your website is eldoradohillsliving.com. Not chriswolfe.com. It has full guides on Serrano, the Promontory, school ratings, the country club membership tiers, the trail miles. Most agents market themselves. You market the zip code. Why, and what did it cost you in year one, in dollars and hours?",
  'What did the site and the content cost you in year one? Give me a number for money and a number for hours.',
  'The market share logic and the actual investment. The audience needs the price tag to judge whether it is copyable.',
  'Individual agents, team leaders.',
  short='Your website is eldoradohillsliving.com, not your name. Why market the town instead of yourself?')

q(10,
  "Who does the work? You have one executive assistant, Lexi. Walk me through how one of those neighborhood guides or a home tour video gets made, from idea to posted. Who touches it, what tools, and how long does it take?",
  'Idea, draft, edit, post. Who does each step, and which steps are software?',
  'The production system, including how much is AI, how much is outsourced and how much is him at a keyboard at night.',
  'Individual agents, team leaders.',
  short='Who actually makes the content? Walk me through one guide from idea to posted.')

q(11,
  "You told me part of this is taking risks. So what's the most money you've spent on a marketing bet that produced nothing, and how long did you let it run before you killed it?",
  'The dollar amount and the number of months.',
  'Risk tolerance with a number attached, and the kill rule. This is what makes the market marketing thesis credible instead of a highlight reel.',
  'Individual agents, team leaders.',
  permission="You can pass on the dollar figure if you want, but I'd love the real number.",
  short='Most money you spent on marketing that produced nothing. How long before you killed it?')

q(12,
  "The agent listening is in a market they didn't grow up in and has no budget. What's the first piece of market content they make this week, where does it go, and how do they know if it worked?",
  'The piece, the platform, and the signal.',
  'The zero dollar version of Block 3, which is the only version most of the audience will run.',
  'New agents, individual agents.',
  short='No budget, new market. First piece of market content this week, and where does it go?')

bridge("So the day produces conversations and the content produces market share. The last piece is the thing you said was the best advice you ever got, and I told you we'd come back to it.")

h2('BLOCK 4: The Grip and the Next Rung (32:00 to 40:00)')
p('Audience note: everybody. This is the block that gets shared, because every agent who had a good year has been told to hire, expand, and scale, and almost nobody says wait out loud.', italic=True, space=2)
p('Arc: the rung he is not reaching for, the software in one question, what he needed from a brokerage, the objection.')

q(13,
  "I've interviewed hundreds of agents on this show, and the ones who blow up in year one very often blow up in year three, because they reached before the grip was solid. You told me the best advice you ever got came from someone worth over 500 million dollars: sometimes you've got to shore yourself up and make sure your grip is solid before you reach for the next rung. You went from 33 deals to 40 by September and doubled the commission. So what is the rung you are deliberately not reaching for right now, and what does a solid grip look like in numbers?",
  'Is it a team, a second market, a bigger price point, a hire? Name the one you said no to, and the number that would make you say yes.',
  'Restraint with a metric behind it. Cash reserve, closings per month, months of pipeline, whatever he actually watches. If there is no number, it is a nice quote and not a system.',
  'All segments.',
  note="This is the load bearing question of the episode and the most dodgeable one in the packet. He can answer it entirely in philosophy. If he does, say \"I love that, and I want the number,\" and ask again. Say the hundreds of agents line here and nowhere else. Do not ask who the 500 million dollar mentor is.",
  short='What rung are you deliberately not reaching for right now, and what does a solid grip look like in numbers?')

q(14,
  "You mentioned Activate Daily in Rapid Fire. I'll give you 90 seconds on what it is and how it runs your day. Then I want the version for the agent who never buys it. If the software vanished tomorrow, what are the five things you would track on a legal pad?",
  'Name the five things on the legal pad. Then tell me which one you look at first every morning.',
  'The underlying operating system, independent of the product. The five items are the takeaway. The software is the callback.',
  'Individual agents, new agents.',
  note="This is the only product question in the packet, and the second half is the part that matters. Do not read the URL or a price on air until he says them. The site would not load from here and a web filter flagged the domain, so open it yourself first. If he runs past 90 seconds on features, say \"and the legal pad version,\" and let him finish there.",
  short='Ninety seconds on Activate Daily, then the legal pad version for the agent who never buys it.')

q(15,
  "You were at Lyon for your first six months and you've been at eXp since March of 2025. Set the splits aside. In month one, as a brand new agent with a 12 year business behind you, what did you need from a brokerage that you were not getting?",
  'What was the specific thing you needed, and did the new one actually deliver it?',
  'Broker owner gold, and the perspective flip. This is what the smallest segment came for, and D.J. can ask it better than anyone because he is not competing with him.',
  'Broker owners, team leaders.',
  permission="You don't have to name anybody, and I'm not asking you to. I'm asking what you needed.",
  note="Both brokerages are on his public license record, so the premise is safe to state. Ask what he needed, never what they did wrong. Second career agents are the recruit every broker owner listening wants, so the answer is a recruiting brief.",
  short='What did you need from a brokerage in month one that you were not getting?')

q(16,
  "There is an agent listening thinking, easy for him. Twelve years of deals, construction knowledge, probably money in the bank, and he landed in a market where the median house is close to a million dollars. I'm in a 300 thousand dollar market with a car payment. What do you say to that person?",
  "What's the version of your year one that works at a third of the price point?",
  'The objection said out loud and answered directly. This decides whether the episode travels past luxury markets.',
  'Individual agents, new agents.',
  permission="Tell me if this is too personal, but I think it matters to the answer: did you start this with money, or without it?",
  note="Ask the money question first, with the permission clause, then the objection. He lost a business over two years, so the honest answer may be that he started closer to zero than the listener assumes. That is the best possible answer for the audience. Count to three after it.",
  short='Agent in a 300 thousand dollar market says easy for you. What do you tell them?')

h2('The Close (40:00 to 43:00)')

h3('Homework assignment (read verbatim)')
p('"Here\'s what I want you to do before the next episode. Tonight, open tomorrow on your calendar and put one 90 minute prospecting block on it before noon. Write in the block the one activity you are doing and the number of conversations that counts as done. Then do not book anything over it. Not next month. Tomorrow."', bold=True)

h3('Guest close')
bullet('', '"Where can people find you, follow you, or work with you?"')
bullet('The software. ', 'Let him say the name and the address out loud. Repeat back exactly what he says. Do not read a URL you found, do not state a price, and do not call it anything he did not call it.')
bullet('The site. ', 'eldoradohillsliving.com is verified and safe to say. Instagram @chriswolfe_realestate is safe to say.')

h3('Optional levity beat, only if the room needs it')
p('He sent two stories: the open house client eating a sandwich over the floor of a vacant listing who then drank Chris\'s iced tea and asked the assistant for a ride home, and the flip where the walls were painted brown with roach feces and a white square appeared when you took a picture down. Both are 20 second laughs. Say "you told me a story about an open house and a sandwich" and let him run. Do not let him identify the client. Do not spend a question on either.', italic=True)

h3('If you are running long, cut these first')
bullet('Cut 1, Q12. ', 'The no budget content question. Q8 already delivers a do it tomorrow, and Q9 carries the thesis of the block.')
bullet('Cut 2, Q2. ', 'What he wrote down before the first deal. Q4 delivers the same lesson pointed at the listener.')
bullet('Cut 3, Q10. ', 'Who makes the content. Useful, but it is process, and Q9 and Q11 carry the block without it.')
bullet('Cut 4, Q15. ', 'The brokerage question. It serves the smallest segment and it is the most severable from the core topic.')

h3('Never cut')
bullet('', 'Q1, the two years of losing. It is the cold open paid off and the credibility of the entire episode.')
bullet('', 'Q5, yesterday hour by hour. The artifact the audience came for.')
bullet('', 'Q6, the ranked sources. The only place the 33 gets decomposed into things a listener can copy.')
bullet('', 'Q13, the rung he is not reaching for. The thesis, the callback and the most shareable moment.')
bullet('', 'Q16, easy for you to say. Without it this is an episode for luxury agents with a construction background.')

doc.add_page_break()

# =====================================================================
# 4. RESEARCH BRIEF
# =====================================================================
h1('4. Research Brief')
p('Reference material. Read this the morning of, not during the interview.', italic=True, space=8)

h2('4A. Background')
p('Christopher Thomas Wolfe. Mortgage lending and loan modification work at Charter Funding and Federal Home Counselors, per LinkedIn, then CEO of The Wolfe Firm from April 2011, a boutique investment firm buying single family, multifamily and development projects in Austin and Tampa. Managing member of Wolfe Capital Investments LLC, a Texas flipper entity formed July 2020 with 11 recorded Austin area acquisitions worth about 5.84 million, the last in August 2023. By his own account the Austin correction that began in mid 2022 took two years to kill the business. He moved to El Dorado Hills, California, was licensed September 20, 2024, spent six months at Lyon Real Estate, moved to eXp in March 2025, and runs a two person operation with executive assistant Lexi Mercado. He built his own daily planning software, Activate Daily, and now sells it to other agents.')

h2('4B. Career timeline (verified entries only)')
table(['Year', 'Role / Company', 'Notable'], [
    ['Before 2011', 'Mortgage lending and loan modifications, Charter Funding and Federal Home Counselors', 'Per his LinkedIn headline data as indexed by search. Not independently confirmed.'],
    ['April 2011', 'CEO, The Wolfe Firm, Austin TX and Tampa FL', 'Value add single family, multifamily and development. The 12 years of flipping.'],
    ['July 28, 2020', 'Managing member, Wolfe Capital Investments LLC, Austin', 'Texas incorporation. 11 recorded acquisitions totaling about 5.84 million, average about 531 thousand.'],
    ['Aug 2023', 'Last observed Wolfe Capital acquisition', 'Consistent with his account of the business winding down through 2022 to 2024.'],
    ['Sept 20, 2024', 'California salesperson license issued, DRE 01894853', 'Lyon Real Estate from 9/25/24.'],
    ['March 23, 2025', 'Moves to eXp Realty of Northern California', 'Current brokerage, per DRE.'],
    ['2025 to 2026', 'Chris Wolfe Real Estate, El Dorado Hills', '33 deals year one and 40 year two to date, per his intake. Instagram bio says 42 million plus sold this year.'],
], widths=[1.3, 2.9, 2.85])
p('The exact date he moved to California and the date the flipping business formally closed are not verified. Ask, do not state.', italic=True)

h2('4C. What makes him interesting for this audience')
bullet('He has never been interviewed. ', 'No podcast, no trade press, no YouTube interview. Every story in this episode is being told for the first time, which is worth more than a polished guest with a talk track.')
bullet('The plan came before the production. ', 'He defined the agent he wanted to be and built the day before he had a client. That is the exact inverse of how most agents operate, and it is copyable at any price point.')
bullet('Sixteen times the median rookie. ', 'NAR data has the median agent with two years or less at about two sides a year. He says 33. Even if the number is off by a third, it is a different category of first year.')
bullet('The brand is the town. ', 'His site is named after the town, not him, and it is stacked with neighborhood guides most agents would never write. It is a live case study in marketing the market for market share.')
bullet('He lost first. ', 'Twelve years of business gone over two years, and he is willing to say so in an intake form. The failure makes the rest of the episode believable, and it is the part second career agents will forward to each other.')

h2('4D. Key data points')
table(['Stat', 'Source', 'Confidence'], [
    ['Licensed September 20, 2024, DRE 01894853, no discipline', 'California DRE public lookup, pulled 9/3/2026', 'High'],
    ['Lyon Real Estate 9/25/24 to 3/23/25, then eXp Realty of Northern California', 'California DRE public lookup', 'High'],
    ['CEO, The Wolfe Firm, Austin and Tampa, since April 2011', 'LinkedIn company and profile pages as indexed', 'Medium, LinkedIn would not load directly'],
    ['Wolfe Capital Investments LLC, Texas, incorporated 7/28/2020, 11 acquisitions about 5.84 million', 'OpenCorporates, Bizapedia, SFR Analytics investor profile', 'Medium, aggregator data'],
    ['Team of two, Chris plus executive assistant Lexi Mercado', 'eldoradohillsliving.com team page', 'High'],
    ['Instagram about 8,600 followers, bio says 42 million plus sold this year', 'Instagram, pulled 9/3/2026', 'High on the count, self reported on the volume'],
    ['5.0 from 8 Zillow reviews', 'Zillow as indexed by search', 'Medium'],
    ['Active listings 575 thousand to 5.95 million', 'MetroList agent roster', 'High'],
    ['El Dorado Hills, Aug 2026: median list about 980 thousand, 72 to 74 days on market, down 1 to 2 percent year over year', 'Market trackers, August 2026', 'Medium, aggregator data'],
    ['Austin prices down 10.2 percent July 2022 to April 2023; Austin flippers median loss 18,640 dollars per flip in 2023', 'Newsweek citing index data; ATTOM via CultureMap Austin', 'High'],
    ['NAR: agents with two years or less, median about two transaction sides and about 8,100 dollars income', 'NAR 2026 member profile as reported by HousingWire and others', 'Medium, secondary reporting'],
    ['33 transactions, 510 thousand net year one; 40 transactions, 1,038,095 net year two to date', 'His KIR intake only', 'Unverified'],
    ['13 years of experience, 600 million bought and sold, 700 transactions, Top 250 nationwide, 20 plus years living in El Dorado Hills', 'His website, Zillow, Luxury Home Magazine, MetroList', 'Conflicts with the DRE record. Do not state.'],
    ['Activate Daily, activatedaily.com, "very inexpensive"', 'His intake only. Site would not load and a web filter flagged the domain.', 'Unverified'],
], widths=[3.3, 2.4, 1.35])

h2('4E. Previous media appearances')
bullet('None found. ', 'Searched podcast, interview, Inman, HousingWire, RISMedia, eXp and keepingitrealpod.com. No appearances. His YouTube channel is his own market content, not interviews.')
bullet('His own content. ', 'eldoradohillsliving.com carries long form neighborhood guides for 2026 (Serrano, the neighborhoods guide, buying and selling guide, pricing guide, family neighborhoods guide, a "top agents" page that names no competitors). YouTube: "El Dorado Hills Real Estate | Expert Realtor Chris Wolfe." Instagram: listing carousels, home tours, market updates.')
bullet('Overasked, listed on the Quick Reference Card. ', 'Nothing is overasked because nobody has asked. The risk is his own website language and his flipping stories, not a prior interview.')

h2('4F. Their own words')
table(['Quote', 'Where and when', 'Confidence', 'How D.J. uses it'], [
    ['"Sometimes you\'ve got to shore yourself up and make sure your grip is solid before you reach for the next rung on the ladder."',
     'His KIR intake, September 2026, attributed to a mentor worth 500 million plus', 'Verbatim',
     'This is the Q13 callback. Read it back word for word after Rapid Fire spends it, then ask for the rung he is not reaching for and the number.'],
    ['"Don\'t go into business for yourself!"',
     'His KIR intake, September 2026, his worst advice', 'Verbatim',
     'Rapid Fire will spend it. If he needs a runway into Q1, hand it back: "somebody told you not to go into business for yourself, and you did it twice."'],
    ['"I have a lot of people asking me how I got to doing so much business so fast, and this is how."',
     'His KIR intake, September 2026, about Activate Daily', 'Verbatim',
     'Hold for Q14. If the 90 seconds turns into a pitch, read this back and ask for the legal pad version, since "this is how" is the claim the audience needs unpacked.'],
    ['"In 2022 the market crashed in Austin TX where he was based out of, and he spent 2 years losing his business and trying to figure out what to do."',
     'His KIR intake, September 2026', 'Verbatim',
     'Q1. "Two years losing his business and trying to figure out what to do" is his sentence. Read it and ask for the month to month.'],
    ['"After 20+ years of living here as a Realtor and a neighbor, I can tell you that\'s not marketing, that\'s the actual community."',
     'His Serrano lifestyle guide, eldoradohillsliving.com, 2026', 'Verbatim',
     'Do not read this on air. It contradicts his license date and his own intake, and it is almost certainly ghostwritten or AI drafted site copy. Raise it in the green room only if you want to, and only to ask how he wants his tenure described.'],
], widths=[2.7, 1.55, 0.95, 1.85])

h2('4G. Audience relevance')
table(['Segment', 'What they get from this episode'], [
    ['Individual agents', 'A real day with real times, the ranked list of what produced 33 closings, and a kill rule for marketing spend.'],
    ['Team leaders', 'A two person content production system and a case for marketing the territory instead of the team.'],
    ['Broker owners', 'What a second career agent needed from a brokerage in month one and did not get, in his own words, with no vendor pitch attached.'],
    ['New agents', 'Permission to design the agent before the first deal, the first block to protect, and the zero budget version of every tactic in the episode.'],
], widths=[1.5, 5.55])

h2('4H. Landmines')
bullet('The public bio. ', '"13 years of experience," "15 years" on Zillow, "600 million bought and sold," "700 transactions," "Top 250 agent nationwide," and "20+ years living here as a Realtor." Licensed 9/20/24. The 13 years and 600 million are plausibly the flipping and development career counted as real estate experience. The Top 250 has no visible source. The 20 years is not explainable. Do not state any of them. Q3 gives him the chance to say how he frames the prior career, and the green room brief tells you to ask first.')
bullet('The business loss. ', 'Twelve years of business, two years dying, in the worst flipping market in the country. There may be partners, lenders or investors involved. Ask what it looked like and what he did. Do not ask about lawsuits, bankruptcy, who lost money or what happened to the firm legally. He raised it, so he is willing, but he sets the depth.')
bullet('Activate Daily. ', 'Could not be loaded or verified. A web filter flagged the domain as a security threat, which is most likely a false positive on a new domain, but it means you should open the site yourself before you read the address on air. One question, no price, no URL until he says it. Product rule: this is the whole product segment.')
bullet('The net commission numbers. ', 'His intake says 510 thousand and 1,038,095 in "net commissions." Confirm in the green room whether net means after the brokerage split. Say the number the way he defines it, and say "you told me" in front of it.')
bullet('The brokerage move. ', 'Lyon to eXp at six months is on the public license record, so the premise of Q15 is safe. Ask what he needed. Never ask why he left, and do not let it become an eXp commercial either.')
bullet('The mentor. ', 'Worth 500 million plus, unnamed. Do not ask who it is. If he volunteers, fine.')
bullet('The client stories. ', 'The sandwich client is a real person in a small luxury market. Do not let him identify the client, the property or the street.')
bullet('Sponsors. ', 'Real Geeks sells agent websites and IDX. Block 3 is about his website. Do not steer any question toward what platform he uses or whether it has IDX. Courted is a data product and does not intersect.')

h2('Drift guardrail')
p('This episode will try to become a flipping war stories hour, because he has a decade of them and they are genuinely good. The moment it does, use this line: "I want to stay on the agent side of that, because the part people cannot get anywhere else is what you did with it." Then ask the block question again. The roach house and the sandwich client are 20 seconds each, at most, and only if the room needs a laugh.', bold=True)

h2('Green room brief (two minutes before you record)')
bullet('', 'Tell him the show is standardized rapid fire, then one topic explored properly, then homework. About 43 minutes.')
bullet('', 'Tell him the topic is the daily system behind the 33 and the 40, and that you will be asking for yesterday, numbers and names, not philosophy. He built a software product around his day, so he will welcome it.')
bullet('', 'Confirm the numbers: 33 and 510 thousand in year one, 40 and 1.04 million year two to date, and whether "net" means after the split. Ask if you can put 33 in the live title. If he hesitates, run the backup title.')
bullet('', 'Ask how he wants the prior career described on air. His website says 13 years and 600 million. Ask if that is the flipping career and what words he uses, then use his words.')
bullet('', 'Tell him you plan to ask what the two years of losing the business looked like, and ask what is off limits before you are live.')
bullet('', 'Ask him to say the Activate Daily address and price out loud, and write both down. Tell him it gets one question and he gets 90 seconds. Do not guess on air.')
bullet('', 'Confirm the pronunciation of Serrano and Kalithea, and confirm whether he wants eXp Realty of Northern California said in full or just eXp.')

doc.add_page_break()

# =====================================================================
# 5. LIVE STREAM TITLE, DESCRIPTIONS AND HASHTAGS
# =====================================================================
h1('5. Live Stream Title, Descriptions and Hashtags')

h2('5A. Live stream title')
p('Primary: ' + LIVE_PRIMARY + ' (' + str(len(LIVE_PRIMARY)) + ' characters)', bold=True)
p('Backup: ' + LIVE_BACKUP + ' (' + str(len(LIVE_BACKUP)) + ' characters)', bold=True)
p('The primary uses his intake number, so confirm the 33 with him in the green room before you paste it. The backup uses only what the DRE record and the Texas corporate filings verify, and works cold. This is not the published episode title. Pick that from 2A after you hear the interview.', italic=True)

h2('5B. Platform descriptions')

h3('Facebook Live')
p('Chris Wolfe flipped houses in Austin for 12 years, lost the business when that market turned, and started over as a brand new agent in El Dorado Hills, California, where he knew nobody. He says he closed 33 transactions in his first year. Today we are getting into the actual day that produced it: the time blocks, the prospecting, and why his website is named after the town instead of him. Drop your questions in the comments!')

h3('Instagram Live')
p('Lost a 12 year business. Moved to a town where he knew nobody. 33 deals in year one. Chris Wolfe on the daily schedule behind it. #RealEstatePodcast #NewAgent #KeepingItReal')

h3('TikTok Live')
p('He lost everything in the Austin crash and then closed 33 deals as a rookie agent. Here is his actual day. #realestate #realtorlife #newrealtor #realestateagent')

h3('YouTube Live')
p('Chris Wolfe is a real estate agent with eXp Realty in El Dorado Hills, California, licensed in September 2024 after 12 years running a house flipping business in Austin, Texas. On this episode of the Keeping It Real Podcast he breaks down his first year daily schedule, time blocking and prospecting, how he markets the market instead of himself with eldoradohillsliving.com, and why he is not reaching for the next rung yet. Real estate agent daily schedule, prospecting, time blocking, new agent, second career.')

h3('LinkedIn Live')
p('Most agents build the business first and design the system later. Chris Wolfe did it in the other order. After 12 years running a house flipping firm in Austin and two years watching it close, he moved to El Dorado Hills, defined the agent he wanted to be before his first deal, and built a day around it. He says year one produced 33 transactions. We cover the schedule, the prospecting mix, the decision to market the town instead of himself, and why he is refusing to scale yet.')

h2('5C. Hashtag sets')
bullet('Universal: ', '#KeepingItReal #RealEstatePodcast #DJParis #RealtorLife #RealEstateAgent')
bullet('Episode specific: ', '#NewAgent #TimeBlocking #Prospecting #ElDoradoHills #SacramentoRealEstate #SecondCareer #HouseFlipping')
bullet('Guest tag: ', '@chriswolfe_realestate on Instagram and YouTube. Chris Wolfe Real Estate on Facebook. Tag eXp Realty only if he asks.')

doc.add_page_break()

# =====================================================================
# 6. YOUTUBE CHAPTER MARKERS
# =====================================================================
h1('6. YouTube Chapter Markers')
p('Estimates. Adjust after recording. Every title is written to be independently searchable.', italic=True, space=8)
table(['Timestamp', 'Chapter title'], [
    ['0:00', 'He Lost a 12 Year Flipping Business, Then Did 33 Deals as a Rookie Agent'],
    ['2:00', 'Rapid Fire: Best and Worst Real Estate Advice'],
    ['4:00', 'What Two Years of Losing a Business Actually Looked Like'],
    ['8:00', 'Defining the Agent You Want to Be Before Your First Deal'],
    ['10:00', 'How to Talk About a Prior Career at a Listing Appointment'],
    ['12:00', 'His Actual Day, Hour by Hour'],
    ['15:00', 'Ranking the Prospecting Activities That Produced 33 Closings'],
    ['18:00', 'The Time Block He Kept Breaking'],
    ['20:00', 'The First Block Every Agent Should Put on the Calendar'],
    ['22:00', 'Why His Website Is Named After the Town, Not Him'],
    ['25:00', 'How a Neighborhood Guide Gets Made With a Team of Two'],
    ['28:00', 'The Most Expensive Marketing Bet That Produced Nothing'],
    ['30:00', 'Market Content With Zero Budget'],
    ['32:00', 'Make Sure Your Grip Is Solid Before the Next Rung'],
    ['35:00', 'Activate Daily and the Legal Pad Version'],
    ['37:00', 'What a New Agent Needs From a Brokerage in Month One'],
    ['39:00', 'Easy for You: Starting Over in a 300K Market'],
    ['41:00', 'Homework and Where to Find Chris Wolfe'],
], widths=[1.1, 5.95])

doc.add_page_break()

# =====================================================================
# 7. STRESS TEST, COUNCIL REVIEW AND EP POLISH
# =====================================================================
h1('7. Stress Test, Council Review and EP Polish')

h2('7A. Stress test (pass 2, part one)')
table(['#', 'What broke', 'Fix applied'], [
    ['1', 'Fact conflict, the big one. His website, Zillow, Luxury Home Magazine and MetroList say 13 years (Zillow says 15), 600 million sold, 700 transactions, Top 250 nationwide, and 20 plus years living in El Dorado Hills. The California DRE record says licensed 9/20/24. His own intake says he moved here after Austin.',
     'Pulled the DRE record and made it the anchor fact on the card. Moved every public marketing claim into a do not state row in 4D and a landmine in 4H. Built Q3 so he can frame the prior career himself, and added a green room question so D.J. uses his words. Removed 600 million from the cold open and every title.'],
    ['2', 'Unverifiable production numbers. 33 and 510 thousand, 40 and 1,038,095 come from the intake only. Instagram says 42 million plus sold this year, which is directionally consistent but not the same fact.',
     'Labeled guest supplied everywhere. Every on air use is prefixed with "you told me." Live title primary uses the 33 with a green room confirmation step, and a fully verified backup title was written for the case where he hesitates.'],
    ['3', 'Unverifiable product. activatedaily.com would not load on any attempt, and a web filter flagged the domain as a security threat, most likely a new domain false positive.',
     'Removed the URL and any price from every on air line. One question only, Q14, with the second half built to deliver value without the product. Instruction in the card, 4H and the green room brief to open the site personally before saying the address.'],
    ['4', 'Spent question. His best advice, the ladder rung line, is the spine of Block 4. Standardized Rapid Fire will burn it in minute one with no follow up.',
     'Added exception one under Rapid Fire, the "we\'re coming back to that" line, and a callback at Q13 that reads the intake quote back verbatim.'],
    ['5', 'Second spent question. His tool answer will be Activate Daily, which invites a product description in Rapid Fire.',
     'Added exception two under Rapid Fire with the exact deflection line, and pointed it at Q14.'],
    ['6', 'Drift risk. He has "a book" of flipping stories and they are good. The episode becomes a war stories hour by minute 15.',
     'Added the Drift Guardrail with the exact line. Demoted both intake stories to a single optional levity beat in the close. Put the flipping career on the overasked list.'],
    ['7', 'Dodgeable question. Q5 can be answered with "every day is different." Q6 can be answered with "referrals and hard work."',
     'Rewrote Q5 to demand yesterday with a producer note and the exact re-ask line. Rewrote Q6 to demand a ranking by closings produced, with the NAR median as the setup and a count per category as the if vague.'],
    ['8', 'Second dodgeable question. Q13 can be answered entirely in philosophy about patience.',
     'Rewrote it to demand the specific rung he refused and the number that would make him reach. Added the "I love that, and I want the number" re-ask.'],
    ['9', 'Gotcha risk. Q15 could land as why did you leave Lyon after six months.',
     'Added a permission clause, rewrote the ask to what he needed rather than what they failed to do. Both brokerages are on the public license record, so the premise is safe to state.'],
    ['10', 'Legal and personal exposure. A 12 year investment firm dying over two years may involve partners, lenders or investors who lost money.',
     'Added a permission clause at Q1 and a producer note listing what not to ask. Question targets the month to month experience and the decision, not the liabilities.'],
    ['11', 'Client identification. The sandwich story is about a real client in a small, expensive market.',
     'Levity beat only, with an explicit instruction not to let him identify the client, property or street.'],
    ['12', 'Definition risk. "Net commissions" in the intake could be gross commission or after split. Saying "a million dollars in GCI" on air could be wrong in either direction.',
     'Added to the card, 4H and the green room brief. D.J. says the number the way Chris defines it.'],
    ['13', 'Sponsor collision. Real Geeks sells agent websites. Block 3 is entirely about his website.',
     'Added a sponsor landmine. No question asks what platform the site runs on or whether it has IDX.'],
    ['14', 'Secondary sourcing. The NAR new agent median and the Austin flipper loss both come from trade press reporting, not the primary reports.',
     'Marked Medium and High respectively in 4D. Producer note at Q6 tells D.J. to say "about two" rather than read it as a precise statistic.'],
    ['15', 'Runtime overrun. Block 3 originally had five questions and Block 4 had five, including a second product question.',
     'Cut Block 3 to four by folding the credibility question into Block 1 as Q3, cut the second product question entirely, and wrote the four item cut list and the five item never cut list.'],
    ['16', 'Perspective imbalance. Every implementation question assumed a luxury price point.',
     'Q16 names the 300 thousand dollar market listener and asks him to answer them, with a permission clause that lets D.J. ask whether he started with money.'],
], widths=[0.3, 3.2, 3.55])

h2('7B. Council review (pass 2, part two)')
p('Convened Hormozi, MrBeast, Kane, Miller, Lazine, Eric Simon, Welsh and Youshaei, with Heath and Berger called as witnesses. He is a numbers guest with a product and no prior interviews, which is the combination that produces either the most useful episode of the quarter or a 40 minute infomercial, depending on the questions.', italic=True)

h3('Member notes')
table(['Member', 'What they would change'], [
    ['Alex Hormozi', 'The whole value is Block 2, and it was sitting behind ten minutes of backstory. Keep the loss as the hook, but Q5 and Q6 have to demand a clock time and a count, or the audience gets a motivational episode with numbers in the title. Fixed with the re-ask lines.'],
    ['MrBeast', 'The sag is Q10, the content production process, right in the middle. Put the expensive failed bet directly behind it so the block re-hooks on a loss before the bridge. Applied at Q11.'],
    ['Brendan Kane', 'The first live title used the software. Nobody clicks a product. The hook is 33 deals as a rookie, and the second variant is the Austin loss. Run the number live, A/B the loss as published.'],
    ['Donald Miller', 'The hero kept being Chris. Q4, Q8, Q12 and Q16 are the only places the listener is the hero. Never cut Q16, and end every block on the listener.'],
    ['Byron Lazine', 'The current moment question needed a real number. NAR just said the median rookie does two deals and makes eight grand. Put it in front of the 33 at Q6 and let the gap do the work.'],
    ['Eric Simon', 'The forwardable moment is the time block he kept breaking, Q7. Every agent has one. Keep it in the middle of the block where people actually reach it.'],
    ['Justin Welsh', 'Do not break Rapid Fire to protect the ladder quote or the software. Use the callback mechanism the format already has, twice. Agreed and applied.'],
    ['Jon Youshaei', 'Half the chapter titles were about Chris. Rewrote them as searches an agent would type: daily schedule, prospecting activities, time block, listing appointment, new agent brokerage.'],
], widths=[1.5, 5.55])

h3('Title')
table(['#', 'Title', 'Ingredient', 'Curiosity mechanism'], [
    ['1', 'He Lost His Flipping Business in the Austin Crash. Then He Did 33 Deals as a Rookie Agent. (Chris Wolfe)', 'Contradiction plus stat', 'The loss and the rookie number should not coexist. The listener has to hear how.'],
    ['2', '33 Deals in Year One, 40 by September of Year Two: Chris Wolfe on the Daily Schedule That Built It', 'Stat plus concept format', 'Two numbers and a promise of the mechanism. The gap is the word schedule: what could a schedule possibly do that produces this.'],
    ['3', 'Chris Wolfe Started Over at Zero in a Market Where He Knew Nobody. Here Is What He Did Every Day.', 'Relatability plus promise', 'The listener who has moved or restarted supplies the tension themselves.'],
], widths=[0.3, 3.1, 1.15, 2.5])
p('Recommended: #2 for the published episode, contingent on him confirming the numbers on air. Run #1 as the live title if you want the story cold. Both carry unverified production numbers, so the verified backup exists for a reason.', italic=True)

h3('Cold open (sharpened)')
p('"Chris Wolfe spent 12 years flipping houses in Austin. Then the market turned in 2022 and he spent the next two years watching that business die. He moved to El Dorado Hills, California, where he knew nobody, got his license in September of 2024, and closed 33 transactions in his first year. We\'re going to talk about exactly what his day looks like, and why. Stay tuned."', bold=True)
p('The gap opens on the loss and closes on the number. The license date is verified. The 33 is his, and the cold open is spoken after the green room confirmation, so it is safe by the time it airs.', italic=True)

h3('The clip engine')
table(['Q#', 'Question', 'Berger emotion', 'Heath gap'], [
    ['1', 'What did two years of losing the business actually look like, month to month?', 'Anxiety, then awe. High arousal. Every agent who has had a bad year is in this answer.', 'The listener assumes a comeback story has a clean bottom. The month to month is the part nobody has heard.'],
    ['5', 'Take me through yesterday. Hour by hour.', 'Practical value, the classic forwarded clip.', 'The gap is the promise in the title. The schedule is the thing they clicked for and it only closes when he says the times.'],
    ['13', 'What rung are you not reaching for, and what does a solid grip look like in numbers?', 'Validation and a little anger. Refusing to scale out loud gives the sharing agent social currency.', 'Everybody has heard reach for the next rung. Nobody has heard an agent with these numbers name the one he turned down.'],
], widths=[0.35, 2.5, 2.1, 2.1])

h3('Live description scrub')
table(['Platform', 'Verdict and fix'], [
    ['Facebook', 'Keep. Conversational, leads with the loss, hedges the 33 with "he says," ends on the comment prompt. Zero em dashes.'],
    ['Instagram', 'Fixed. The draft opened with his title and brokerage. Rewrote to four fragments and the promise, which is what actually reads on a phone.'],
    ['TikTok', 'Fixed. Removed the software mention. TikTok is the least forgiving place to name a product nobody can look up, and the crash to rookie line is the punchier hook anyway.'],
    ['YouTube', 'Keep. Loaded with the searchable terms: daily schedule, time blocking, prospecting, new agent, second career, El Dorado Hills, eXp.'],
    ['LinkedIn', 'Fixed. Opened on the contrast, most agents build first and design later, which is the only line in the set that reads native to that feed.'],
], widths=[1.1, 5.95])

h3('Arc fix')
p('The draft had five blocks, with the flipping career as its own block between the loss and the day. That is a second episode. Dissolved it. The prior career now shows up in exactly one question, Q3, as the listing appointment scene, where it does work for second career agents instead of entertaining everyone. The software also had its own block. Cut to one question at Q14, inside the restraint block, where the legal pad version turns it into a system question. Sixteen questions, four blocks, one topic.')

h3('Why it should work')
bullet('Curiosity mechanism, Heath: ', 'The cold open states a loss and a number that should not follow it, and never explains how. The explanation is Block 2 and the payoff is Q13.')
bullet('Share driver, Berger: ', 'A schedule with real times is practical value, the most reliable share driver there is. The refusal to scale at Q13 adds social currency on top.')
bullet('Retention move, MrBeast: ', 'Every block ends on an implementation question aimed at the listener, and Block 3 re-hooks on a failed bet before the bridge.')

h3('The dissent')
p('Hormozi still objects that Block 2 is the episode and it should open the show, with the loss saved for the back half where it has earned an emotional payoff. The counter is that the loss is what makes the day believable, and a schedule from a stranger is a listicle. Kane sides with Hormozi on the live title, which is why the primary leads with the 33 and not the crash. The experiment for the next numbers guest is to run the mechanism block first and the story second, and compare retention at the 20 minute mark.')

h2('7C. EP polish (pass 3)')
bullet('', 'Pulled the California DRE record and rebuilt the Quick Reference Card around it, since every public bio number about him conflicts with the license date. Every marketing claim moved to do not state.')
bullet('', 'Rewrote Q1 with a permission clause and a producer note listing what not to ask, because a failed investment firm may carry partners and lenders who are not in the room.')
bullet('', 'Merged the prior career advantage question into Q3 as a listing appointment scene, which gives him the chance to frame the 13 years and 600 million himself and gives second career agents a script.')
bullet('', 'Rewrote Q5 and Q6 to demand yesterday and a ranked count, with the NAR rookie median as the setup for Q6 and the exact re-ask lines in the producer notes.')
bullet('', 'Moved the expensive failed marketing bet to Q11, directly behind the production process question, so Block 3 re-hooks on a loss.')
bullet('', 'Cut the second product question and the software block. Activate Daily now gets Rapid Fire plus Q14, and Q14 is built so the legal pad version is the answer the audience keeps.')
bullet('', 'Added both Rapid Fire exceptions with the exact deflection lines, after the stress test caught that the ladder quote and the software would both be spent in minute one.')
bullet('', 'Added permission clauses at Q1, Q11, Q15 and Q16 and removed them everywhere else, so they still carry weight when they appear.')
bullet('', 'Rewrote Q16 to name the 300 thousand dollar market listener and to ask, with permission, whether he started with money, because the answer decides whether the episode travels.')
bullet('', 'Shortened every SAY THIS line to under 20 words and cut the bridge lines to one clause each side.')
bullet('', 'Rewrote the homework so it produces a calendar block with a definition of done before the listener goes to bed, since the original stopped at intending to prospect.')
bullet('', 'Wrote a fully verified backup live title and added the green room confirmation step in front of the primary, so the 33 never goes on a stream title without him saying it first.')
bullet('', 'Added the Drift Guardrail line and the green room brief after the stress test flagged that a decade of flipping stories is the default gravity of this interview.')

doc.save("/Users/djparis/GitHub Projects/keeping-it-real-content-system/guest-prep/Chris_Wolfe_Interview_Prep.docx")
print("Saved Chris_Wolfe_Interview_Prep.docx")
print("Live primary:", len(LIVE_PRIMARY), "chars | backup:", len(LIVE_BACKUP), "chars")
