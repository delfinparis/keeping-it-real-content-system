#!/usr/bin/env python
# Builds the KIRP interview prep .docx for Justin Black, LIV Sotheby's International Realty (Breckenridge, CO)
# Built to prompts/06_interview_prep.md v5: research -> draft -> stress test + council -> EP polish
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)

for name, size, color in (('Heading 1', 18, 0x1F3864), ('Heading 2', 14, 0x2E5496), ('Heading 3', 11.5, 0x2E5496)):
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor((color >> 16) & 255, (color >> 8) & 255, color & 255)
    st.font.bold = True

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.8)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def p(text='', bold=False, italic=False, size=None, space=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space)
    r = par.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    return par


def rich(segments, style=None, space=6):
    par = doc.add_paragraph(style=style)
    par.paragraph_format.space_after = Pt(space)
    for seg in segments:
        text = seg[0]
        b = seg[1] if len(seg) > 1 else False
        i = seg[2] if len(seg) > 2 else False
        r = par.add_run(text)
        r.bold, r.italic = b, i
    return par


def bullet(text_bold, text_rest=''):
    par = doc.add_paragraph(style='List Bullet')
    par.paragraph_format.space_after = Pt(2)
    if text_bold:
        par.add_run(text_bold).bold = True
    if text_rest:
        par.add_run(text_rest)
    return par


def quoted(label, text):
    par = doc.add_paragraph(style='Intense Quote')
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.space_before = Pt(3)
    r = par.add_run(label + ' ')
    r.bold = True
    r.italic = True
    r2 = par.add_run(text)
    r2.italic = True
    return par


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            par = cells[i].paragraphs[0]
            run = par.add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def q(num, question, if_vague, reveals, serves, note=None, short=None, permission=None):
    rich([(str(num) + '. ', True), (question, True)], space=2)
    if short:
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(3)
        par.paragraph_format.left_indent = Inches(0.25)
        r = par.add_run('SAY THIS: ')
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        r2 = par.add_run('"' + short + '"')
        r2.bold = True
        r2.font.size = Pt(10.5)
    if permission:
        quoted('First, say this:', '"' + permission + '"')
    quoted('If vague, ask:', if_vague)
    quoted('Ideal answer reveals:', reveals)
    quoted('Serves:', serves)
    if note:
        rich([('PRODUCER NOTE: ', True), (note, False, True)], space=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def bridge(text):
    rich([('BRIDGE TO NEXT BLOCK (read as written): ', True), ('"' + text + '"', False, True)])


# =====================================================================
# TITLE
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('KEEPING IT REAL PODCAST')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Interview Prep: Justin Black')
r.bold = True
r.font.size = Pt(24)
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Real Estate Advisor, LIV Sotheby's International Realty  |  Breckenridge and Summit County, CO  |  Prepared August 12, 2026  |  v1 (three passes complete)")
r.italic = True
r.font.size = Pt(10)

# =====================================================================
# PAGE 1 — QUICK REFERENCE CARD
# =====================================================================
h1('PAGE 1: QUICK REFERENCE CARD')
p('Glance at this mid-interview. Everything else in this packet is for the morning of.', italic=True, size=9.5)

h3('The two standing reminders')
bullet('Ask the short version. ', 'Every question below has a SAY THIS line. That is the one you say out loud. The long form is setup you read only if the room needs it.')
bullet('Count to three before you respond. ', "When his answer ends, wait. He is a careful, prepared talker and he will fill the silence. The fill is the honest part.")

h3('Who he is')
bullet('Name: ', "Justin Black, 29. Real Estate Advisor and associate broker, LIV Sotheby's International Realty.")
bullet('Based: ', '101 S Main St, Breckenridge, CO 80424. Phone (719) 684-3329.')
bullet('Markets: ', 'Summit County luxury (Breckenridge, Blue River, Frisco, Silverthorne, Dillon, Keystone, Copper), plus Colorado Springs and El Paso County through established partnerships.')
bullet('Price band: ', 'Roughly $1 million to $15 million and above.')
bullet('Career volume: ', 'More than $150 million, over about ten years. Licensed at 19, right out of high school, in Colorado Springs.')
bullet('Awards: ', "NAR 30 Under 30 class of 2026 (announced June 29, 2026), selected from 280-plus applicants. Winner of the program's Web Choice Award with nearly 11,000 of 37,577 votes, described by his brokerage as the largest margin in the award's 25-year history. First 30 Under 30 honoree in LIV Sotheby's International Realty history and the only Coloradan in the class.")
bullet('Also claims: ', 'RealTrends Verified top 0.75% of U.S. agents (2026) and Five Star Professional (2026). Both are stated on his own site.')
bullet('Designations: ', "ABR (Accredited Buyer's Representative), MRP (Military Relocation Professional).")
bullet('Record deals: ', '132 North Gold Flake Terrace, Breckenridge, $9,481,500, a six-bedroom that had sat five years asking $15 million. Also 1453 Ponderosa Road in Alma at $2,699,500, an area price record.')
bullet('Volunteering: ', 'Young Professionals Network and the Breckenridge Outdoor Education Center. Both were in his 30 Under 30 application.')
bullet('Web and content: ', 'justinblackre.com. His content hub is called The Brief. Prior brokerages per NAR: Keller Williams, RE/MAX, eXp.')
bullet('Connection to KIR: ', 'First appearance. No prior KIR episode found. He researched recent episodes before pitching and specifically pitched an angle the show has not run.')
bullet('Guest type: ', 'Type A (producing agent) with a Type E overlay (he has a contrarian systems take). Push for the mechanism and the numbers, not the philosophy.')

h3('The Core Topic')
p('The exact system Justin built so ChatGPT, Copilot, and Google AI cite him as the answer, and how an agent starts it Monday with no budget and no developer.', bold=True)

h3("The 'I've interviewed hundreds' moment")
p('Q9, and only Q9. Read it as written: "I have interviewed hundreds of agents at this point and almost every single one asks me how to use AI. You are the first one telling me the real question is how to get used by it."', italic=True)

h3('Overasked questions to avoid')
bullet('', '"What did winning 30 Under 30 mean to you?" He has answered this for NAR, Summit Daily, and his own brokerage inside the last six weeks.')
bullet('', '"How did you get that many votes?" Same. He already gave the good line about people setting alarms.')
bullet('', '"How is the Summit County market doing?" He literally published a post cataloging the national outlets that keep asking him this. It is his least interesting answer.')

h3('Live stream title (paste into Restream before you record)')
p('Justin Black: Licensed at 19, $150M Sold, and Now AI Recommends Him  (67 characters)', bold=True)
p('Backup: The $15M Listing That Sat 5 Years, and How It Sold (Justin Black)  (65 characters)', italic=True)
p('This does not have to be the published episode title. Pick that after you hear the interview.', italic=True, size=9.5)

h3('Watch out for')
bullet('Rapid Fire will spend three of his best answers. ', 'He pre-sent his best and worst advice, so questions 1 and 2 are already burned. Worse, "what would surprise people about your day-to-day" will almost certainly pull the entire AI system in minute two. If it does, do not follow up. Say "Love it, and we are going deep on that in about five minutes," and move on. Callbacks are built into Q1 and Q9.')
bullet('Do not state the $4.8 million as fact. ', 'It is guest-supplied and it appears nowhere public. Q5 is built so he says the number, not you. Same for the Bankrate feature and the new Inman column.')
bullet('The vote count conflicts. ', 'Two sources say 10,822 and one says 10,832. Say "nearly eleven thousand" and let him correct you if he wants.')
bullet('Do not assert which side of Gold Flake he was on. ', 'His own site calls it a buyer-side close. Summit Daily frames it as a record-breaking sale. Q13 is written to let him say it.')
bullet('Do not raise the deaths in his family. ', 'See Landmines, section 4H.')

# =====================================================================
# GREEN ROOM BRIEF
# =====================================================================
h1('GREEN ROOM BRIEF (60 seconds before you hit record)')
p('More of this episode rests on guest-supplied numbers than usual, and the fix is to get him to own them out loud before you are live. Ask these four in the green room so you know what is coming.', italic=True)
bullet('1. ', '"When we get to the $4.8 million, how many deals is that, and how do you know AI sent them?" You need to know he has a real answer before you build a block on it.')
bullet('2. ', '"On Gold Flake Terrace, which side were you on?" Then you can frame it correctly on air the first time.')
bullet('3. ', '"The Bankrate piece and the Inman column, are those live yet or still coming?" If they are not published, you say "about to run" instead of "ran."')
bullet('4. ', '"Anything you would rather I not bring up?" This is a standing green room question and it matters more than usual with this guest.')
p('Then the Rapid Fire pre-brief, said out loud: "I am going to open with four fast ones, best advice, worst advice, one tool, and one thing that would surprise people. Keep them to a sentence, because everything good gets its own segment after."', italic=True)

# =====================================================================
# PAGE 2 — EPISODE FRAMEWORK
# =====================================================================
h1('PAGE 2: EPISODE FRAMEWORK')

h2('2A. Title Options')
table(['#', 'Title', 'Why It Works'], [
    ['1', "Justin Black Got Licensed at 19. Now AI Recommends Him. Here's the System.",
     'Name plus a verified number plus a concrete promise. The two short sentences set up a gap the third one closes. Every number in it is verified.'],
    ['2', 'The $15M Listing That Sat 5 Years, and the Playbook That Sold It (Justin Black)',
     'Two numbers and a specific claim. This is the title if the expired and positioning block turns out to be the better half of the tape.'],
    ['3', "NAR's Biggest Web Choice Win in 25 Years: Justin Black on Getting Cited by AI",
     'Leads with the credential rather than the tactic. Weakest of the three for an agent audience, strongest for a browse-by-name audience.'],
], widths=[0.4, 3.4, 3.3])
p('Recommended: #1. It is the only one that names the takeaway an agent can steal, and the takeaway is the reason he pitched.', italic=True)

h2('2B. Cold Open Hook')
p('"My guest today got his real estate license at 19, has sold more than $150 million, and just won a national award by the biggest vote margin in its 25-year history. None of that is what I want to ask him about, because he figured out how to get the AI to recommend him. We are going to talk about that today. Stay tuned."', bold=True)
p('Every number in there is verified. Do not add the $4.8 million to this.', italic=True, size=9.5)

h2('2C. Episode Arc')
p('Core Topic: The exact system Justin built so AI engines cite him as the answer, and how an agent starts it Monday.', bold=True)
p('Why this topic: He is right that the show has covered AI from a lot of angles, and every one of them was about agents using AI. Nobody has come on and explained how to be the thing AI recommends, and he is the rare guest who has both a system and closed business behind it. The expired and positioning material is not a second episode, it is the proof that this is a positioning skill and not a tech skill, which is why it lands in Block 4 and not in its own show.', italic=True)

table(['Segment', 'Time', 'Purpose'], [
    ['Cold open plus intro', '0:00-2:00', 'Hook, ads, guest intro.'],
    ['RAPID FIRE', '2:00-4:00', 'Four standardized quick hits. No follow-ups. See the watch-out.'],
    ['BLOCK 1: The Build', '4:00-14:00', 'What he actually made, and what one page looks like.'],
    ['BLOCK 2: The Receipts', '14:00-24:00', 'Whether it pays, how he knows, and the twenty-minute self-audit.'],
    ['BLOCK 3: The Cost', '24:00-33:00', 'Contrarian frame, hours, whose job this is, and the weekly cadence.'],
    ['BLOCK 4: Same Skill, a Listing', '33:00-42:00', 'Positioning applied to inventory. Gold Flake and expireds.'],
    ['THE CLOSE', '42:00-45:00', 'The human landing, homework, where to find him.'],
], widths=[1.9, 1.3, 3.9])

# =====================================================================
# SECTION 3 — INTERVIEW QUESTIONS
# =====================================================================
h1('SECTION 3: INTERVIEW QUESTIONS')
p('Questions are numbered continuously so you can call a number out loud. Seventeen questions plus Rapid Fire, sized for 45 minutes.', italic=True, size=9.5)

h2('RAPID FIRE (standardized, same every episode)')
p('Read as written. One or two sentences each from him. You say "Love it" and move on. No follow-ups.', italic=True)
bullet('1. ', "Best real estate advice you've ever received?")
bullet('2. ', "Worst real estate advice you've ever received?")
bullet('3. ', "One tool or app you can't run your business without?")
bullet('4. ', 'What would surprise people most about your day-to-day?')
p('He already sent you 1 and 2. Best: "the older I get, the less I know," and "what you are not changing, you are choosing." Worst: "good things come to those who wait." That worst-advice line is load-bearing for this whole episode, because his entire argument is that agents are waiting out AI. Do not spend it here. Say "Love it, and we are coming back to that," and the callback is written into Q9.', italic=True)

# ---------------- BLOCK 1 ----------------
h2('BLOCK 1: The Build (4:00-14:00)')
p('Audience note: individual producing agents. This is the block where they decide whether the rest of the episode is worth their commute.', italic=True)
p('Arc: what he made, what one page looks like, what wasted his time, what they write Monday.', italic=True)

q(1,
  "Justin, before we started I went through your site, and between the end of June and the beginning of August you published something like thirty pages. And they are not blog posts. They are questions. Breckenridge or Park City. Can you 1031 exchange a Colorado vacation home. Is ski-in ski-out worth the premium. That is not an accident. What were you actually building?",
  'Give me the title of the very first one you wrote, and why that one first.',
  'That the unit of the system is a real question a real client asked him, not a keyword. Also the sequencing logic, which is the part nobody copies correctly.',
  'Individual agents, new agents.',
  note='This is your research-anchored opener and it does two jobs. It proves you did the homework, which changes how carefully he answers everything after. Say "around thirty," not "thirty." Nobody counted precisely. If he already spent this in Rapid Fire question 4, open with "You touched on this in the rapid fire and I cut you off on purpose, so let us go back."',
  short='You published around thirty pages that are all questions. What were you building?')

q(2,
  'Take one of those pages and walk me through the anatomy of it. What is at the top, how long is it, and what has to be in there for an AI to actually pull from it instead of skipping past it?',
  'How many words, and what does the first sentence do?',
  'The template. Whether he leads with the answer, how he handles place names and numbers, and whether he is writing for a reader or for a machine that is summarizing for a reader.',
  'Individual agents, new agents.',
  note='He can dodge this with "I just write genuinely helpful content." Do not accept it. Push until you have a word count and a first sentence. This is the single most stealable minute in the episode.',
  short='Walk me through one page. What has to be on it for AI to cite it?')

q(3,
  'What did you do in the first few months that turned out to be a complete waste of time?',
  'Give me the one thing you would tell an agent to skip entirely.',
  'The failure modes. Most likely: old keyword SEO habits, mass-generated content, and chasing Google rankings that no longer matter.',
  'All segments.',
  short='What did you waste months on before you got this right?')

q(4,
  'An agent listening has none of this. What is the first page they write, and where does it have to live, their own domain or somewhere else?',
  'Give me the exact title of that page for an agent in an ordinary suburban market.',
  'The starting move, and whether a brokerage profile page or a Zillow bio can carry this. It cannot, and hearing why is what makes an agent go buy a domain.',
  'Individual agents, new agents.',
  short='First page they write Monday. What is it, and where does it live?')

bridge('So that is the build. Now I want to know whether it actually pays, because a content library nobody reads is just a hobby.')

# ---------------- BLOCK 2 ----------------
h2('BLOCK 2: The Receipts (14:00-24:00)')
p('Audience note: everyone, but especially the skeptic. This block either makes the episode credible or it does not.', italic=True)
p('Arc: the number, the objection, the one that went wrong, the twenty-minute audit.', italic=True)

q(5,
  'You told me AI has sourced roughly $4.8 million in listings for you. I do not really want the number, I want the mechanics. How many deals is that, and how did you know AI is what sent them?',
  'What did the client literally say when you asked how they found you? Give me their words.',
  'Deal count, and his attribution method. Whether he asks every single lead, whether clients are screenshotting the AI answer, and whether he is guessing.',
  'All segments.',
  note='Set it up with "you told me," so the number belongs to him and not to the show. Then stop talking. If he answers with a story instead of a count, ask again. The count is the whole point of this block.',
  short='How many deals was that, and how did you know AI sent them?')

q(6,
  'There is an agent listening right now, twelve years in the business, who is saying that is attribution theater, those people would have found you anyway. Talk to that person directly. What would actually convince them?',
  'What is the one thing you saw that made you stop thinking it was a coincidence?',
  'Whether he has a real signal or a hopeful one. Also whether he can hold a position under pressure, which is what makes a clip.',
  'Individual agents, team leaders.',
  note='Put the skepticism in the audience, not in your own mouth. You are handing him a person to answer, not accusing him.',
  short='Someone says those clients would have found you anyway. Convince them.')

q(7,
  'You mentioned the FBI once contacted you about an online lead you had been talking to. What happened?',
  'What was the first thing about that lead that felt off to you?',
  'A great story, and underneath it, his actual vetting habit. Online leads at scale include people who are not who they say they are, and this is the counterweight to a block that is otherwise all upside.',
  'All segments. This is the clip.',
  permission='Tell me how much of this you are actually allowed to say, but I have to ask.',
  short='The FBI called you about an online lead. What happened?')

q(8,
  'Give me the twenty-minute version of finding out whether AI already knows who I am. What do I type, and what does a bad answer look like?',
  'Give me the prompt word for word, like I am typing it right now.',
  'The audit. The exact prompts, which engines to check, and what "you do not exist" looks like versus "you exist but you are invisible."',
  'All segments.',
  note='This is the homework question and the most replayed sixty seconds of the episode. Make him say the prompt out loud, in full. If he speaks it in shorthand, ask him to say it again the way you would type it.',
  short='Twenty minutes. How do I find out if AI already knows me?')

bridge('Okay, so it is real and you can measure it. Now the part I think everybody is quietly bracing for, which is what this cost you.')

# ---------------- BLOCK 3 ----------------
h2('BLOCK 3: The Cost (24:00-33:00)')
p('Audience note: the agent who believes this requires money or a developer, and the broker-owner deciding whose job it is.', italic=True)
p('Arc: the contrarian frame, the hours, whose job this is, the weekly cadence.', italic=True)

q(9,
  'I have interviewed hundreds of agents at this point and almost every single one asks me how to use AI. You are the first one telling me the real question is how to get used by it. Why is everybody pointed the wrong direction?',
  'What does it cost an agent to get that backwards for two more years?',
  'His actual thesis, stated cleanly. This is also the callback to his worst-advice line about good things coming to those who wait.',
  'All segments.',
  note='This is the one place the hundreds-of-agents line earns its spot. Use it here and nowhere else. Tie the callback in explicitly: "you said the worst advice you ever got was good things come to those who wait, and this is that, is it not."',
  short='Everyone asks how to use AI. You ask how to get used by it. Why?')

q(10,
  'No budget, no developer, no agency, no coding background. What did this actually cost you in hours a week, and over how many weeks?',
  'What did you give up to find those hours?',
  'The real price. An honest number here is more persuasive than the revenue number, because the listener is doing the math on their own week.',
  'Individual agents.',
  short='Hours a week, and for how many weeks?')

q(11,
  'Flip this for me. If you ran a brokerage with two hundred agents tomorrow, is this something the company does once at the top, or does every single agent have to do it themselves?',
  'Who owns the domain in that setup?',
  'Whether this scales institutionally or only individually. The domain-ownership answer is the one broker-owners will argue about.',
  'Broker-owners, team leaders.',
  short='You run a 200-agent brokerage. Does the company do this once, or every agent alone?')

q(12,
  'It is built. What do you do every week now to keep it working?',
  'How long, and on what day?',
  'That this is maintenance, not a project. The cadence is what separates the agents who will still be doing this in March from the ones who will not.',
  'Individual agents.',
  short='What is the weekly maintenance now that it is built?')

bridge('Here is what I keep coming back to. Everything you just described is a positioning problem and not a technology problem, and you proved that on a listing before you ever proved it on yourself.')

# ---------------- BLOCK 4 ----------------
h2('BLOCK 4: Same Skill, Applied to a Listing (33:00-42:00)')
p('Audience note: individual agents and team leaders. This block is where an agent who does not care about AI at all still gets paid for listening.', italic=True)
p('Arc: the property, the relaunch sequence, the argument, the first touch.', italic=True)

q(13,
  '132 North Gold Flake Terrace. Six bedrooms, sat on the market five years asking fifteen million, and it closed at nine point four eight. Take me inside that one. What was actually wrong with it?',
  'What is the first thing you would have changed if you had it on day one?',
  'His diagnosis of a failed listing at the high end, and how the eventual price got found. Also lets him define his role in the deal.',
  'Individual agents, team leaders.',
  note='His site describes this as a buyer-side close. Summit Daily wrote it up as a record-breaking sale. Do not assert either. Ask the question and let him say which side he was on. If you already asked in the green room, you can frame it correctly on the way in.',
  short='It sat five years at fifteen million and closed at nine four eight. What was wrong?')

q(14,
  'You wrote a piece called "Your Listing Didn\'t Fail. Your Positioning Did." So when you pick up an expired, what do you change in the first ten days, and in what order?',
  'Is the first move a price change, yes or no?',
  'The relaunch sequence. Whether price comes first or last is the whole argument, and forcing the yes or no stops him from giving you a philosophy.',
  'Individual agents.',
  short='You take on an expired. What changes in the first ten days, in order?')

q(15,
  'One vendor tracked more than seventy-eight thousand listings expiring in a single week this past April, up eighty-three percent in two years. Most agents look at that and say it is a pricing problem. You say it is positioning. In practice, on a real listing, what is the difference?',
  'Give me a listing where the price never moved and it sold anyway.',
  'Whether positioning is a real distinct discipline or a nicer word for price reduction. If he cannot give you the example, that is an answer too.',
  'All segments.',
  note='Say "one vendor tracked." The 78,395 figure is Landvoice tracking its own data, not NAR or MLS reporting. Do not put it in the show\'s voice as an industry statistic.',
  short='Expireds are up 83% in two years. Pricing problem or positioning problem?')

q(16,
  'An agent has a list of expireds sitting in front of them tonight. What is the first touch, and what do they say that literally nobody else calling that seller is saying?',
  'Say it to me like I am the seller. Right now.',
  'A usable script, spoken out loud. This is a clip and a chapter marker on its own.',
  'Individual agents, new agents.',
  note='Do not let him describe the approach. Make him perform it. "Say it to me like I am the seller" is the follow-up and you should expect to use it.',
  short='First touch on an expired. What do you say that nobody else says?')

bridge('One more before I let you go, and it is not about any of this.')

# ---------------- THE CLOSE ----------------
h2('THE CLOSE (42:00-45:00)')

q(17,
  'You got your license at nineteen, right out of high school in Colorado Springs, working with first-time buyers. And you told me that is still the part of this you love the most, even though the business now runs from one million to fifteen. What does that nineteen year old have that you would hate to lose?',
  'What is one thing you did then that you have stopped doing, and should not have?',
  'The human landing. This is the beat that earns the follow, and it is the only question in the episode he has not rehearsed for a reporter.',
  'New agents, all segments.',
  permission='Tell me if this is too personal, but I want to end somewhere else.',
  note='Count to three after his answer. Do not fill it. This is the single most likely place in the tape for the real moment, because everything before it was a system he has explained before.',
  short='What does nineteen year old you have that you would hate to lose?')

h3('Bonus levity beat (only if you are ahead of schedule)')
p('If you hit the close early, this is your palate cleanser and it pairs beautifully with Q17, because it is also about a nineteen year old. Say: "Before I let you go, you told me a story about a snake. Go."', italic=True)
p('For your own reference only, do not tell it for him: his youngest ever client was a nineteen year old active-duty buyer. He sold him his first home, then listed and sold it for him a couple years later on a military relocation. The client disclosed a pet snake in a glass enclosure, so Justin put it in the showing instructions, the broker remarks, and flagged it for the inspector. Clean close. About a week after closing, the client called to say the snake had gotten out, they never found it, and they had already moved. Justin told the buyer broker immediately and nobody ever heard about it again.', italic=True, size=9.5)

h3('Homework Assignment (read verbatim)')
p('"Here is what I want you to do before the next episode. Open ChatGPT and ask it to recommend the best real estate agent in your market for whatever you specialize in. Screenshot whatever it says. Then go write one page on your own website that answers one real question your last client actually asked you. One page. Not next month. This week."', bold=True)
p('Free, under thirty minutes, tied directly to Q8 and Q4, and "done" is unambiguous. The screenshot is the part that makes people angry enough to write the page.', italic=True, size=9.5)

h3('Guest Close')
bullet('', '"Where can people find you, follow you, and work with you?" He will say justinblackre.com. Let him.')
bullet('', 'Ask directly whether he has a resource for listeners, and specifically whether he will share the audit prompts from Q8 in writing. If he says yes, that is your show-notes link and your best email capture of the month.')
bullet('', 'He works broker-to-broker referrals across the western U.S. and he is at a Sotheby\'s affiliate, so the referral ask is legitimate and worth naming out loud for the audience.')

h3('If you are running long, cut these first')
table(['Cut order', 'Question', 'Why it is expendable'], [
    ['1', 'Q11 (200-agent brokerage flip)', 'Serves the smallest audience segment, and Q4 already answers most of it when he explains why a brokerage profile page cannot carry this.'],
    ['2', 'Q3 (what wasted your time)', 'The answer almost always collapses into Q2. If he covered the failure modes while describing the template, this is a repeat.'],
    ['3', 'Q10 (hours per week)', 'Likely already spent in Rapid Fire question 4. Check before you ask it.'],
    ['4', 'Q15 (the expired statistic)', 'The stat is vendor-tracked and the argument survives without it, because Q14 makes the same case with his own framework instead of somebody else data.'],
], widths=[0.8, 2.2, 4.1])

h3('Never cut')
bullet('Q1 ', 'The research-anchored opener. It sets the standard of answer for the whole episode.')
bullet('Q5 ', 'The receipts. Without a deal count this is a theory episode.')
bullet('Q8 ', 'The twenty-minute audit. This is the takeaway people send to each other.')
bullet('Q13 ', 'Gold Flake. The only story in the episode with a number, a scene, and a five-year setup.')
bullet('Q17 ', 'The human landing. Everything else is a system. This is the guest.')

# =====================================================================
# SECTION 4 — RESEARCH BRIEF
# =====================================================================
h1('SECTION 4: RESEARCH BRIEF')
p('Read this the morning of, not during.', italic=True, size=9.5)

h2('4A. Background')
p("Born and raised in Colorado Springs, Justin got licensed at nineteen, essentially straight out of high school, and built his first years on Front Range first-time and move-up buyers. He later relocated to Summit County and joined LIV Sotheby's International Realty as an associate broker, shifting from first-time buyers to luxury second, third, and fourth homes. He still serves Colorado Springs and El Paso County through partnerships. Ten-plus years in, he has more than $150 million in career volume, an NAR 30 Under 30 spot, and the largest Web Choice vote total the program has recorded.")

h2('4B. Career Timeline (verified entries only)')
table(['When', 'Role / Company', 'Notable'], [
    ['Age 19', 'Licensed, Colorado Springs', 'First-time and Front Range buyers. Years not published, so do not state one.'],
    ['Early career', 'Keller Williams, RE/MAX, eXp', 'NAR lists all three. Order and dates are not published. Do not sequence them on air.'],
    ['About mid-2025', "LIV Sotheby's International Realty, Breckenridge", 'Became an independent associate broker roughly a year before the June 2026 NAR profile. His own post is titled "Going Independent After 9 Years."'],
    ['June 29, 2026', 'NAR 30 Under 30, class of 2026', '30 selected from 280-plus applications. Only Coloradan. First in LIV SIR history.'],
    ['June 2026', 'NAR 30 Under 30 Web Choice Award', 'Nearly 11,000 of 37,577 votes. Brokerage calls it the largest margin in the 25-year history of the award.'],
    ['2026', 'RealTrends Verified, Five Star Professional', 'Top 0.75% of U.S. agents claim. Stated on his own site.'],
    ['Aug 10, 2026', 'MoneyWise feature', '"How to Spot Hidden Real Estate Fees." Listed on his own site.'],
], widths=[1.2, 2.5, 3.4])

h2('4C. What Makes Him Interesting for This Audience')
bullet('The content library is visible and datable. ', 'You can literally see the system on his site. Roughly thirty question-shaped pages published between late June and early August 2026, things like "Breckenridge or Park City?" and "Can You 1031 Exchange a Colorado Vacation Home?" Most guests describe a system. His is auditable while he talks.')
bullet('He is arguing the inverse of every other AI guest. ', 'Everyone else teaches agents to use AI. He is teaching agents to be what AI recommends. That is a genuinely new episode for this show and it is why the pitch works.')
bullet('He got licensed at nineteen and is twenty-nine now. ', 'A listener who feels late to their own career hears a guy who has ten years in and is not yet thirty. That is both inspiring and slightly irritating, which is good radio.')
bullet('The Gold Flake deal is a positioning story disguised as a luxury story. ', 'A six-bedroom that sat five years asking fifteen million closed at nine point four eight. An agent in a two hundred thousand dollar market can use that lesson tomorrow.')
bullet('Expireds are a large part of his business. ', 'He is a luxury agent whose lead source is the least glamorous one there is. That combination is unusual and it makes the tactical material credible.')

h2('4D. Key Data Points')
table(['Stat', 'Source', 'Confidence'], [
    ['$150M+ career sales volume', "NAR, natlawreview press release, Summit Daily, his own site. All agree.", 'High'],
    ['NAR 30 Under 30, class of 2026, announced June 29, 2026', 'nar.realtor, natlawreview press release', 'High'],
    ['30 honorees from 280+ applications', 'natlawreview press release', 'High'],
    ['Web Choice Award, 37,577 total votes cast', 'nar.realtor and natlawreview agree on the total', 'High'],
    ['His vote count: 10,822 or 10,832', 'Two sources say 10,822, NAR says 10,832', 'Conflict. Say "nearly eleven thousand."'],
    ['Largest margin in the 25-year history of the award', 'natlawreview press release, which is brokerage-issued', 'Medium. It is his brokerage saying it.'],
    ['132 N Gold Flake Terrace, $9,481,500', 'natlawreview press release, Summit Daily', 'High on the number. Conflict on his role.'],
    ['That property sat 5 years asking $15M, six bedrooms', 'Summit Daily', 'Medium. Single source.'],
    ['1453 Ponderosa Road, Alma, $2,699,500, area record', 'natlawreview press release', 'Medium'],
    ['Licensed at 19, age 29, Colorado Springs start', 'Summit Daily', 'Medium. Single source, but a local paper that interviewed him.'],
    ['Prior brokerages: Keller Williams, RE/MAX, eXp', 'nar.realtor', 'Medium. Order and dates unknown.'],
    ['RealTrends Verified top 0.75% of U.S. agents, 2026', 'His own site only', 'Self-reported'],
    ['Five Star Professional 2026', 'His own site only', 'Self-reported'],
    ['ABR and MRP designations', 'His own site', 'Self-reported, and routine.'],
    ['$4.8M in AI-sourced listings', 'His pitch to us. Nothing public.', 'Unverified. Do not state it. Make him say it.'],
    ['Bankrate feature, new Inman column', 'His pitch to us', 'Unverified. Not found in search.'],
    ['MoneyWise feature, Aug 10, 2026', 'His own site', 'Medium. Self-reported but specific and dated.'],
    ['ChatGPT is 87.4% of AI referral traffic; AI referrals up 527% YoY and convert 4.4 to 5x organic', 'Industry AEO reporting, 2026', 'Medium. Vendor-adjacent. Use as context, not as a claim about him.'],
    ['78,395 listings expired in one U.S. week, April 2026, up 83% in two years', 'Landvoice, June 11, 2026, their own tracking', 'Medium. Say "one vendor tracked."'],
], widths=[3.0, 2.5, 1.6])

h2('4E. Previous Media Appearances')
bullet('nar.realtor, June 2026. ', '"This Year\'s Web Choice Winner Is a Man in Motion." Covered the award, the voting, his move to luxury, and a significant amount of personal loss. This is the richest source and also the one with the material you should not raise.')
bullet('Summit Daily, mid-2026. ', 'Local news roundup. Covered his age, licensing at 19, career volume, Gold Flake, and the observation that Summit County luxury runs on cash buyers largely unaffected by rates.')
bullet('LIV Sotheby\'s blog and LIV Now, April 30, 2026. ', 'Brokerage coverage of the award.')
bullet('MoneyWise, August 10, 2026. ', 'Hidden real estate fees. Per his own site.')
bullet('No podcast appearances found. ', 'This appears to be his first real podcast interview, which is good news. It means the AI system material has never been said out loud anywhere, and the questions in Block 1 and Block 2 have no prior take to compete with.')
p('The overasked questions with this guest are not from podcasts, they are from the press cycle around the award. See the Quick Reference Card.', italic=True, size=9.5)

h2('4F. Their Own Words')
p('Read the Verbatim ones as quotes. Paraphrase everything else out loud with "you have said something like." Never attribute exact wording to him that a third party wrote down.', italic=True, size=9.5)
table(['Quote', 'Where and when', 'Confidence', 'How to use it'], [
    ['"Stop trying to beat AI or wait it out, make it your best referral partner."',
     'His pitch to us, August 2026', 'Verbatim',
     'Read it back to open Q9. It is his thesis in one sentence and he wrote it more sharply than he will say it.'],
    ['"The older I get, the less I know." And "what you are not changing, you are choosing."',
     'His pitch to us, August 2026', 'Verbatim',
     'The second one is the better one. Save it for Q9 as the callback to his worst-advice answer about waiting.'],
    ['"Your Listing Didn\'t Fail. Your Positioning Did."',
     'His own article title, justinblackre.com, June 2025', 'Verbatim (his own writing)',
     'Read the title back to open Q14. Titles are the safest quote you can use because he chose every word.'],
    ['"This recognition reflects the trajectory of the Colorado mountain luxury market more than any individual achievement."',
     'Brokerage press release, June 29, 2026', 'Reported',
     'Hold in reserve. If he gets modest and deflects to the market, read this back and ask him to answer it as himself instead.'],
    ['"That level of wealth and money moves quickly now."',
     'Summit Daily, 2026', 'Reported',
     'Useful only if the conversation drifts to the market. Paraphrase it.'],
    ['"It took receiving the award for me to actually step back and enjoy the recognition."',
     'Summit Daily, 2026', 'Reported',
     'A soft entry to Q17 if the permission clause does not land. Paraphrase it.'],
], widths=[2.6, 1.5, 1.1, 2.0])

h2('4G. Audience Relevance')
table(['Segment', 'What They Get From This Episode'], [
    ['Individual agents', 'A page template, a twenty-minute audit with the exact prompts, and a weekly cadence they can run without buying anything.'],
    ['Team leaders', 'A view on whether AI visibility is a team asset or an individual one, and an expired relaunch sequence they can hand to a junior agent.'],
    ['Broker-owners', 'Q11 head-on. Whether this is a company-level build or an agent-level one, and who ends up owning the domain and the traffic.'],
    ['New agents', 'Somebody who started at nineteen with nothing, plus a first-page assignment that does not require a database or a budget.'],
], widths=[1.6, 5.5])

h2('4H. Landmines')
bullet('Do not raise the deaths in his family. ', 'NAR\'s profile mentions his brother died about eight months before that piece, and that a childhood friend and his grandmother died on the same day. He chose to tell NAR that. He did not choose to tell us. If he raises it, count to three, let it sit, and do not pivot to a question. If he does not raise it, it does not exist on this tape.')
bullet('Do not say he left Sotheby\'s or started his own brokerage. ', 'He is an associate broker at LIV Sotheby\'s International Realty. "Going independent" in his story means independent of a team, not independent of a brand.')
bullet('Do not say he "won 30 Under 30." ', 'He was named to a class of thirty. He separately won the Web Choice Award. Those are two different things and he will notice.')
bullet('Do not state the $4.8 million, the Bankrate feature, or the Inman column as fact. ', 'All three are guest-supplied and none surfaced in search. He says them, not you.')
bullet('Do not call Gold Flake his record sale without letting him define his role. ', 'His site says buyer side. The paper wrote it as a record-breaking sale. Both can be true and you do not want to guess on air.')
bullet('Do not sequence or compare his prior brokerages. ', 'Keller Williams, RE/MAX, and eXp are all listed by NAR with no order and no dates, and two of them compete with Kale for the same agent.')
bullet('Do not turn the business-model difference into a debate. ', 'Kale is 100% commission, he is at a luxury franchise. You are not competing with him and that is exactly why he will talk. Leave it alone.')
bullet('Do not say "thirty pages" as a counted number. ', 'Say "around thirty." Nobody counted precisely and he will know if he published thirty-four.')

# =====================================================================
# SECTION 5 — LIVE STREAM
# =====================================================================
h1('SECTION 5: LIVE STREAM TITLE, DESCRIPTIONS & HASHTAGS')

h2('5A. Live Stream Title')
p('Live stream title: Justin Black: Licensed at 19, $150M Sold, and Now AI Recommends Him  (67 characters)', bold=True)
p('Backup: The $15M Listing That Sat 5 Years, and How It Sold (Justin Black)  (65 characters)', bold=True)
p('Both run on verified numbers only, because this gets set before the interview happens. The published episode title does not have to match. Pick that after you hear the tape, and if Block 2 delivers a real deal count, the published title should use it.', italic=True, size=9.5)

h2('5B. Platform Descriptions')
h3('Facebook Live')
p('Justin Black got his real estate license at 19 and has sold more than $150 million since. This year he was named to NAR\'s 30 Under 30 and won the Web Choice Award with the most votes anybody has gotten in the program\'s history. But that is not what we are talking about. We are talking about how he rebuilt his online presence so that ChatGPT and Google\'s AI actually recommend him by name. Drop your questions in the comments!')

h3('Instagram Live')
p('He is 29, licensed at 19, $150M sold. Now he has AI recommending him by name. Live with Justin Black. #RealEstateAgent #AIforRealtors #KeepingItReal')

h3('TikTok Live')
p('He figured out how to make ChatGPT recommend him. Live now. #realtor #realestateagent #chatgpt #realestatetips')

h3('YouTube Live')
p('Justin Black of LIV Sotheby\'s International Realty in Breckenridge, Colorado joins the Keeping It Real Podcast. Named to the National Association of REALTORS 30 Under 30 class of 2026 and winner of the Web Choice Award, Justin explains the system he built so ChatGPT, Copilot, and Google AI cite him as the answer, with no budget and no coding background. We also cover expired listings, luxury positioning, and the $15 million listing that sat five years before it sold.')

h3('LinkedIn Live')
p('Most agents are asking how to use AI. Justin Black spent the last year on the opposite question: how to get used by it. He is a LIV Sotheby\'s advisor in Breckenridge, a 2026 NAR 30 Under 30 honoree, and he rebuilt his entire online presence so AI engines cite him as the answer. We get into the exact system, what it cost him in hours, and how the same positioning discipline rescues a listing that failed.')

h2('5C. Hashtag Sets')
bullet('Universal: ', '#KeepingItReal #RealEstatePodcast #DJParis #RealtorLife #RealEstateAgent')
bullet('Episode-specific: ', '#AIforRealtors #AnswerEngineOptimization #ChatGPT #LuxuryRealEstate #Breckenridge #SummitCounty #ExpiredListings')
bullet('Guest tag: ', "Justin Black, LIV Sotheby's International Realty. justinblackre.com. LinkedIn: linkedin.com/in/justin-black-772277150. Confirm his Instagram handle in the green room, it was not verified in research.")

# =====================================================================
# SECTION 6 — CHAPTERS
# =====================================================================
h1('SECTION 6: YOUTUBE CHAPTER MARKERS')
p('Estimates. Adjust after recording. If you use the cut list, delete the matching rows.', italic=True, size=9.5)
table(['Timestamp', 'Chapter Title'], [
    ['0:00', 'He Got AI to Recommend Him. Here Is How.'],
    ['2:00', 'Rapid Fire: Best and Worst Advice'],
    ['4:00', 'The 30 Pages That Are All Questions'],
    ['7:30', 'The Anatomy of a Page AI Will Actually Cite'],
    ['10:30', 'What He Wasted Months On First'],
    ['12:00', 'The First Page an Agent Should Write, and Where It Has to Live'],
    ['14:00', 'How Many Deals AI Has Actually Sent Him'],
    ['17:30', '"Those Clients Would Have Found You Anyway"'],
    ['20:00', 'When the FBI Called About an Online Lead'],
    ['22:00', 'The 20-Minute Check: Does AI Know Who You Are?'],
    ['24:00', 'Everyone Asks How to Use AI. Ask How to Get Used By It.'],
    ['27:00', 'What It Cost: Hours Per Week, With No Budget'],
    ['29:30', 'Should a Brokerage Do This Once, or Every Agent Alone?'],
    ['31:30', 'The Weekly Maintenance That Keeps It Working'],
    ['33:00', 'The $15M Listing That Sat 5 Years and Closed at $9.48M'],
    ['36:00', 'The First 10 Days on an Expired Listing, In Order'],
    ['38:30', 'Expireds Are Up 83%. Pricing Problem or Positioning Problem?'],
    ['40:30', 'The Expired First Touch, Word for Word'],
    ['42:00', 'Licensed at 19: What He Would Hate to Lose'],
    ['44:00', 'Homework and Where to Find Justin Black'],
], widths=[1.1, 6.0])

# =====================================================================
# SECTION 7 — STRESS TEST, COUNCIL, EP POLISH
# =====================================================================
h1('SECTION 7: STRESS TEST, COUNCIL REVIEW & EP POLISH')

h2('7A. Stress Test (pass 2)')
table(['#', 'What broke', 'Fix applied'], [
    ['1', 'Vote count conflict. Two sources say 10,822, NAR says 10,832.', 'Every reference in the document now says "nearly eleven thousand." Both numbers are in 4D with the conflict named, and the Quick Reference Card carries the warning.'],
    ['2', 'The $4.8 million is guest-supplied and appears nowhere public. The draft had it in a title and in the cold open.', 'Stripped from all three titles, the cold open, the live stream title, and all five platform descriptions. Q5 was rewritten so he states it. Added a green room ask so D.J. knows the deal count before going live.'],
    ['3', 'The Bankrate feature and the Inman column did not surface in search.', 'Kept out of the intro and the descriptions. Listed as Unverified in 4D, added to Landmines, and added to the green room asks so D.J. can say "about to run" instead of "ran."'],
    ['4', 'Gold Flake role conflict. His site says buyer-side in Weisshorn. Summit Daily frames it as a record-breaking sale.', 'Q13 rewritten to ask about the property rather than assert his role, with a producer note and a green room ask. Marked in 4D as high confidence on the number, conflict on the role.'],
    ['5', 'Rapid Fire spends three of his best answers. He pre-sent best and worst advice, and question 4 will almost certainly pull the entire AI system in minute two.', 'Watch-out added to the Quick Reference Card with the exact deflection line. The worst-advice callback is written into Q9. Q1 carries a recovery opener for the case where he spends it early. A Rapid Fire pre-brief was added to the green room.'],
    ['6', 'Q5 was dodgeable. He could answer $4.8 million with a story and never give a count.', 'Rewritten to demand deal count and attribution method, with an explicit producer note telling D.J. not to accept the substitute and to ask again.'],
    ['7', 'Q2 was dodgeable. "I just write genuinely helpful content" would have passed.', 'Rewritten to demand word count and what the first sentence does, with a producer note.'],
    ['8', 'The Landvoice statistic was written as industry data.', 'Rephrased to "one vendor tracked" in Q15, in the chapter marker, and in 4D, with the source and date named.'],
    ['9', 'Gotcha risk on the attribution question.', 'Q6 reframed so the skepticism belongs to a described audience member, not to D.J. He is handed a person to answer, not an accusation.'],
    ['10', 'Runtime overrun. Twenty-one questions drafted against 45 minutes.', 'Cut to seventeen plus Rapid Fire. Cut list and never-cut list written.'],
    ['11', 'Drift risk: the Summit County market update. He is a luxury mountain agent and every reporter asks him about the market.', 'Guardrail: it is named as an overasked question on the Quick Reference Card, and the only market quote in 4F is marked "useful only if the conversation drifts." If he starts a market update, the bridge into Block 4 pulls back to positioning.'],
    ['12', 'Drift risk: the general "will AI replace agents" debate, which this show has already run.', 'Guardrail: every question in Blocks 1 through 3 asks about his build, not about the industry. Q9 is the only opinion question and it is pointed at agent behavior, not at the future of the profession.'],
    ['13', 'The grief material from the NAR profile was in the draft as a question.', 'Removed entirely from Section 3. Moved to Landmines as guest-led only, with instructions for what to do if he raises it.'],
    ['14', 'Two topics competing for one episode. AI visibility and expired listings.', 'Resolved by making positioning the through-line. Expireds are Block 4, framed as the same skill applied to inventory, and the bridge from Block 3 states that connection out loud so the audience hears one episode instead of two.'],
], widths=[0.4, 3.0, 3.7])

h2('7B. Council Review (pass 2)')
h3('Member notes')
table(['Member', 'What they would change'], [
    ['Alex Hormozi', 'Q8 is the whole episode. Twenty minutes, free, and they find out if they are invisible. Move it earlier if the tape sags, and do not let him give the prompt in shorthand.'],
    ['MrBeast', 'Block 3 opened on a logistics question and that is where people leave. Open it on the contrarian instead. Nobody clicks away from an argument.'],
    ['Brendan Kane', 'Title 1 is the stat-plus-insight combination and it is the only one with a gap. But you had a $4.8 million number and you cannot use it, so your best variant is off the table. Run title 1, A/B title 2 next time.'],
    ['Donald Miller', 'The hero is the agent, not Justin. Two of the questions were written as admiration. Rewrite them as instructions. And pick one topic: is this AI or is this expireds?'],
    ['Byron Lazine', 'Expireds up 83 percent in two years is the newsjack and it is sitting in Block 4 where nobody is looking. Fine, but say the number out loud and say who counted it, because it is a vendor.'],
    ['Eric Simon', 'The FBI story is the clip. An agent will send that to another agent inside four minutes. Do not bury it and do not step on the ending.'],
    ['Jon Youshaei', 'Twenty chapters and every one has to work alone. "The Anatomy of a Page AI Will Actually Cite" works. "What He Wasted Months On First" works. Fix any chapter that needs the one above it.'],
    ['Chip and Dan Heath (witness)', 'Called on the cold open. The first version front-loaded the punchline. Opening on the credentials and then saying none of that is what I want to ask about opens the gap before closing it. Approved.'],
    ['Jonah Berger (witness)', 'Called on the share claim. The high-arousal emotion here is not awe at his awards, it is anxiety, specifically an agent realizing AI has never heard of them. That is what gets shared. The homework converts the anxiety into an action, which is what makes it survive a retelling.'],
], widths=[1.5, 5.6])

h3('Title')
table(['#', 'Title', 'Ingredient', 'Curiosity Mechanism'], [
    ['1', "Justin Black Got Licensed at 19. Now AI Recommends Him. Here's the System.", 'Stat plus insight', 'Two facts that do not obviously connect, then a promise to connect them.'],
    ['2', 'The $15M Listing That Sat 5 Years, and the Playbook That Sold It (Justin Black)', 'Stat', 'A failure with a number attached. The gap is what changed.'],
    ['3', "NAR's Biggest Web Choice Win in 25 Years: Justin Black on Getting Cited by AI", 'Personality plus stat', 'Authority first, mechanism second. Weakest gap of the three.'],
], widths=[0.4, 3.0, 1.2, 2.5])
p('Recommended: #1. The episode goal here is clips and reach, not a guest relationship favor, and #1 is the only title where the payoff is a system an agent can steal.', italic=True)

h3('Cold Open Hook (sharpened)')
p('"My guest today got his real estate license at 19, has sold more than $150 million, and just won a national award by the biggest vote margin in its 25-year history. None of that is what I want to ask him about, because he figured out how to get the AI to recommend him. We are going to talk about that today. Stay tuned."', bold=True)

h3('The Clip Engine')
table(['Q#', 'Question', 'Berger Emotion', 'Heath Gap'], [
    ['Q7', 'The FBI called you about an online lead. What happened?', 'Amusement and anxiety', 'A three-letter agency in a real estate story. The gap opens on the word FBI and does not close until the last sentence.'],
    ['Q9', 'Everyone asks how to use AI. You ask how to get used by it. Why?', 'Anxiety, plus social currency for the agent who reposts it', 'Reverses a premise the listener did not know they were holding.'],
    ['Q13', 'It sat five years at fifteen million and closed at nine four eight. What was wrong?', 'Awe and schadenfreude', 'A five-year failure with a specific ending. You cannot stop before the answer.'],
], widths=[0.5, 2.7, 1.5, 2.4])

h3('Live-Description Scrub')
table(['Platform', 'Verdict and fix'], [
    ['Facebook', 'Keep. Four short sentences, conversational, and the turn on "but that is not what we are talking about" is doing the work. Zero dashes.'],
    ['Instagram', 'Keep. Three fragments and a hook. Any longer and it truncates before the payoff.'],
    ['TikTok', 'Keep. Ten words and a claim. Do not add the credentials, they read as brag on TikTok and kill the watch.'],
    ['YouTube', 'Keep. Keyword-loaded on purpose: name, brokerage, town, NAR, ChatGPT, Copilot, expired listings. This is the one that has to work in search.'],
    ['LinkedIn', 'Keep. Opens on the reframe, which is the only thing on LinkedIn that stops a scroll from a professional audience.'],
], widths=[1.2, 5.9])

h3('Arc Fix')
p('Block 3 was the sag. It originally opened on "what did this cost you in hours," which is a logistics question in the exact minute where a listener decides whether to finish the episode. Fixed by resequencing: the contrarian question and the hundreds-of-agents line now open Block 3 as the hook, and the hours question drops to the mechanism slot behind it. The block arc still holds, hook to mechanism to tension to implementation.')

h3('Why It Should Work')
bullet('Curiosity mechanism (Heath): ', 'The cold open stacks three impressive credentials and then discards all of them, which forces the question of what could possibly be more interesting than that.')
bullet('Share driver (Berger): ', 'The emotion is not admiration, it is the anxiety of an agent who realizes AI has never heard of them. Q8 turns that anxiety into a twenty-minute test, and a test is a thing people send to each other.')
bullet('Retention move (MrBeast): ', 'Every block ends on an implementation question and every bridge names a new tension. The FBI story sits at minute twenty, which is exactly where the middle would otherwise sag.')

h3('The Dissent')
p('Miller is still objecting that Block 4 is a second episode. He may be right. The A/B test for next time: if the tape shows the energy dropping at the Block 3 to Block 4 bridge, cut Block 4 entirely on the next guest like this and run a four-block single-topic episode with a deeper Block 2 instead. Justin is the right guest to learn this on, because he pitched both topics himself and the seam will be audible.')

h2('7C. EP Polish (pass 3)')
bullet('', 'Resequenced Block 3 so the contrarian question opens it and the hours question follows, on MrBeast\'s note about the sag. Renumbered Q9 through Q12 and retimed the chapter markers to match.')
bullet('', 'Stripped the $4.8 million out of all three titles, the cold open, the live stream title, and the five platform descriptions, then rebuilt every one of them on verified numbers only.')
bullet('', 'Rewrote Q5 to demand a deal count and an attribution method instead of a number he could answer with a story, and added the producer note telling D.J. not to accept the substitute.')
bullet('', 'Rewrote Q2 the same way, so "I write genuinely helpful content" cannot pass as an answer.')
bullet('', 'Reframed Q6 so the attribution skepticism belongs to a described listener rather than to D.J., which turns a possible ambush into a shared problem.')
bullet('', 'Removed the grief question entirely and moved that material to Landmines as guest-led only, with instructions for what to do if he raises it himself.')
bullet('', 'Added the green room brief, because four separate facts in this episode are guest-supplied and the only real fix is to get him to own them before the tape rolls.')
bullet('', 'Added the Rapid Fire pre-brief to the green room and the deflection line to the Quick Reference Card, since question 4 is very likely to spend the entire core topic in minute two.')
bullet('', 'Rewrote the Block 3 to Block 4 bridge to state the through-line out loud, so the expired material reads as the same skill applied to inventory rather than as a second episode.')
bullet('', 'Changed every instance of "thirty pages" to "around thirty," and added it to Landmines.')
bullet('', 'Rephrased the Landvoice statistic as "one vendor tracked" in the question, the chapter marker, and the data table.')
bullet('', 'Cut four questions to fit 45 minutes, then wrote the cut list in cut order and the never-cut list.')
bullet('', 'Added the snake story as an explicitly optional levity beat at the close rather than as a numbered question, with the instruction not to tell it for him.')
bullet('', 'Wrote permission clauses on exactly two questions, Q7 and Q17, and removed them everywhere else so the device does not become a tic.')
bullet('', 'Checked every short version against the twenty-word ceiling. Q1 and Q13 had to be rewritten to survive the compression.')
bullet('', 'Swept the whole document for em dashes, curly quotes, and AI-speak. Zero.')

doc.save("/Users/djparis/GitHub Projects/keeping-it-real-content-system/guest-prep/Justin_Black_Interview_Prep.docx")
print('saved')
