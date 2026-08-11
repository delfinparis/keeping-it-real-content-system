#!/usr/bin/env python
# Builds the KIRP interview prep .docx for Bobby Kerr, founder of LOC8 My Business
# Built to prompts/06_interview_prep.md v4: draft -> stress test + council -> EP polish
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)

for name, size, color in (('Heading 1', 18, 0x1F3864), ('Heading 2', 14, 0x2E5496), ('Heading 3', 11.5, 0x2E5496)):
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor((color >> 16) & 255, (color >> 8) & 255, color & 255)
    st.font.bold = True

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.8)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def p(text='', bold=False, italic=False, size=None, space=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space)
    r = par.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    return par


def rich(segments, style=None, space=6):
    par = doc.add_paragraph(style=style)
    par.paragraph_format.space_after = Pt(space)
    for seg in segments:
        text = seg[0]
        b = seg[1] if len(seg) > 1 else False
        i = seg[2] if len(seg) > 2 else False
        r = par.add_run(text)
        r.bold, r.italic = b, i
    return par


def bullet(text_bold, text_rest=''):
    par = doc.add_paragraph(style='List Bullet')
    par.paragraph_format.space_after = Pt(2)
    if text_bold:
        par.add_run(text_bold).bold = True
    if text_rest:
        par.add_run(text_rest)
    return par


def quoted(label, text):
    par = doc.add_paragraph(style='Intense Quote')
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.space_before = Pt(3)
    r = par.add_run(label + ' ')
    r.bold = True
    r.italic = True
    r2 = par.add_run(text)
    r2.italic = True
    return par


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            par = cells[i].paragraphs[0]
            run = par.add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def q(num, question, if_vague, reveals, serves, note=None):
    rich([(str(num) + '. ', True), (question, True)], space=2)
    quoted('If vague, ask:', if_vague)
    quoted('Ideal answer reveals:', reveals)
    quoted('Serves:', serves)
    if note:
        rich([('PRODUCER NOTE: ', True), (note, False, True)], space=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def bridge(text):
    rich([('BRIDGE TO NEXT BLOCK (read as written): ', True), ('"' + text + '"', False, True)])


# =====================================================================
# TITLE
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('KEEPING IT REAL PODCAST')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Interview Prep: Bobby Kerr')
r.bold = True
r.font.size = Pt(24)
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Founder, LOC8 My Business  |  Google Business Profile optimization for real estate  |  Prepared August 11, 2026  |  v2')
r.italic = True
r.font.size = Pt(10)

# =====================================================================
# PAGE 0 — WHAT CHANGED IN v2
# =====================================================================
h1('WHAT CHANGED IN v2')
p('A second stress test, a council re-run, and a full EP rewrite ran on v1. If you already read v1, these are the only things you need to re-learn.', italic=True, size=9.5)
bullet('The 83% stat is now verified, and a better number came out of the check. ', 'v1 told you to say 83% on air off a source I never actually opened. It holds up, but the sharper finding is Pew: when Google shows an AI answer, people click a cited source about 1% of the time. Q3 was rebuilt around it.')
bullet('Q3 now asks about more than Google. ', 'v1 titled the episode around AI search and then only interrogated Google. Gemini reads the profile directly. ChatGPT largely does not, it leans on Bing Places, Yelp, and Foursquare. That is in his own pitch and v1 buried it in the research brief.')
bullet('There is now a scripted guest intro. ', 'v1 referenced the intro four separate times in the landmines and never wrote one. That was the biggest hole in the document. See the next page.')
bullet('There is now a green-room brief. ', 'Three things have to be said before you record, including the one that stops Rapid Fire from burning his best story.')
bullet('Block 3 lost two minutes and Block 4 gained them. ', 'v1 gave the block with both never-cut questions the least room on the clock. That was backwards. See the council, which is also how the v1 dissent got resolved.')
bullet('The stat rule got precise. ', 'v1 said no guest-supplied numbers in the descriptions and then put one in the YouTube description. The rule now separates performance claims, which are banned, from volume claims, which are fine when attributed. See Landmine 5.')
bullet('Q10 moved onto the never-cut list. ', 'v1 called the Rapid Fire callbacks load-bearing and then left both of them cuttable. Contradiction fixed.')

doc.add_page_break()

# =====================================================================
# THE INTRO AND THE GREEN ROOM
# =====================================================================
h1('THE INTRO (read as written)')
p('This is the safe version. Every landmine in this document is already engineered out of it: no business list, no years, no city, no performance numbers, and no chance of introducing the other Bobby Kerr.', italic=True, size=9.5)
rich([('"My guest today is Bobby Kerr. He is the founder of LOC8 My Business, which is a marketing agency built around exactly one thing, the Google Business Profile that every real estate agent has and almost nobody has finished. He has rebuilt more than a thousand of them for agents and teams across the United States and Canada. Before this he spent years running businesses inside our industry, and I am going to let him tell you which ones. Bobby, welcome to Keeping It Real."', False, True)])
p('The line "I am going to let him tell you which ones" is doing real work. It gets his resume on tape in his own words without you asserting a version of it that his own published bios contradict. See Landmine 2.', italic=True, size=9.5)

h2('GREEN ROOM (say all three before you hit record)')
bullet('1. "How do you say your name, and are you Bobby or Robert on air?" ', 'Also confirm his pronouns. He uses he and him throughout his own materials but he has never said so in public.')
bullet('2. "Where are you sitting right now?" ', 'Solves Landmine 3 off air. His company address is Florida and most of his history is Kansas City, and you do not want to find that out live.')
bullet('3. THE IMPORTANT ONE. "I open with four rapid fire questions, including best and worst advice. Keep them short, because I am not going to follow up on them there. Both of your answers are too good to spend in minute one, so we are coming back to them later on purpose." ', 'This is the fix for the single biggest structural problem with this guest. Without it, he tells his two best stories in ninety seconds to a host who then changes the subject, and the callbacks at Q9 and Q10 land as repetition instead of as a return.')

doc.add_page_break()

# =====================================================================
# PAGE 1 — QUICK REFERENCE CARD
# =====================================================================
h1('PAGE 1: QUICK REFERENCE CARD')
p('Glance at this during the interview. Everything else is morning-of reading.', italic=True, size=9.5)

bullet('Name: ', 'Bobby Kerr. Pronouns not publicly stated; he/him used throughout his own materials. Confirm in the green room.')
bullet('Title: ', 'Founder, LOC8 My Business, a real estate marketing agency built entirely around Google Business Profile optimization. Business address on the site: 2237 SW 4th Ct, Cape Coral, FL 33991. Phone 239.268.3852.')
bullet('Calls himself: ', '"The Google Guy" for real estate pros. Also uses "Rockstar Realtor." Certified through Google\'s marketing certification program, which he has described publicly as a 240-hour program.')
bullet('Also does: ', 'Co-host of the Always Be Cool (ABC) Podcast with Darren Copeland of Summit Lending. Performs as Bob Jovi, a 1980s Bon Jovi tribute act. Listed as COO of The Shaun Ashley Team at RE/MAX Heritage in Blue Springs, Missouri. See Landmines 2 and 3 before you say any of this on air.')
bullet('What LOC8 sells: ', 'GBP Build and Rebuild from $499 one time. Done-For-You GBP Build from $1,250 one time, which includes 25 custom services written with local keywords. Authority Membership from $297 a month, which is 32 content pieces a month, 8 posts, 8 products or offers, 8 services. Also an AI search visibility service.')
bullet('The Authority Engine: ', 'His three-step system. 1) Profile audit and build. 2) Review acceleration. 3) Local authority growth, meaning monthly optimization, citations, and AI-ready content.')
bullet('Claimed stats: ', '1,000+ agent profiles optimized across the US and Canada. 160+ markets. LOC8 itself shows 248+ five-star reviews at a 4.9 average. Read Landmine 5 before you repeat any number.')
bullet('Media he claims: ', 'FOX 4 News, The Kansas City Star, WINK News, BBB accredited. Verified outside appearance: SpeakerFlow\'s Technically Speaking, season 3 episode 42, "Google Business Is More Relevant Than You Think."')
bullet('Free stuff he gives away: ', 'A 5-Step Local Search System checklist, a free SEO training video, a complete GBP guide, and a download called "From Forgotten to Found." He also teaches live webinars and workshops weekly.')
bullet('Find him: ', 'loc8mybusiness.com | thebobbykerr.com | LinkedIn /in/thebobbykerr | Facebook, Instagram and YouTube @thebobbykerr | Facebook @loc8mybiz | info@loc8mybiz.com')
bullet('KIR connection: ', 'FIRST TIME GUEST. No prior KIR appearance found.')
bullet('Guest type: ', 'C with a D overlay. Vendor first, educator second. Product cap applies: no more than 20% of the questions can be about what he sells, and the episode has to fully pay off for an agent who never gives him a dollar. Every product answer gets a "now give me the version I do myself on a Sunday night" follow-up.')
bullet('Episode length: ', 'NOT CONFIRMED. Everything here is built to 43 minutes, four blocks. If it is 30, drop Block 3 to two questions and cut Q12 and Q4. If it is 60, Q5 and Q13 are the two that reward more room.')

rich([('THE CORE TOPIC: ', True),
      ('What an agent actually changes on their Google Business Profile this month so that Google and the AI engines hand them over as the answer, field by field, in the order he would do it.',)])

rich([('THE ONE THING THAT MAKES THIS AN EPISODE AND NOT AN AD: ', True),
      ('Everything he says will be true and almost none of it will be specific unless you make it specific. "Optimize your profile" is not an answer. "Fill out your services" is not an answer. The answer is the words he would type into the box. Every single time he gives you a category, ask for the sentence. If you leave with a checklist and no example copy, the episode failed.', False, False)])

rich([('WATCH OUT FOR (this one will bite you in minute two): ', True),
      ('He already sent his best advice and his worst advice in the intake, and both of them are load-bearing for this episode. The standardized Rapid Fire will spend them in the first two minutes and you do not follow up on Rapid Fire. Pre-brief him in the green room, sentence 3 on the intro page, then after each of those two answers say "Love it, and we are coming back to that." Then come back to them on purpose: the worst-advice callback is Q9 and the best-advice callback is Q10. They are written out. Do not improvise the callbacks.', False, False)])

rich([('THE "I\'VE INTERVIEWED HUNDREDS" MOMENT (use once, Q15): ', True),
      ('"I have interviewed hundreds of agents on this show, and when Google comes up the answer is almost always the same. I claimed it years ago, it is fine, and anyway my business comes from referrals. That agent is listening right now and he is not wrong about his referrals. Talk to him directly."', False, False)])

rich([('LIVE STREAM TITLE (paste into Restream before you hit record): ', True),
      ('Bobby Kerr: How Agents Get Picked by Google and AI Search in 2026', False, True)])
p('65 characters. Backup: "Your Google Profile Is Your Front Door: Bobby Kerr on Winning AI Search" (71 characters). This does not have to match the published episode title. Pick that one after you hear the interview. See Section 5A.', italic=True, size=9.5)

rich([('OVERASKED, DO NOT ASK: ', True)])
bullet('', '"What is a Google Business Profile and why does it matter?" This is the opening question on every show he has been on. Your audience already knows what it is. Start at what is broken on theirs.')
bullet('', '"How do I get more reviews?" His published answer already exists: QR codes at events, and a nine-touch drip across email, text, and voicemail. Q13 and Q14 go at reviews from the angle nobody asks about, which is what the review has to say.')
bullet('', '"What is the biggest mistake agents make?" His answer is "an incomplete profile," which is a category, not a mistake. Q5 gets the same territory with a list you can actually use.')

doc.add_page_break()

h2('LANDMINES (read all seven before you record)')

rich([('1. THERE ARE TWO BOBBY KERRS AND THE OTHER ONE IS MORE FAMOUS. ', True),
      ('Search his name and the top results are frequently Bobby Kerr the Irish entrepreneur, the Insomnia Coffee founder, a Dragons\' Den investor, and the host of Down To Business on Newstalk. That is a completely different person. If any part of your intro came out of a search summary or an AI assist, check it. Do not say Dragons\' Den, do not say Insomnia Coffee, and do not say anything about Ireland.', False, False)])

rich([('2. HIS RESUME HAS TWO PUBLISHED VERSIONS AND THEY DO NOT MATCH. ', True),
      ('The bio he sent you says he started, scaled, and sold multiple real estate businesses: home inspections, mortgage, property and casualty insurance, and real estate sales. The bio he gave SpeakerFlow says his career runs casino gaming sales, middle school teaching, then real estate. His own personal site lists a home inspection company and an investment company, and describes him as co-leader of The Shaun Ashley Team at RE/MAX Heritage; a BBB listing for that team names him as COO. The mortgage brokerage and the insurance agency are in his intake and nowhere public that I could find.', False, False)])
rich([('   WHAT YOU DO: ', True),
      ('Do not read the list as verified and do not say "sold." Do not attach a year to anything. Q1 is deliberately built on a client result rather than his origin story so you never have to. If you want the resume on tape, let him say it: "Before LOC8 you ran a stack of businesses in this industry. Name them for me." Then take what he says at face value and move.', False, False)])

rich([('3. TWO HOME BASES. ', True),
      ('LOC8 lists a Cape Coral, Florida address. His real estate team, his tribute act, and his Kansas City Star mention are all Kansas City, and his old Twitter handle is KC_Bob. Both can be true and it is not the story. Do not geolocate him in the intro. If it comes up, ask "where are you sitting right now" and let him answer it.', False, False)])

rich([('4. GOOGLE KILLED THE Q&A SECTION AND HE STILL TEACHES IT. ', True),
      ('His topic pitch lists Q&A as one of the fields to fill on a complete profile. Google discontinued the Q&A API on November 3, 2025, began phasing the public Q&A section out on December 3, 2025, and new profiles are now created without it. The replacement is the Gemini-powered "Ask about this place" answer box, which generates answers in real time from the profile, the reviews, the photos, and the website. This is the single most interesting question in the episode and also the easiest one to blow. Q7 is written to land it as two practitioners comparing notes, not as a correction. If he says Q&A is still on profiles, do not argue and do not fact-check him on air. Your line is "on some it still is, on new ones it is gone, so what fills that job now," and you move. The listener hears the update either way and he keeps his dignity.', False, False)])

rich([('5. THE NUMBERS DO NOT AGREE WITH EACH OTHER. ', True),
      ('He has publicly used "1065% more likely to get calls and clicks" for profiles with 100+ photos, and his site uses "10x more client calls," "3 to 5x more profile views," and a client who got $1.2M in listings in two weeks. None of those have a traceable primary source. Separately, his SpeakerFlow bio says 1,200+ five-star reviews across his own Kansas City businesses while LOC8\'s site says 248+ reviews at 4.9. Those are two different things and stacking them sounds like a contradiction.', False, False)])
rich([('   THE RULE, and v2 made it precise because v1 was sloppy about it: ', True),
      ('Performance claims are banned from your mouth. That is 1065%, 10x, 3 to 5x, and the $1.2M. Let him say them, do not repeat them back, never say two in one sentence, and keep all of them out of the intro, the titles, and the descriptions. Volume claims are fine when they are attributed to him: more than a thousand profiles rebuilt, 160+ markets, US and Canada. Those are consistent across every source of his and they describe what he did, not what it produced. That is why the intro and the YouTube description both carry the thousand. The only research number you assert as fact is in Q3, and it is not his.', False, False)])

rich([('6. SPONSOR EXPOSURE. ', True),
      ('His whole thesis is that the free Google page beats the paid website, and that agents who outsource content to third-party posting companies get buried. Real Geeks sells agent websites and lead generation. Courted sells agent data. Let him make his case in full, but do not co-sign it, do not add your own shot at the category, and do not let a clip go out that reads as the show saying websites are dead. The frame that keeps everybody whole is written into Q2: the profile is the front door, the website is the house, and most agents built the house and boarded up the door. Q2 also forces him to name what the website still does that the profile cannot, which is the sentence that makes the clip safe.', False, False)])

rich([('7. SOMEONE DIED IN ONE OF HIS STORIES. ', True),
      ('His intake includes a closing where the seller passed away at the table, with a loan officer named Alan in the room. A real person died in front of a room full of professionals who are all still working. HARD RULES: this is not a question in the run of show and you do not go fishing for it. If he brings it up, let him tell it, do not ask what happened to the estate, do not ask for the city, the year, the brokerage, or the title company, and do not laugh into it. Land it with one flat sentence and move: "That will stay with everybody who was in that room." If he starts identifying anyone, your line is "let us keep everybody anonymous on this one." Never cut it as a standalone clip with a caption.', False, False)])

rich([('TONE READ, DO THIS IN THE FIRST THREE MINUTES: ', True),
      ('He is a performer. He fronts a Bon Jovi tribute band, he hosts his own show, and his materials run hot on energy. That means two risks. One, he will answer in slogans, because slogans are what he does on stage and on webinars every week. Two, he will be very comfortable talking, so a soft question will cost you four minutes. The counter for both is the same and it is the only interviewing note that matters on this guest: after every answer, ask for the words. Not the principle, the words. "Say the actual line." "Type it out loud for me." If you do that six times in this episode, it is a great episode. If you do it zero times, it is a webinar.', False, False)])

doc.add_page_break()

# =====================================================================
# PAGE 2 — EPISODE FRAMEWORK
# =====================================================================
h1('PAGE 2: EPISODE FRAMEWORK')

h2('2A. Title Options')
table(['#', 'Title', 'Ships To', 'Why It Works'],
      [['1', 'Ranking Is Over. Bobby Kerr on Getting Handed to the Buyer as the Answer.',
        'Podcast feed, social cuts',
        'Kills a belief in four words, then names the replacement. No number to defend, no jargon to decode, and it works on an agent who has never heard of him.'],
       ['2', 'Google Cites You and Nobody Clicks. Bobby Kerr on Getting Picked Inside the Answer.',
        'YouTube',
        'Built on the verified Pew finding that a cited source gets clicked about 1% of the time. The threat is in the first four words and the escape is withheld. v1 used the 83% version of this, which is true but is the number every SEO account already posted.'],
       ['3', 'The Free Google Page You Claimed and Forgot Is Now Your Front Door (Bobby Kerr)',
        'Hold for a clip caption',
        'Best emotional hit of the three because it accuses the listener gently, but "front door" is his metaphor and it needs the episode around it to mean anything.']],
      widths=[0.3, 2.9, 1.1, 2.6])

h2('2B. Cold Open Hook')
p('Read before the ads.', italic=True, size=9.5)
rich([('"My guest today says the most valuable piece of real estate you own is not a listing, it is a free page on Google you probably have not opened since the day you claimed it. He has rebuilt more than a thousand of them for agents, and he says the ones winning right now are not ranking higher, they are getting handed to the buyer as the answer. We are going to talk about that today. Stay tuned."', False, True)])

h2('2C. Episode Arc')
rich([('Core topic: ', True), ('What an agent actually changes on their Google Business Profile this month so that Google and the AI engines hand them over as the answer.',)])
rich([('Why this topic and not the others: ', True),
      ('He could carry a whole episode on AI search theory, and it would be the fourth one your audience heard this quarter. He is the only guest you will book who has personally rebuilt a thousand of these profiles, so the thing only he can give you is the field-by-field build. Theory is free everywhere. The words he types in the box are not.',)])
rich([('The four angles: ', True),
      ('Block 1 is why this surface and not the website. Block 2 is what complete actually means, field by field, including what to do now that Q&A is gone and how not to get suspended. Block 3 is why generic content buries you and what hyperlocal looks like at the sentence level. Block 4 is reviews, specifically why the text beats the star count once an AI engine is doing the reading.',)])

doc.add_page_break()

# =====================================================================
# SECTION 3 — INTERVIEW QUESTIONS
# =====================================================================
h1('SECTION 3: INTERVIEW QUESTIONS')

h2('Structure (43 minute target)')
table(['Segment', 'Time', 'Purpose'],
      [['Rapid Fire', '0:00 to 2:00 post-intro', '4 standardized questions, no follow-ups. Read the Watch Out For note first.'],
       ['Block 1: The Front Door', '2:00 to 12:00', 'Why this surface, and why "I claimed it years ago" is the whole problem.'],
       ['Block 2: What Complete Actually Means', '12:00 to 22:00', 'The build, field by field. The Q&A change. The suspension rules.'],
       ['Block 3: Commodity vs Non-Commodity', '22:00 to 30:00', 'Why syndicated content buries a profile and what hyperlocal looks like in a sentence.'],
       ['Block 4: Reviews and Being the Answer', '30:00 to 40:00', 'Review text over star count, and the exact ask.'],
       ['The Close', '40:00 to 43:00', 'Homework, his resource, where to find him.']],
      widths=[1.9, 1.6, 3.4])

h3('RAPID FIRE (standardized, read as written, no follow-ups)')
p('1. Best real estate advice you have ever received?', bold=True, space=2)
p('2. Worst real estate advice you have ever received?', bold=True, space=2)
p('3. One tool or app you cannot run your business without?', bold=True, space=2)
p('4. What would surprise people most about your day to day?', bold=True, space=2)
rich([('PRODUCER NOTE: ', True),
      ('He pre-sent 1 and 2 and they are both central to this episode. After each, say "Love it, and we are coming back to that," then move. Do not follow up here. The callbacks are Q9 and Q10 and they are written out. For reference so you recognize them coming: his best advice is a recruiting coach telling him to stop advertising that he owned three companies, and his worst advice is "you have to be everywhere."', False, True)])

doc.add_page_break()

h2('BLOCK 1: THE FRONT DOOR (2:00 to 12:00)')
p('Audience note: every segment, but aimed at the solo producing agent who has a website they paid for and a Google profile they have not opened in two years.', italic=True, size=9.5)
p('Arc: the client result, the hierarchy, what changed in search, the first field.', italic=True, size=9.5)

q(1,
  "You have a client in Annapolis who says she picked up a brand new listing the first week after you rebuilt her profile, and a second one the week after that. I want the boring version of that story, not the highlight. What was actually on that profile before you touched it, and what was on it after?",
  "Name three fields that were empty on Monday and full on Friday.",
  "Whether the result came from a mechanism or from luck, and it hands the listener a before and after picture in the first two minutes.",
  "Individual agents, new agents.")

q(2,
  "Most of the agents listening spent real money on a website and almost nothing on the free page Google handed them. You are telling them they built the house and boarded up the front door. Make that case, and then be specific about what the website still does that the profile cannot.",
  "If an agent has 500 dollars and one weekend, which one gets it, and why?",
  "A defensible hierarchy instead of a pitch, and a sentence that keeps the clip from reading as the show saying websites are dead.",
  "All segments.",
  "This is your sponsor protection. Do not let him skip the second half of the question. If he does not name something the website still does, ask it again as its own question.")

q(3,
  "Here is the number that made me want to have this conversation. Pew looked at what people actually do when Google shows them an AI answer, and they click one of the sources it cites about one percent of the time. One percent. So getting cited is not the win anymore, being the sentence is. What does that change about what an agent puts on the profile?",
  "Walk me through it screen by screen. A buyer opens their phone, types the thing, and then what happens between there and my phone ringing?",
  "The mechanism behind being the answer instead of ranking, which is the promise in the title.",
  "All segments.",
  "SECOND HALF, ask it as its own question and do not skip it: your pitch says Gemini reads the profile and ChatGPT reads the internet's opinion of you. Those are two different jobs. From what I have seen, ChatGPT leans on Bing Places, Yelp, and Foursquare more than it touches a Google profile at all. So which engine reads what, and what does an agent do about the ones that are not Google? This is the honest version of his own claim and it is the difference between an AI search episode and a Google episode. If he waves it off, that tells you something too.")

q(4,
  "Somebody is listening to this on a treadmill right now and has not opened their profile in two years. When they get home tonight, what is the one field they open, and what do they type into it?",
  "Give me the field name and then say the example sentence out loud.",
  "The first domino, and the first time in the episode he gives you actual copy instead of a category.",
  "Individual agents, new agents.")

bridge("So the profile is the front door and most agents have it boarded up. Let's open it. I want to go field by field through what a finished one actually looks like, because I think most people hear complete and picture the hours being right.")

doc.add_page_break()

h2('BLOCK 2: WHAT COMPLETE ACTUALLY MEANS (12:00 to 22:00)')
p('Audience note: this is the value-density block. It is the one an agent screenshots and the one a team lead sends to twelve people.', italic=True, size=9.5)
p('Arc: the list in order, the field everyone fakes, the thing that changed, the thing that gets you suspended.', italic=True, size=9.5)

q(5,
  "Define complete for me the way Google defines it, not the way an agent defines it. If you sat down at my computer with my profile open, what is the list you work down, in order?",
  "Rank the top five by impact, and tell me which one is blank most often.",
  "The actual checklist. This is the single most stealable answer in the episode and everything after it is a variation on it.",
  "All segments.",
  "NEVER CUT. If he gives you eleven items in thirty seconds, stop him and say give me the top three again slower.")

q(6,
  "Services is the field I hear you talk about more than any other, and you write twenty-five of them for a client instead of the four an agent would write. Give me a real one. For an agent who works Logan Square here in Chicago, what does a services entry look like, word for word?",
  "Say the actual line. Not the category, the line.",
  "The difference between commodity and non-commodity at the field level, in language a listener can copy tonight.",
  "Individual agents.",
  "He sells a 25-services build. That is fine, let him reference it once. Then ask: how many can an agent write themselves in an hour, and where do they get the neighborhood words from?")

q(7,
  "Here is where I want your honest read, because this stuff moves under all of us. Google pulled the public Q and A section off profiles starting at the end of last year and pushed everybody toward the AI answer box on the listing instead. Q and A has been part of the build you teach. What survived that change, what did not, and how do you feed that AI box now that you cannot seed the questions yourself?",
  "If a brand new profile does not have Q and A at all, what field takes over that job? Name it.",
  "Whether his system is current or a checklist from three years ago. This is the credibility moment of the episode, and if he answers it well it is also the most forwardable minute in it.",
  "All segments.",
  "NOT A GOTCHA, and it will read as one if you deliver it flat. Ask it leaning in, as two people comparing notes. If he says Q and A is still there, do not argue and do not correct him on air. Say: on some profiles it still is, on new ones it is gone, so what do you do there. Then move.")

q(8,
  "Last thing on the build, and it is the one that actually scares people. Agents get suspended. The name rules, the brokerage address you do not personally sit at, the service area setup. What is the safe configuration for a solo agent at a big brokerage, and what is the thing that gets people taken down that they never see coming?",
  "Give me the exact business name format for an agent named Maria Lopez at a 500 agent brokerage. Type it out loud.",
  "The compliance layer almost nobody teaches, and the reason a broker owner keeps listening.",
  "Individual agents, new agents, broker owners.",
  "So you can tell whether the answer is right: Google treats agents as individual practitioners, which means the profile name is the person's name and nothing else. Maria Lopez. Not Maria Lopez Realtor, not Maria Lopez Real Estate, not Maria Lopez at Anywhere Realty. Adding any of that is a suspension trigger. The other one is the address. An agent who lists a brokerage address they do not personally staff gets flagged, and the setup most solo agents actually qualify for is a service area business with the street address hidden. If he says all of that, he is the real thing. If he says put your keyword in the name because it helps you rank, that is against Google's own guidelines, and your follow-up is: how many of your clients have been suspended doing that?")

bridge("Okay, so it's built and it's not going to get pulled down. Now it's empty, and it stays empty unless somebody feeds it. That's the part I want to spend the most time on, because I think this is where most agents outsource themselves into invisibility.")

doc.add_page_break()

h2('BLOCK 3: COMMODITY VS NON-COMMODITY (22:00 to 30:00)')
p('Audience note: this block carries both Rapid Fire callbacks. It is also the most contestable claim in the episode, so it needs the most evidence. v2 cut it from ten minutes to eight and gave them to Block 4. Four questions in eight minutes means you cannot let him run here, and Q12 is the one that goes if he does.', italic=True, size=9.5)
p('Arc: the callback, the proof, the specialty problem, the calendar, the broker version.', italic=True, size=9.5)

q(9,
  "I want to come back to your worst advice, because we did not finish it. You said the be-everywhere advice pushes agents into hiring a company to post for them, and those companies run the same content across thousands of profiles. I want to know if Google can actually tell. Not in theory. What have you watched happen to a profile running syndicated content next to one that is not?",
  "Two agents, same market, one on a content service and one writing their own. What do the profile views look like at ninety days?",
  "Whether Google knows is an observation or a belief. This is his biggest claim and the one an SEO person in the audience will push back on, so make him show his work.",
  "All segments.",
  "If he cannot produce a before and after, do not rescue him. Say: so it is a judgment call more than a measured thing. Let the answer be what it is. That honesty is the clip, and it costs you nothing.")

q(10,
  "Earlier you told me a recruiting coach made you stop advertising that you owned three companies, and that you fought him on it and he was right. Same principle down here. What does pick one thing look like on a Google profile for an agent who genuinely does buyers, sellers, rentals, and investors, and who cannot afford to turn any of it away?",
  "Does the specialty live in the services, the posts, or the reviews? Pick one and tell me why the other two are wrong.",
  "How to be specific without amputating the business, which is the objection every agent has to niching and the reason most of them never do it.",
  "Individual agents, team leaders.")

q(11,
  "Give me the calendar. How often does something new have to hit that profile before Google treats it as alive, what are the formats, and how long does it actually take an agent to do it themselves on a Sunday night?",
  "Posts per month, photos per month, videos per month. Give me the numbers.",
  "The recency mechanism, plus an honest read on whether do-it-yourself is realistic or whether this is a thing you hire out.",
  "All segments.",
  "This is where the membership pitch wants to come out. Product cap. Let him name it once, then say: now give me the version I do myself with no budget.")

q(12,
  "Flip this for a second. I recruit agents for a living, so I am listening to this as a broker. If I have 800 agents and most of their profiles are blank or wrong, is that my problem or theirs? And is there a brokerage-level version of this that does not step on the agents' own profiles?",
  "Should the brokerage profile and the agent profile say the same thing or deliberately different things?",
  "The team and brokerage architecture, which nobody covers and which is the reason a broker owner shares the episode.",
  "Broker owners, team leaders.")

bridge("So it's built, it's fed, and it's specific. The last piece is the one agents think they already understand, and I don't think they do. Let's talk about reviews, because you're telling me the number I've been chasing for years is the wrong number.")

doc.add_page_break()

h2('BLOCK 4: REVIEWS AND BEING THE ANSWER (30:00 to 40:00)')
p('Audience note: the most immediately actionable block in the episode, the one with the highest chance of producing a screenshot, and the only block holding two never-cut questions. v1 gave it the tightest clock in the episode, which was backwards. It now gets ten minutes.', italic=True, size=9.5)
p('Arc: the wrong metric, the exact ask, the objection, the honest timeline.', italic=True, size=9.5)

q(13,
  "You have said the text inside a review matters more than the star count. Every agent I know is chasing five stars and counting them. What is an AI engine actually doing with the words in a review, and what does a review have to say to be worth more than one more five-star Bobby was great?",
  "Write me the review you would want, out loud, for an agent who sells two-bedroom condos in one neighborhood.",
  "The most stealable tactic in the episode, and the one that reframes a metric the whole audience already tracks.",
  "All segments.",
  "NEVER CUT.")

q(14,
  "Then the hard part is asking. Most agents send please leave me a review and get great to work with. What do you say to a client, word for word, that gets them to name the neighborhood, the property type, and the problem you solved, without it feeling like you wrote it for them?",
  "Give me the text message. The actual text message, the way it would show up on their phone.",
  "A script listeners will screenshot before the episode ends.",
  "Individual agents, new agents.",
  "NEVER CUT. If he gives you a philosophy, ask again. You want a message with a send button on it.")

q(15,
  "I have interviewed hundreds of agents on this show, and when Google comes up the answer is almost always the same. I claimed it years ago, it is fine, and anyway my business comes from referrals. That agent is listening right now and he is not wrong about his referrals. Talk to him directly. Why does he care about this?",
  "He does thirty deals a year off his sphere. What does the profile do for him that his sphere does not already do?",
  "Whether this is universal or a fix for agents without a network. It names the biggest objection in the room instead of dodging it, and the answer decides whether the episode converts.",
  "All segments.",
  "NEVER CUT. This is your one hundreds moment. Do not use the phrase anywhere else.")

q(16,
  "Last one. Somebody does everything you just said. Realistically, when do they see the first thing they can point at and say that came from this, and what does that first thing usually look like?",
  "Weeks or months? And is the first signal a call, a message, or a direction request?",
  "An honest timeline, which protects him from the too-good-to-be-true read and protects your audience from quitting at week three.",
  "All segments.")

h3('THE HUMAN LANDING (optional, right before the close)')
q('17a',
  "One thing I have to ask you about before I let you go. You front a Bon Jovi tribute band. You go by Bob Jovi. Everything you just told me about being specific and not looking like everybody else is not a marketing theory for you, it is how you have run your own name for years. What did standing on a stage teach you about being memorable that no marketing book ever did?",
  "What is the moment you knew the name was working?",
  "The vulnerable beat that earns the follow, and the one moment in the episode that is about him instead of about a field on a form.",
  "All segments.",
  "ALTERNATE if the energy is low or you are tight: his home inspection story. Empty house, top closet shelf, a bank envelope with a few thousand dollars in it that only a tall guy would ever have seen, handed straight to the buyer's agent. The button is that the listing agent owned a home inspection company himself and had walked that house. Ask it as: what is the strangest thing you ever found in an empty house?")

h2('THE CLOSE')
rich([('HOMEWORK (read verbatim): ', True),
      ('"Here is what I want you to do before the next episode. Open your Google Business Profile, go to the services section, and write ten services. Not four. Ten. Each one named the way a client would say it out loud, and each one with a neighborhood in it. If you have got four in there right now, you are six short. Not next month. This week."', False, True)])
p('Why this one: it is free, it takes under thirty minutes, and done is a number you can count. v1 said "write out every service you offer, and if you have fewer than ten you have homework," which left the listener unsure whether done meant counting or writing. Ten is the assignment. If Q6 produces a better field than services, swap the field, but keep the number. The number is what makes it done.', italic=True, size=9.5)

rich([('GUEST CLOSE: ', True), ('"Where can people find you, and what do you want them to do first?" He has several free giveaways: the 5-Step Local Search System checklist, a free training video, a full GBP guide, and the From Forgotten to Found download. Ask him to name ONE. Four offers is the same as none, and it is also the exact mistake this episode is about.',)])

h3('If you are running long, cut these first')
bullet('1. Q12, the broker flip. ', 'It is the furthest from your largest audience segment, and Block 3 still ends on an implementation question without it. Block 3 is eight minutes now, so assume you are cutting this one.')
bullet('2. Q17a, the human landing. ', 'Warmest thing in the episode, but the close survives without it and the homework does not.')
bullet('3. Q4. ', 'If Q2 already produced the first field, this is the same answer twice. The homework covers it anyway.')
bullet('4. Q11, the calendar. ', 'Only if Q9 already produced the cadence, which happens more often than you would think, because he cannot explain why syndicated content fails without describing what feeding a profile properly looks like.')
p('Q3 came off this list in v2. It is now carrying the verified research and the multi-engine question, which means it is the only place in the episode where the phrase AI search gets tested instead of assumed.', italic=True, size=9.5)

h3('Never cut')
bullet('Q5. ', 'The list in order. Without it there is no episode, only opinions.')
bullet('Q7. ', 'The Q and A change. This is the only minute of this episode that could not have been recorded two years ago.')
bullet('Q10. ', 'The best-advice callback. Added in v2. Rapid Fire spends his best story in minute one, and this is the only thing in the run of show that buys it back. v1 called the callbacks load-bearing and then left both of them cuttable.')
bullet('Q13. ', 'Review text over star count. The reframe.')
bullet('Q14. ', 'The exact ask. The screenshot.')
bullet('Q15. ', 'The objection said out loud. This is the one that decides whether anybody acts.')

doc.add_page_break()

# =====================================================================
# SECTION 4 — RESEARCH BRIEF
# =====================================================================
h1('SECTION 4: RESEARCH BRIEF')
p('Morning-of reading. Not for during the interview.', italic=True, size=9.5)

h2('4A. Background')
p('Kerr is the founder of LOC8 My Business, a marketing agency built entirely around Google Business Profile optimization for real estate professionals, run from a Cape Coral, Florida address. His own materials are consistent on one thing and inconsistent on almost everything else: he has run a stack of small businesses before this, he collected an unusual number of Google reviews doing it, he went and got Google certified, and he turned that into a service business for agents. The parts that do not line up across sources are which businesses, in what order, and whether he still owns any of them. See Landmine 2. He is also a working performer, which is relevant rather than trivia: the tribute act and the podcast are the proof case for the thing he teaches, which is that being specific and unmistakable beats being everywhere.')

h2('4B. Career Timeline (verified entries only)')
table(['When', 'Role / Company', 'Source and note'],
      [['Undated', 'Founder, home inspection company', 'His own site. Intake says he later sold it. No public record of the sale.'],
       ['Undated', 'Founder, investment company', 'His own site.'],
       ['Current or recent', 'COO and co-leader, The Shaun Ashley Team, RE/MAX Heritage, Blue Springs MO', 'BBB listing plus his own site. Conflicts with the intake line that he sold his sales business. Do not say either version.'],
       ['Undated', 'Google marketing certification, described by him as a 240-hour program', 'Guest-supplied via SpeakerFlow.'],
       ['Ongoing', 'Co-host, Always Be Cool (ABC) Podcast, with Darren Copeland of Summit Lending', 'Apple, Spotify, Amazon, YouTube. 100+ episodes. Start year not verified, so do not say one.'],
       ['Ongoing', 'Founder, LOC8 My Business, Cape Coral FL', 'Company site. Founding year not published.'],
       ['Ongoing', 'Performs as Bob Jovi, 1980s Bon Jovi tribute act', 'His own site.']],
      widths=[1.2, 2.6, 3.1])
p('Mortgage brokerage and property and casualty insurance agency: intake only. Both are load-bearing in his best-advice story. Let him tell it. Do not introduce them as fact.', italic=True, size=9.5)

h2('4C. What Makes Him Interesting to This Audience')
bullet('He has done the reps, not just the research. ', '1,000+ agent profiles rebuilt across 160+ markets in the US and Canada. That means he can answer at the field level instead of the strategy level, which is exactly the altitude your audience wants and almost never gets from a marketing guest.')
bullet('He is on the right side of a real change. ', 'Whitespark\'s 2026 local search survey puts Google Business Profile signals at roughly a third of local pack ranking, more than reviews, on-page SEO, and backlinks combined, and it added AI search visibility as its own ranking category this year. This is not a trend piece, it is a scoreboard change.')
bullet('The surface is free and almost nobody has finished it. ', 'Every other marketing episode you run ends with "and it costs this much." This one ends with a field an agent fills in tonight for nothing. That is rare and it makes the homework land.')
bullet('He ran the businesses his clients run. ', 'Inspections, lending, insurance, sales. Whatever the exact list turns out to be, he is not a marketer who discovered real estate. He talks like somebody who has been on a closing.')
bullet('He is the proof of his own thesis. ', 'A guy who calls himself the Google Guy and performs as Bob Jovi is not going to get confused with any other vendor. Q17a turns that into the human beat of the episode instead of a fun fact.')

h2('4D. Key Data Points')
table(['Stat', 'Source', 'Confidence'],
      [['1,000+ agent profiles optimized, 160+ markets, US and Canada', 'LOC8 site, matches his intake', 'Medium, self-reported but consistent'],
       ['248+ five-star reviews, 4.9 average on LOC8', 'LOC8 site', 'Medium, checkable live'],
       ['1,200+ five-star reviews across his own Kansas City businesses', 'SpeakerFlow bio', 'Medium, and a different claim than the one above'],
       ['Pricing: $499 build, $1,250 done-for-you, $297 a month membership', 'LOC8 site', 'High'],
       ['10x more client calls, 3 to 5x more profile views', 'LOC8 site', 'Unverified marketing claim. Do not repeat.'],
       ['$1.2M in listings within two weeks for one client', 'LOC8 site', 'Unverified. The Annapolis two-listings version in Q1 is the safer telling.'],
       ['1065% more likely to get calls and clicks with 100+ photos', 'Said by him on SpeakerFlow', 'Unverified, widely circulated with no primary source. Do not echo it.'],
       ['People click a cited source about 1% of the time when an AI Overview is shown, and click through on 8% of those searches versus 15% without one', 'Pew Research behavioral tracking', 'High. Verified in the v2 pass. This is the number you say on air, and it is now the basis for Q3.'],
       ['83% of searches with an AI Overview end with no click, 93% inside AI Mode', 'Seer Interactive; AI Mode figure via Semrush, Sept 2025', 'Medium to high. Safe to say, but the Pew number is sharper and lands harder.'],
       ['68% of all US Google searches end without a click', 'SparkToro and Similarweb, 2026', 'High. Useful for context if he asks where the number comes from.'],
       ['GBP signals are roughly 32% of local pack ranking', 'Whitespark 2026 Local Search Ranking Factors', 'High'],
       ['Public Q and A discontinued: API Nov 3 2025, public section phasing out from Dec 3 2025', 'Google support documentation and industry coverage', 'High. This is the basis for Q7.'],
       ['Gemini grounds local answers in Google Maps data across 250M+ verified places', 'Google, 2026', 'High'],
       ['ChatGPT local answers lean on Bing Places, Yelp, and Foursquare more than on GBP directly', 'Multiple local search analyses, 2026', 'Medium. Useful if he overclaims about ChatGPT reading his GBP. Do not correct him, just ask which engines read what.']],
      widths=[2.9, 2.4, 1.6])

h2('4E. Previous Media Appearances')
bullet('SpeakerFlow, Technically Speaking, season 3 episode 42, "Google Business Is More Relevant Than You Think." ', 'His fullest published interview. Covered: Google Business as a social platform rather than a review box, photos and video, posts, offers, events, Q and A, QR codes at live events, and a nine-touch review drip across email, text, and voicemail. Note that this appearance predates the Q and A removal, which is why Q7 exists.')
bullet('Always Be Cool (ABC) Podcast, his own show with Darren Copeland. ', '100+ episodes, Kansas City business, sports, and real estate guests, including a former Royals player. Tagline is Always Be Cool equals live by the golden rule. Worth one sentence of acknowledgment, not a segment.')
bullet('Claimed but not independently verified: FOX 4 News, The Kansas City Star, WINK News. ', 'All three appear on his site with no links. Do not read them in the intro as verified. "He has been featured in local press in Kansas City and Florida" is safe if you want the line at all.')
bullet('Weekly live webinars and workshops on Google-first visibility. ', 'Guest-supplied. This is why he will answer in trained slogans. See the tone read.')

h2('4F. Audience Relevance')
table(['Segment', 'What they get from this episode'],
      [['Individual agents', 'A free surface they already own, a field-by-field build, example copy they can type tonight, and a review ask with a send button.'],
       ['Team leaders', 'How to make a team specific without amputating the business, and how team and individual profiles should differ.'],
       ['Broker-owners', 'Whether 800 blank agent profiles are the brokerage\'s problem, and the suspension rules that keep agents from getting taken down at scale.'],
       ['New agents', 'The cheapest visibility available to somebody with no database, plus the name and address rules before they set the profile up wrong.']],
      widths=[1.5, 5.4])

doc.add_page_break()

# =====================================================================
# SECTION 5 — LIVE STREAM TITLE, DESCRIPTIONS & HASHTAGS
# =====================================================================
h1('SECTION 5: LIVE STREAM TITLE, DESCRIPTIONS & HASHTAGS')

h2('5A. Live Stream Title')
rich([('Live stream title: ', True), ('Bobby Kerr: How Agents Get Picked by Google and AI Search in 2026', False, True), (' (65 characters)',)])
rich([('Backup: ', True), ('Your Google Profile Is Your Front Door: Bobby Kerr on Winning AI Search', False, True), (' (71 characters)',)])
p('Both work cold, both use only verified research, and neither carries a number he would have to defend. This does not have to match the published episode title. Pick that one from Section 2A after you hear the interview.', italic=True, size=9.5)

h3('Facebook Live')
p('Bobby Kerr has rebuilt Google Business Profiles for more than a thousand real estate agents, and he says the ones winning right now are not the ones ranking higher. They are the ones Google and the AI engines hand over as the answer. We are going through the profile field by field, live. Drop your questions in the comments!')

h3('Instagram Live')
p('The free Google page you claimed and forgot is now your front door. Bobby Kerr on how agents get picked. #KeepingItReal #RealEstateAgent #GoogleBusinessProfile')

h3('TikTok Live')
p('Your Google profile is doing more for your business than your website, and you have not opened it in two years. #realtor #realestateagent #localseo')

h3('YouTube Live')
p('Bobby Kerr, founder of LOC8 My Business, has optimized Google Business Profiles for more than 1,000 real estate agents across the US and Canada. On this episode of the Keeping It Real Podcast, he walks D.J. Paris through what a complete profile actually looks like field by field, why the text inside a review beats the star count, what changed when Google removed the Q and A section, and which AI engines actually read a Google profile versus which ones do not. Google Business Profile optimization, local SEO, and AI search visibility for realtors.')

h3('LinkedIn Live')
p('Search is changing faster than most agents have changed their marketing. Bobby Kerr, founder of LOC8 My Business, joins D.J. Paris to break down how a Google Business Profile now feeds both the local map pack and AI answer engines. We also get into the brokerage question: what a broker does about hundreds of agent profiles that are blank, wrong, or one policy violation away from suspension.')

h3('Hashtag Sets')
bullet('Universal: ', '#KeepingItReal #RealEstatePodcast #DJParis #RealtorLife #RealEstateAgent')
bullet('Episode-specific: ', '#GoogleBusinessProfile #LocalSEO #AISearch #RealEstateMarketing #MapPack #GBP #AnswerEngineOptimization')
bullet('Guest tags: ', '@thebobbykerr on Facebook, Instagram and YouTube. LinkedIn /in/thebobbykerr. Company page @loc8mybiz on Facebook. Sites: loc8mybusiness.com and thebobbykerr.com. He also has an X account, which we do not post to.')

doc.add_page_break()

# =====================================================================
# SECTION 6 — YOUTUBE CHAPTER MARKERS
# =====================================================================
h1('SECTION 6: YOUTUBE CHAPTER MARKERS')
p('Estimates. Adjust after recording. Each title is written to be searchable on its own. Retimed in v2 for the Block 3 and Block 4 rebalance. If you use the cut list, the 29:00 and 40:00 chapters are the two that disappear, so delete those rows rather than stretching the ones around them.', italic=True, size=9.5)
table(['Timestamp', 'Chapter Title'],
      [['0:00', 'Ranking Is Not the Game Anymore. Being the Answer Is.'],
       ['2:00', 'Rapid Fire: Best and Worst Advice Bobby Kerr Ever Got'],
       ['4:00', 'The Agent Who Got Two Listings Two Weeks After Fixing Her Google Profile'],
       ['7:00', 'Google Business Profile vs Your Website: Where the 500 Dollars Should Go'],
       ['9:30', 'People Click a Cited Source 1% of the Time. So How Do Agents Get Chosen?'],
       ['11:30', 'The One Field to Open Tonight'],
       ['12:00', 'What a Complete Google Business Profile Actually Looks Like, In Order'],
       ['15:30', 'How to Write Services for Your Neighborhood, Word for Word'],
       ['18:00', 'Google Removed the Q and A Section. What Feeds the AI Answer Box Now?'],
       ['20:00', 'How Real Estate Agents Get Their Google Profile Suspended'],
       ['22:00', 'Can Google Tell When Your Content Was Copied to a Thousand Other Agents?'],
       ['25:00', 'How to Pick One Specialty Without Turning Away Business'],
       ['27:30', 'The Posting Calendar: How Often a Profile Has to Be Fed'],
       ['29:00', '800 Agents, 800 Blank Profiles. Whose Problem Is That?'],
       ['30:00', 'Why the Words in a Review Beat the Star Count'],
       ['33:30', 'The Exact Text Message That Gets a Client to Write a Useful Review'],
       ['36:30', 'For the Agent Who Says My Business Comes From Referrals Anyway'],
       ['38:30', 'How Long Before Any of This Actually Shows Up'],
       ['40:00', 'Bob Jovi: What a Tribute Band Taught Him About Being Memorable'],
       ['41:30', 'Homework and Where to Find Bobby Kerr']],
      widths=[1.1, 5.8])

doc.add_page_break()

# =====================================================================
# SECTION 7 — STRESS TEST, COUNCIL, EP POLISH
# =====================================================================h1('SECTION 7: STRESS TEST, COUNCIL REVIEW & EP POLISH (v2)')
p('This is the record of the second full round. v1 shipped after its own three passes; this section is what a fresh attack found in that finished document, what the council said on the re-run, and what actually changed. Rows here describe fixes that are visible in this document, not suggestions.', italic=True, size=9.5)

h2('7A. Stress Test, round two')
table(['#', 'What broke in v1', 'Fix applied in v2'],
      [['1', 'THE BIG ONE. v1 instructed D.J. to say 83% on air and called it the one number he could assert. That figure came out of a search summary. The underlying article returned a 403 and was never actually opened, so the document was telling the host to state a number nobody had verified. This is the exact failure the research rules exist to prevent, and v1 committed it while lecturing the guest about unsourced stats.',
        'Verified in the v2 pass. It holds: 83% for AI Overviews via Seer Interactive, 93% inside AI Mode via Semrush. But the check surfaced a better number. Pew tracked real behavior and found a cited source gets clicked about 1% of the time. Q3 was rebuilt on it, title 2 was rebuilt on it, and 4D now carries all three figures with their actual sources.'],
       ['2', 'The episode was titled around AI search and only ever interrogated Google. His own pitch says Gemini reads the profile and ChatGPT reads the internet\'s opinion of you, and v1 knew from research that ChatGPT leans on Bing Places, Yelp, and Foursquare rather than reading a Google profile at all. That finding was parked in a research-brief footnote marked useful if he overclaims.',
        'It is now the second half of Q3, written out as its own question with an instruction not to skip it. If his answer is thin, that is a real finding about the depth of his system, and D.J. can hear it in real time.'],
       ['3', 'There was no intro. v1 referenced the intro four times across the landmines, telling D.J. what not to put in it, and never wrote one. The document policed a thing it had not supplied.',
        'Written, on its own page, engineered around every landmine, including a line that gets his contested resume on tape in his own voice instead of D.J.\'s.'],
       ['4', 'No green room brief. The single biggest structural problem with this guest is that Rapid Fire spends both of his best stories in minute one, and v1 tried to solve it entirely with an in-the-moment holding line.',
        'Three green-room sentences added. The third one pre-briefs him that the callbacks are coming, which turns Q9 and Q10 from repetition into a return.'],
       ['5', 'Clock allocation was backwards. Block 4 held two of the five never-cut questions and had the tightest window in the episode at eight minutes for four questions. Block 3, the block the council already wanted shortened, had ten.',
        'Block 3 is now 22:00 to 30:00 and Block 4 is 30:00 to 40:00. Chapter markers retimed to match.'],
       ['6', 'Internal contradiction. Page 1 called the two Rapid Fire callbacks load-bearing, then the never-cut list contained neither of them, which meant the document told D.J. both that he had to do them and that he could drop them.',
        'Q10 added to never-cut, with the reason written next to it. Q9 stays cuttable on purpose, because its value depends on him producing evidence he may not have.'],
       ['7', 'The stat rule contradicted itself. Landmine 5 banned guest-supplied numbers from the descriptions, and then the YouTube description said more than 1,000 agents.',
        'The rule now separates performance claims, which stay banned, from volume claims, which are fine when attributed. The thousand describes what he did. The 10x describes what it supposedly produced. Only one of those is safe.'],
       ['8', 'The homework did not define done. Write out every service you offer, and if you have fewer than ten you have homework, leaves the listener unsure whether the assignment is counting or writing.',
        'Rewritten to a flat number. Write ten, each with a neighborhood in it. Counting is the check, not the task.'],
       ['9', 'Q8 asked the suspension question without arming D.J. with the answer, so a wrong answer would have gone straight to air unchallenged on the one topic in this episode that can actually cost a listener their profile.',
        'Producer note now carries Google\'s actual rules: practitioner name only, no keyword in the name, and service-area setup with the address hidden for agents who do not staff the brokerage office. Plus the follow-up to use if he recommends the thing that gets people suspended.'],
       ['10', 'Q3 sat on the cut list while carrying the episode\'s central premise.',
        'Off the cut list. Q11 took its place, with the reason: the cadence usually falls out of Q9 anyway.']],
      widths=[0.3, 3.2, 3.4])

h2('7B. Council Review, re-run on v2')
p('The board was reconvened on the rebuilt document. Members who signed off on v1 without changes are not repeated here.', italic=True, size=9.5)

h3('Member notes')
table(['Member', 'What they would change'],
      [['Alex Hormozi', 'The dissent is resolved and I will take the trade. Block 3 loses two minutes, Block 4 gains them, and the block with the screenshot question finally has room to breathe. I still think Q9 is the softest question in the episode. The difference now is that if it collapses, it collapses inside eight minutes instead of ten.'],
       ['Brendan Kane', 'Title 2 got better because you stopped reaching for the number everybody else already posted. 83% is on every SEO account in the country. One percent is the same finding with a knife on it. That is the variant.'],
       ['Chip and Dan Heath', 'One percent beats 83% on the mechanism, not just the drama. 83% tells the listener that people do not click, which they already suspect. One percent tells them the thing they were working toward, getting cited, does not pay, which they do not suspect at all. That is a gap opening, not a fact landing.'],
       ['Byron Lazine', 'Now the AI part of this episode has a second engine in it. Every real estate show this year has done Google and AI. Almost none of them have said out loud that ChatGPT is mostly not reading your Google profile. If he confirms that, that is the take and it belongs in the clip pile.'],
       ['MrBeast', 'The intro page is the fix I did not ask for and needed. v1 opened cold on a bio the host had been told he could not state. Now the first ninety seconds are scripted and the guest fills in the risky part himself.'],
       ['Donald Miller', 'The homework got clearer and the homework is the promise. Write ten is a thing a person does. Write out every service you offer is a thing a person means to do.'],
       ['Jon Youshaei', 'You retimed the chapters, which most people would have skipped. Add the note about which rows die when the cut list is used, because a chapter marker pointing at a question that never got asked is worse than no chapter at all.'],
       ['Justin Welsh', 'Two rounds in, the format survived. Standardized Rapid Fire intact, four blocks, bridges written, homework at the end. The green-room brief is a format improvement worth carrying to every guest who pre-sends their advice answers, not just this one.']],
      widths=[1.4, 5.5])

h3('Witnesses called')
bullet('Jonah Berger, re-called on the share claim: ', 'The share driver moved. In v1 it was Q14, the review script, which is social currency. It is now split, because one percent is an anxiety number with a specific target: the agent who has been doing SEO correctly and is about to learn that correct is no longer the same as chosen. That is high arousal and it has a villain. Valuable virality check on the new framing: retell it as "getting cited is not the win, being the sentence is," and it survives the retelling intact. Q14 is still the screenshot. Q3 is now the argument.')
bullet('Chip and Dan Heath, on the curse of knowledge in the new intro: ', 'The intro assumes the listener knows what a Google Business Profile is and assumes nothing about who Bobby Kerr is. Correct on both. The phrase every agent has and almost nobody has finished does the whole job in nine words, because it makes the listener check their own profile mentally before the guest has said a word.')

h3('Title (unchanged recommendation, sharpened runner-up)')
table(['#', 'Title', 'Ingredient', 'Curiosity mechanism'],
      [['1', 'Ranking Is Over. Bobby Kerr on Getting Handed to the Buyer as the Answer.', 'Insight', 'Kills a belief in the first four words, then withholds the mechanism.'],
       ['2', 'Google Cites You and Nobody Clicks. Bobby Kerr on Getting Picked Inside the Answer.', 'Stat, rebuilt in v2', 'Names a win the listener is chasing and reveals it does not pay, before offering the replacement.'],
       ['3', 'The Free Google Page You Claimed and Forgot Is Now Your Front Door (Bobby Kerr)', 'Text and personality', 'Second person accusation. Strongest emotionally, weakest cold.']],
      widths=[0.3, 3.1, 0.9, 2.6])
p('Recommended: still #1 for the feed and the social cuts, #2 for YouTube. The pairing is stronger in v2 because #2 stopped being a louder version of #1 and became a different argument.', italic=True)

h3('The Clip Engine (updated)')
table(['Q#', 'Question', 'Berger emotion', 'Heath gap'],
      [['3', 'A cited source gets clicked one percent of the time, so what does an agent actually put on the profile? Plus: which AI engines read a Google profile and which ones do not?', 'Anxiety with a target. The listener who has been doing this right learns that right is not the same as chosen.', 'Reveals that the goal they were working toward does not pay, before naming the replacement.'],
       ['14', 'The exact text message that gets a client to write a review naming the neighborhood and the problem.', 'Social currency, high. Forwarding it makes the sender look useful.', 'Names the failed version everybody sends first, then withholds the fix until he says it.'],
       ['13', 'Why the words in a review beat the star count.', 'Anxiety plus surprise. Devalues a metric the whole audience tracks.', 'Reopens something the listener thought was finished.']],
      widths=[0.4, 2.6, 2.0, 1.9])
p('Q7, the Q and A removal, dropped off the clip engine in v2. Not because it got weaker, but because Q3 now carries the same anxiety with a number attached, and two versions of the same feeling compete with each other in a clip pile. Q7 stays never-cut for the episode. It is a credibility beat, not a clip.', italic=True, size=9.5)

h3('Live-Description Scrub')
table(['Platform', 'Verdict and fix'],
      [['Facebook', 'Keep, unchanged from v1. Field by field, live is the promise that earns a comment.'],
       ['Instagram', 'Keep. One accusation plus three tags, fixed in v1.'],
       ['TikTok', 'Keep. Opens on the accusation, not on a name nobody recognizes.'],
       ['YouTube', 'Fixed again. The line about getting cited by AI search engines was replaced with which AI engines actually read a Google profile versus which ones do not, because that is the searchable question and it is the one the episode now answers.'],
       ['LinkedIn', 'Keep. The broker angle is the only one native to the platform.']],
      widths=[1.1, 5.8])

h3('Arc Fix')
p('The v1 arc fix held: the Block 2 to Block 3 bridge accuses instead of summarizing, and Q15 stays late. The remaining risk in v2 is different and it is at the front. Q3 is now the strongest question in the first third of the episode, and it sits third, behind a client anecdote and a website comparison. MrBeast pushed to move it to the top of Block 1. It stays at position three for one reason: Q1 has to establish that he rebuilds these things for a living before he is allowed to make a claim about the future of search, or he is just another guy with a stat. The mitigation instead is the cold open, which now front-loads the being handed to the buyer as the answer language, so the listener is holding the premise before Q1 starts.')

h3('Why It Should Work')
bullet('Curiosity mechanism (Heath): ', 'Two gaps stacked. The title takes away ranking, and Q3 takes away getting cited. The listener does not get a replacement for either until the guest supplies one.')
bullet('Share driver (Berger): ', 'Q14 gets forwarded because it makes the sender look useful. Q3 gets argued about, which travels differently and further.')
bullet('Retention move (MrBeast): ', 'Four blocks, each ending on something to do tonight, a bridge that accuses at the sag point, and the most valuable block now has the clock to actually finish.')

h3('The Dissent')
p('Hormozi stood down on Block 3. The new dissent is Byron Lazine, and it is about restraint. His position: the multi-engine finding in Q3 is the most newsworthy thing in this episode, and the document deliberately keeps D.J. from pressing it, because Landmine 4 and the Q7 note both instruct him not to correct the guest on air. Byron argues that on a claim this central, protecting a first-time vendor guest costs the audience the actual answer. The counter is that a guest who feels caught stops giving you mechanism and starts giving you defense, and mechanism is the entire episode. THE EXPERIMENT: run it as written this time and listen to what the second half of Q3 produces. If he engages with it honestly, the no-correcting rule was unnecessary caution and the next vendor guest gets pressed directly. If he deflects, the rule saved the interview and we keep it.')

h2('7C. EP Polish, pass 3 of round two')
bullet('', 'Verified the 83% claim that v1 told D.J. to assert on air, found it holds, and found a better number in the process. Rebuilt Q3, title 2, and three rows of 4D on the Pew finding.')
bullet('', 'Promoted the multi-engine finding from a research-brief footnote to the second half of Q3, with an instruction not to skip it.')
bullet('', 'Wrote the guest intro the whole document had been referring to, with a line that hands his contested resume back to him.')
bullet('', 'Added the green-room brief, and put the Rapid Fire pre-brief in it, which is the only real fix for a guest who pre-sends his best material.')
bullet('', 'Moved two minutes from Block 3 to Block 4 and retimed all twenty chapter markers to match.')
bullet('', 'Added Q10 to the never-cut list, resolving a contradiction where the document called the callbacks essential and cuttable in the same breath.')
bullet('', 'Rewrote the stat rule to separate performance claims from volume claims, which is what v1 meant and not what v1 said.')
bullet('', 'Rewrote the homework to a countable number, and rewrote the note explaining why.')
bullet('', 'Armed Q8 with Google\'s actual naming and address rules so a wrong answer does not go to air unchallenged on the one topic that can cost a listener their profile.')
bullet('', 'Swapped Q3 off the cut list and Q11 onto it.')
bullet('', 'Added the chapter-marker note about which rows to delete if the cut list gets used, on Youshaei\'s note.')
bullet('', 'Re-swept the document for em dashes, curly quotes, and AI-speak after every change above. Still zero.')

doc.save('/Users/djparis/GitHub Projects/keeping-it-real-content-system/guest-prep/Bobby_Kerr_Interview_Prep.docx')
print('saved')

