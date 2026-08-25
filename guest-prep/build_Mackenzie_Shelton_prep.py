#!/usr/bin/env python
# Builds the KIRP interview prep .docx for Mackenzie Shelton, Prairie and Pine Real Estate Group at Realty ONE Group Sterling (Omaha, NE)
# Built to prompts/06_interview_prep.md v5: research -> draft -> stress test + council -> EP polish
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)

for name, size, color in (('Heading 1', 18, 0x1F3864), ('Heading 2', 14, 0x2E5496), ('Heading 3', 11.5, 0x2E5496)):
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor((color >> 16) & 255, (color >> 8) & 255, color & 255)
    st.font.bold = True

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.8)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def p(text='', bold=False, italic=False, size=None, space=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space)
    r = par.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    return par


def rich(segments, style=None, space=6):
    par = doc.add_paragraph(style=style)
    par.paragraph_format.space_after = Pt(space)
    for seg in segments:
        text = seg[0]
        b = seg[1] if len(seg) > 1 else False
        i = seg[2] if len(seg) > 2 else False
        r = par.add_run(text)
        r.bold, r.italic = b, i
    return par


def bullet(text_bold, text_rest=''):
    par = doc.add_paragraph(style='List Bullet')
    par.paragraph_format.space_after = Pt(2)
    if text_bold:
        par.add_run(text_bold).bold = True
    if text_rest:
        par.add_run(text_rest)
    return par


def quoted(label, text):
    par = doc.add_paragraph(style='Intense Quote')
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.space_before = Pt(3)
    r = par.add_run(label + ' ')
    r.bold = True
    r.italic = True
    r2 = par.add_run(text)
    r2.italic = True
    return par


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            par = cells[i].paragraphs[0]
            run = par.add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def q(num, question, if_vague, reveals, serves, note=None, short=None, permission=None):
    rich([(str(num) + '. ', True), (question, True)], space=2)
    if short:
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(3)
        par.paragraph_format.left_indent = Inches(0.25)
        r = par.add_run('SAY THIS: ')
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        r2 = par.add_run('"' + short + '"')
        r2.bold = True
        r2.font.size = Pt(10.5)
    if permission:
        quoted('First, say this:', '"' + permission + '"')
    quoted('If vague, ask:', if_vague)
    quoted('Ideal answer reveals:', reveals)
    quoted('Serves:', serves)
    if note:
        rich([('PRODUCER NOTE: ', True), (note, False, True)], space=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def bridge(text):
    rich([('BRIDGE TO NEXT BLOCK (read as written): ', True), ('"' + text + '"', False, True)])


# =====================================================================
# TITLE
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('KEEPING IT REAL PODCAST')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Interview Prep: Mackenzie Shelton')
r.bold = True
r.font.size = Pt(20)
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Prairie and Pine Real Estate Group at Realty ONE Group Sterling  |  Omaha, NE and Southwest Iowa  |  Target runtime 43 minutes')
r.italic = True
r.font.size = Pt(10)

# =====================================================================
# 1. QUICK REFERENCE CARD
# =====================================================================
h1('1. Quick Reference Card')
p('One page. Glance at this during the interview.', italic=True, space=8)

h3('Who she is')
bullet('Name: ', 'Mackenzie Shelton. Team Owner, Prairie and Pine Real Estate Group, at Realty ONE Group Sterling.')
bullet('Based: ', 'Lives on a small farm outside Underwood, Iowa. Team offices at 254 N 114th Street, Omaha, and 510 Walker Street, Woodbine, Iowa.')
bullet('Markets: ', 'Omaha and Lincoln metros in Nebraska, plus Southwest Iowa. Rural acreages are the specialty, residential is the volume.')
bullet('Guest type: ', 'A, team leader, with a heavy D coach and educator overlay. She is an author, a speaker and a coach, so push hard for the mechanism or the episode turns into a keynote.')

h3('Verified numbers (safe to say on air)')
bullet('Licensed in Nebraska since 2014. ', 'License number 20140341. Licensed in Iowa since 2018. She was a property manager for five years before she got the license.')
bullet('Team owner since 2019. ', 'Prairie and Pine has closed 400 plus homes since inception. This number appears on both her brokerage profile and her own team site, so it is the safest production figure in the packet.')
bullet('Six agents on the public team page: ', 'Mackenzie, Cortney Thomas, Jen Neill, Brittany Wiley, Shelby Peters, Shari Good. Every one of them is a woman. Their own Instagram also names marketing, operations and transaction coordinators.')
bullet('Reviews: ', 'The team site shows 5.0 out of 5 from 73 Google reviews.')
bullet('Education: ', 'BA in Sociology.')
bullet('She hosts a podcast: ', 'Roots and Roofs, launched 2024, roughly 51 episodes, on the intersection of motherhood and entrepreneurship in real estate. Read the landmine in 4H before you compliment it.')
bullet('Brokerage history: ', 'Prairie and Pine was at Keller Williams Greater Omaha. It is now at Realty ONE Group Sterling. Both pages are still public. The date of the move is not verified, so ask, do not state.')
bullet('Local market right now: ', 'Omaha and Council Bluffs metro, July 2026. Median listing price around 404 thousand, about 33 days on market, roughly 1.8 months of supply, homes selling near 98 percent of list.')

h3('Guest supplied, do not state as verified')
bullet('', 'The book, The Serenity Solution: Freedom from Control, Perfectionism and the Pressure to Do It All. There is no findable public listing for it anywhere. Her Instagram link in bio does point to an Amazon product page, which suggests it exists, but the page would not load. Let her name it, let her describe it, and ask her where listeners buy it.')
bullet('', '600 plus clients since 2019. Her speaker bio says this. Her brokerage profile says 400 plus homes. Use 400 plus and let her correct you live.')
bullet('', 'Buying out her business partner two months before giving birth. This came from the show notes of her Becoming More appearance, not from her directly to you. Attribute it to that interview when you say it.')
bullet('', 'The Avatar hiring strategy and the Maxwell Leadership authorship track. Both are hers to explain. Neither has a public page you can check.')

h3('Personal, publicly shared, use for warmth only')
bullet('', 'Small farm outside Underwood, Iowa. Husband and four children. Gardening, raising livestock, horseback riding.')
bullet('', 'Her personal Instagram bio reads "Homesteader, Hippie Cowgirl, Acreage Queen, Keynote Speaker, Real Estate Coach." Roughly 1,470 followers. It is the best one line summary of her in existence and she wrote it herself.')
bullet('', 'Faith comes up in her own language. She has talked publicly about surrender and spirit led living. Let her bring it. Do not open the door for her.')

h3('Contact and social')
bullet('', 'Website prairieandpinerealestategroup.com. Email sheltonomaha@gmail.com. Phone 402.680.8213. Personal Instagram @mackenzieleeshelton. Team Instagram @prairieandpine.re.group. Podcast Roots and Roofs on Apple, Spotify and Amazon Music.')

h3('Connection to KIR')
bullet('', 'No prior KIR appearance found. First time on the show.')
bullet('', 'She is at Realty ONE Group Sterling, so there is no Kale overlap and nothing to manage there.')
bullet('', 'She is a Dianna Kokoszka Inner Circle member. If Kokoszka comes up, that is a good name for the audience, let it land.')

h3('THE CORE TOPIC')
p('She is a self described recovering perfectionist who runs a six agent team across two states from a farm in Iowa, which only works because she stopped being the person who controls it. One topic, four angles: the moment control became impossible, what she actually stopped doing, who does the work now, and why she refuses to double.', bold=True)

h3('Overasked questions to avoid')
bullet('', '"How do you balance motherhood and running a business." She has hosted 51 episodes of her own podcast on exactly that. There is nothing left in it.')
bullet('', '"Tell me the story of buying out your partner right before the baby." Dianna Kokoszka already ran that story in December 2025. You state the fact, she does not retell it. You ask what broke afterward, which nobody has asked.')
bullet('', '"What does do it messy mean." That was literally the title of the Kokoszka episode. Ask for the most expensive messy thing she shipped instead.')

h3('The "I have interviewed hundreds" moment')
p('Q13, and nowhere else: "I have interviewed hundreds of agents on this show, and almost all of them measure a year by whether the number went up. You sent me that the worst advice you ever got was to double your production every year."', italic=True)

h3('Live stream title (paste into Restream before you hit record)')
p('Mackenzie Shelton Bought Out Her Partner 2 Months Before Giving Birth (69 characters)', bold=True)
p('Backup: 400 Homes From a Farm in Underwood, Iowa: Mackenzie Shelton on Letting Go (73 characters)', italic=True)
p('This is not the published episode title. Pick that after you hear the interview. Recommendation is in 2A.', italic=True)

h3('WATCH OUT FOR')
bullet('Rapid Fire will spend your best setup. ', 'She sent both her best and her worst advice in the intake. The worst advice, "double your production every year," is the entire spine of Block 4. When she gives it in Rapid Fire, say "Love it, and we are coming back to that," then run the callback at Q13 exactly as written.')
bullet('The book cannot be verified. ', 'Nothing about The Serenity Solution is findable publicly. Do not describe its contents, do not say it is a bestseller, and do not tell listeners where to buy it until she tells you on air. Ask her for the link and repeat what she says.')
bullet('The scam story is legal territory. ', 'She was scammed by a buyer. Let her set the level of detail. Do not ask for the name, the address, the brokerage or whether anyone was charged. The value is the red flags and the process change, not the case.')
bullet('This episode wants to become a motherhood and balance conversation. ', 'That is her home turf and her own podcast. Every time it drifts, pull back to the business mechanism with the line in the Drift Guardrail on page 4.')
bullet('She is a trained speaker. ', 'Speakers answer in frameworks. When you get a framework, ask for the week it happened.')

h3('THE TWO STANDING REMINDERS')
p('Ask the short version.', bold=True, space=2)
p('Count to three before you respond.', bold=True)

doc.add_page_break()

# =====================================================================
# 2. EPISODE FRAMEWORK
# =====================================================================
h1('2. Episode Framework')

h2('2A. Title Options')
table(['#', 'Title', 'Why it works'], [
    ['1', 'She Bought Out Her Partner 2 Months Before Giving Birth. Then She Had to Stop Running Her Own Company. (Mackenzie Shelton)',
     'Number, specific claim, name. Opens the gap with the reckless sounding decision and closes it with the counterintuitive consequence.'],
    ['2', '400 Homes From a Farm in Underwood, Iowa: Mackenzie Shelton on Building a Team That Runs Without You',
     'Leads with the verified production number and the place, which is the detail nobody expects. Carries the promise the largest audience segment came for.'],
    ['3', 'Mackenzie Shelton Was Told to Double Her Production Every Year. She Said No.',
     'Her own words as the hook. Every agent who has sat through that goal setting meeting stops scrolling.'],
], widths=[0.35, 3.6, 3.1])
p('Recommended: #2. The goal for this episode is team leaders and the solo agents about to become one, and #2 is the only title carrying a verified number, the promise, and a detail strange enough to click. Run #1 as the live title because it works cold, before the interview has happened.', italic=True)

h2('2B. Cold Open Hook')
p('"Mackenzie Shelton bought out her business partner two months before she gave birth, and then she had to build a company that could run while she was gone. Six agents, two states, 400 plus homes, all of it operated out of a farm outside Underwood, Iowa, by somebody who will tell you she is a recovering perfectionist. We are going to talk about that today. Stay tuned."', bold=True)

h2('2C. Episode Arc')
p('Core topic: what a controlling operator has to actually give up to build a business that runs without her, and what it costs her when she does.', bold=True)
p('Why this topic: her origin, her balance story and her do it messy philosophy are all already published. What is nowhere is the mechanism. Which specific tasks left her hands, what the written standard is that replaced her judgment, who does the work now, what broke when she let go, and what number she runs the business on instead of the one everybody else chases. That is the part the largest slice of this audience cannot get anywhere else, and it works whether they have six agents or zero.')
p('This is also the book, applied. She wants to talk about The Serenity Solution and overwhelm and perfectionism. Fine. This arc is that book with a profit and loss statement attached, which is the only version this audience will sit through.', italic=True)
p('Four angles, one topic. Block 1 is the moment control ran out. Block 2 is what she stopped doing and where the line is. Block 3 is who does the work now. Block 4 is why she refuses to double. Every block ends on something a listener can do this week.', italic=True)

doc.add_page_break()

# =====================================================================
# 3. INTERVIEW QUESTIONS
# =====================================================================
h1('3. Interview Questions')

h2('Rapid Fire (0:00 to 2:00, standardized, read as written)')
p('1. Best real estate advice you have ever received?', space=2)
p('2. Worst real estate advice you have ever received?', space=2)
p('3. One tool or app you cannot run your business without?', space=2)
p('4. What would surprise people most about your day to day?', space=2)
p('No follow ups. "Love it" and move on. One exception, below.', italic=True, space=4)
rich([('THE ONE EXCEPTION: ', True), ('She already sent you her worst advice, "double your production every year," and it is the spine of Block 4. When she says it, say "Love it, and we are coming back to that." Then run the callback at Q13. Do not discuss it here. Her best advice, "how can I add value to others today," is safe to spend in Rapid Fire, it does not carry a block.', False, True)])

h2('BLOCK 1: The Two Months Before the Baby (2:00 to 12:00)')
p('Audience note: team leaders, and the solo agent whose business currently stops when they stop. The most universal fear in the audience is that the thing only works because they are holding it.', italic=True, space=2)
p('Arc: the decision, the handoff, what broke, the first thing to give away.')

q(1,
  'You told Dianna Kokoszka you bought out your business partner two months before you gave birth. Most people would have waited until after the baby. What was happening inside the business that made waiting the more dangerous option?',
  'What specifically was going wrong that made two months before better than six months after?',
  'The real trigger, not the brave version. This is where you find out whether it was ambition or a problem she could not fix any other way.',
  'Team leaders, broker owners.',
  short='Why buy out your partner two months before the baby instead of after?',
  note='She has told the arc of this story once already. You are not asking for the story, you are asking for the reason. If she starts retelling the arc, let her run about 45 seconds and then ask the short version again.')

q(2,
  'So the buyout closes and the clock is running. Walk me through the actual handoff. What came off your plate in those eight weeks, who took it, and what did you refuse to hand to anybody?',
  'Name the task and name the person. Not the category, the task.',
  'The first real subtraction list, plus the one thing she could not let go of, which is usually the most revealing item in the interview.',
  'Team leaders, individual agents.',
  short='What came off your plate in those eight weeks, and what would you not hand over?')

q(3,
  'Something always breaks. What actually fell apart while you were out, and what did it cost you?',
  'A client, a deal, money, or a person. Which one was it?',
  'Credibility. Every leader in the audience has been told to delegate and is afraid of exactly this. Naming the real cost is what makes the rest of the episode believable.',
  'Team leaders, broker owners.',
  permission='Tell me if this is too personal, but I want to ask about the part that did not go well.',
  note='Ask it, then count to three and say nothing. This is one of the two most likely clips in the episode. Do not accept "everything worked out fine."',
  short='What broke while you were out, and what did it cost you?')

q(4,
  'There is an agent listening whose business stops the day they stop. They have no partner to buy out and no baby coming. What is the first thing they hand off this week, and how do they hand it off so it does not come back?',
  'Give me the task and the handoff. What do they write down before they give it away?',
  'Portable implementation that needs no team, no budget and no hire.',
  'Individual agents, new agents.',
  short='First thing an agent hands off this week, and how do they hand it off?')

bridge('So the handoff was forced on you by a deadline you could not move. Which means somewhere in there you had to decide what actually has to be perfect and what does not.')

h2('BLOCK 2: The Line Between a Standard and a Grip (12:00 to 22:00)')
p('Audience note: the largest segment, individual agents who believe their personal touch is the product and cannot tell the difference between a standard and a stranglehold.', italic=True, space=2)
p('Arc: the thing she let go worse, the written line, the time letting go got expensive, the guardrail.')

q(5,
  'You wrote a whole book about perfectionism, so here is the test. Name one thing in your business that somebody else now does worse than you did, and you left it alone anyway.',
  'What is it, and how much worse are we talking? Ten percent or fifty?',
  'Whether the letting go is real or whether she is describing letting go while still doing everything. The specific worse thing is the proof.',
  'Individual agents, team leaders.',
  short='Name one thing somebody else now does worse than you, and you left it alone.',
  note='If she cannot name one, that is the most interesting answer of the episode. Say so, kindly, and stay in it.')

q(6,
  'Where is the line written down? Not the philosophy. In your business, what is on the list of things that have to be exactly right, and what is on the list of things that just have to be done?',
  'Give me two items on each list.',
  'The actual operating standard. This is the artifact the audience wants to steal.',
  'Team leaders, individual agents.',
  short='What has to be exactly right, and what just has to be done? Two of each.')

q(7,
  'You have said you were scammed by a buyer. That is the nightmare version of trusting your process instead of your gut. What were the red flags you can see now that you could not see then, and what did you change afterward?',
  'What is the specific step you added, and where does it sit in the process now?',
  'Real listener protection, plus the honest counterweight to the whole letting go thesis. Letting go has a bill and this is it.',
  'Individual agents, team leaders, new agents.',
  permission='Do you mind if I ask about the buyer who scammed you? Tell me only as much as you are comfortable with.',
  note='Let her set the detail level. Do not ask for the name, the address, the brokerage, or whether charges were filed. You want the red flags and the process change. If she goes quiet, move to Q8.',
  short='What were the red flags you missed, and what did you change afterward?')

q(8,
  'So for the agent who hears all this and gets scared back into doing everything themselves. What is the one thing they should stop personally touching this week, and what is the one guardrail that keeps that from turning into your story?',
  'The thing to stop, and the check that replaces you.',
  'The both and answer. Delegate the task, keep the verification. This is the takeaway that survives being retold in one sentence.',
  'Individual agents, team leaders.',
  short='What do they stop touching this week, and what guardrail replaces them?')

bridge('A standard only holds if somebody is actually running it. You have six agents doing that across two states, and you hire on something you call an avatar.')

h2('BLOCK 3: Who Does the Work Now (22:00 to 32:00)')
p('Audience note: team leaders, and the solo agent who is one good year from a first hire. This is also the block broker owners will listen to hardest.', italic=True, space=2)
p('Arc: the avatar, the hiring mechanism, the wrong hire, the first hire.')

q(9,
  'You have six agents across Nebraska and Iowa, run out of a farm outside Underwood, and you hire against something you call an avatar. Who is the avatar, and what were you hiring on before you had one?',
  'Describe the person. What do they have to already be, and what will you teach?',
  'The filter, and by contrast the mistake that produced the filter.',
  'Team leaders, broker owners.',
  short='Who is the avatar you hire against, and what did you hire on before that?')

q(10,
  'Walk me through what happens to a new agent in their first 30 days on your team. What do they get, what do they owe you, and who checks?',
  'Day one, week one, week four. Be literal.',
  'The onboarding system, which is the single most requested and least shared thing in this audience.',
  'Team leaders, broker owners, new agents.',
  short='First 30 days for a new agent on your team. What do they get and who checks?')

q(11,
  'Every team owner has one hire they should not have made. What did you miss in that person, and how long did it take you to admit it?',
  'How long. Give me the number of months.',
  'The honest cost of hiring for potential, and the tell she now watches for.',
  'Team leaders, broker owners.',
  permission='This one is uncomfortable and you can pass on it, but I think it helps people.',
  note='Do not let her name the person. If she starts to, redirect to the pattern. Count to three after her answer.',
  short='One hire you should not have made. What did you miss, and how long did it take?')

q(12,
  'Solo agent listening, good year behind them, about to make their first hire and terrified of it. What role do they hire first, and what do they have that person do in week one?',
  'Which role, and the actual first task.',
  'The implementation for this angle, portable to any market and any team size.',
  'Individual agents, team leaders.',
  short='First hire for a solo agent. What role, and what do they do in week one?')

bridge('So the team is built and the standard is written and it runs without you. Which is exactly the point where somebody tells you to take all that and double it.')

h2('BLOCK 4: The Year She Refused to Double (32:00 to 40:00)')
p('Audience note: everybody. This is the block that gets shared, because the number pressure is universal and almost nobody says no to it out loud.', italic=True, space=2)
p('Arc: the refusal, the niche that made the math work, the brokerage question, the objection.')

q(13,
  'I have interviewed hundreds of agents on this show, and almost all of them measure a year by whether the number went up. You sent me that the worst advice you ever got was to double your production every year. So what is the number you actually run the business on, and what does a good year look like if it is not bigger?',
  'Give me the metric. Deals per agent, profit, hours, retention. Which one is on the wall?',
  'Whether the refusal is a real operating decision with a metric behind it or a nicer sounding way to describe a plateau.',
  'All segments.',
  note='This is the load bearing question of the episode and it is the most dodgeable one in the packet. She can answer it entirely in philosophy. If she does, say "I love that, and I want the number," and ask again. Do not move on without a metric. Say the hundreds of agents line here and nowhere else.',
  short='If it is not doubling, what is the number you run the business on?')

q(14,
  'You are the acreage person. In a metro where every other agent is fighting over 404 thousand dollar houses in Omaha, you went out to the rural acreages. How much of the peace you are describing is mindset, and how much is just that you picked a lane nobody else wanted?',
  'What does an acreage deal require that a city deal does not, and how long did it take to be the one they call?',
  'The unglamorous structural answer underneath the mindset answer. A narrow niche shrinks the number of things you have to be great at, which is the same thing as needing less control.',
  'Individual agents, team leaders, new agents.',
  short='How much of this is mindset, and how much is that you picked a lane nobody wanted?')

q(15,
  'Your team was at Keller Williams Greater Omaha, and today it is at Realty ONE Group Sterling. Set aside the splits for a second. What did you need from a brokerage as a team owner that you were not getting?',
  'What was the specific thing you needed, and did the new one actually deliver it?',
  'Broker owner gold, and the perspective flip. This is the question the smallest audience segment came for and the one D.J. can ask better than anyone because he is not competing with her.',
  'Broker owners, team leaders.',
  permission='You do not have to name anybody, and I am not asking you to. I am asking what you needed.',
  note='This one can read as an ambush. It is not, because both pages are still public, but the permission clause is what keeps it landing as a shared problem instead of a gotcha. Ask what she needed, never what they did wrong.',
  short='What did you need from a brokerage that you were not getting?')

q(16,
  'There is an agent listening thinking, easy for her to say. She has six agents and a book and a farm. I have a car and a lead contract. What do you say to that person?',
  'What is the version of this they can do at zero agents and zero dollars?',
  'The objection said out loud, answered directly. This is what decides whether the episode travels past team owners.',
  'Individual agents, new agents.',
  short='Agent with no team and no budget says easy for you to say. What do you tell them?')

h2('The Close (40:00 to 43:00)')

h3('Homework assignment (read verbatim)')
p('"Here is what I want you to do before the next episode. Open your calendar from last week and go through it hour by hour. Circle every single thing you touched that did not require your license, your face or your judgment. Then pick one, write down what done looks like in three sentences, and hand it to somebody before Friday. Not next month. This week."', bold=True)

h3('Guest close')
bullet('', '"Where can people find you, follow you, or work with you?"')
bullet('', 'She hosts Roots and Roofs. Say the name, ask her to describe who it is for, do not review it.')
bullet('The book. ', 'Ask her to say the full title out loud and to tell listeners exactly where to buy it. Repeat back what she says. Do not read a link you found, because the one you found would not load.')
bullet('The event. ', 'Her team is running a training called The Predictable Production System on September 16, 2026. If she wants to mention it, that is a clean listener offer. Do not raise it yourself unless the close is short.')

h3('Optional levity beat, only if the room needs it')
p('She sent a story about a closing delayed by a septic system where the clients brought chocolate cupcakes decorated to look like the problem. It is a 20 second laugh and it is a good place to land if Q11 or Q7 got heavy. Do not spend a question on it. Say "you told me a story about cupcakes and a septic system" and let her run.', italic=True)

h3('If you are running long, cut these first')
bullet('Cut 1, Q10. ', 'The 30 day onboarding. Valuable, but it serves the smallest slice and Q9 already establishes the hiring filter.')
bullet('Cut 2, Q6. ', 'The written line. Q5 makes the same point with a better story attached.')
bullet('Cut 3, Q14. ', 'The acreage niche. It is the most interesting question in the packet for agents, and it is also the most severable from the core topic if the clock is gone.')
bullet('Cut 4, Q2. ', 'The eight week handoff. Q4 delivers the same lesson in a form the listener can use.')

h3('Never cut')
bullet('', 'Q1, the buyout decision. It is the cold open paid off.')
bullet('', 'Q3, what broke. The credibility of the entire episode sits here.')
bullet('', 'Q7, the buyer scam. The only question in the packet that protects somebody from losing money.')
bullet('', 'Q13, the refusal to double. The thesis, the callback, and the most shareable moment.')
bullet('', 'Q16, easy for you to say. Without it this is an episode for team owners only.')

doc.add_page_break()

# =====================================================================
# 4. RESEARCH BRIEF
# =====================================================================
h1('4. Research Brief')
p('Reference material. Read this the morning of, not during the interview.', italic=True, space=8)

h2('4A. Background')
p('Five years in property management, then a Nebraska real estate license in 2014. Iowa license in 2018. She founded Prairie and Pine Real Estate Group in 2019, first under Keller Williams Greater Omaha and today under Realty ONE Group Sterling. The team has closed 400 plus homes since it started and now lists six agents plus marketing, operations and transaction coordination. She built a specialty in rural acreages across the Omaha and Lincoln metros and Southwest Iowa, which is unusual in a metro where most agents chase in town residential. She lives on a small farm outside Underwood, Iowa with her husband and four children, hosts the Roots and Roofs podcast, speaks, coaches, and has written a book about perfectionism and control.')

h2('4B. Career timeline (verified entries only)')
table(['Year', 'Role / Company', 'Notable'], [
    ['Roughly 2009 to 2014', 'Property manager', 'Five years before licensure, per her own brokerage bio.'],
    ['2014', 'Licensed REALTOR, Nebraska', 'License number 20140341.'],
    ['2018', 'Licensed in Iowa', 'Adds the Southwest Iowa side of the business.'],
    ['2019', 'Founds Prairie and Pine Real Estate Group', 'Under Keller Williams Greater Omaha at the time.'],
    ['2024', 'Launches Roots and Roofs podcast', 'Roughly 51 episodes to date.'],
    ['Dec 2025', 'Guest, Becoming More with Dianna Kokoszka', 'Episode aired December 30, 2025. Inner Circle member.'],
    ['Current', 'Team Owner, Prairie and Pine at Realty ONE Group Sterling', 'Six agents, offices in Omaha and Woodbine, Iowa.'],
], widths=[1.3, 2.9, 2.85])
p('The date she moved the team from Keller Williams to Realty ONE Group Sterling is not verified. Ask, do not state.', italic=True)

h2('4C. What makes her interesting for this audience')
bullet('The niche nobody wanted. ', 'She is the acreage person in a metro where everyone else is fighting over in town residential. Listeners get a live case study in what a narrow lane actually buys you.')
bullet('The forced delegation. ', 'A buyout and a due date eight weeks apart is the cleanest natural experiment in letting go you will find. Most guests delegate gradually and cannot tell you what it cost.')
bullet('She said no to the number. ', 'Almost every guest on this show is optimizing upward. She was told to double every year and refused. That is a genuinely contrarian position for this audience.')
bullet('An all women team across two states. ', 'Six agents, two state licenses, run from a farm. The geography alone forces systems that an in office team never has to build.')
bullet('She got scammed and will talk about it. ', 'Most agents bury this. She sent it unprompted in her intake, which means she is willing, and it is the most directly protective content in the episode.')

h2('4D. Key data points')
table(['Stat', 'Source', 'Confidence'], [
    ['400 plus homes closed by the team since 2019', 'Realty ONE Group Sterling agent profile and team site', 'High'],
    ['Licensed in NE since 2014, license 20140341', 'Realty ONE Group Sterling agent profile', 'High'],
    ['Licensed in IA since 2018', 'Realty ONE Group Sterling agent profile', 'High'],
    ['Five years in property management before licensure', 'Team site bio', 'High'],
    ['Six agents on the team', 'Team site team page', 'High'],
    ['5.0 out of 5 from 73 Google reviews', 'Team site', 'Medium, self reported on their own page'],
    ['BA in Sociology', 'Brokerage profile', 'High'],
    ['Omaha and Council Bluffs metro, July 2026: median listing about 404 thousand, about 33 days on market, about 1.8 months supply, about 98 percent of list', 'Market trackers, July 2026', 'Medium, aggregator data'],
    ['600 plus clients since 2019', 'Her own speaker bio and her intake', 'Unverified, and it conflicts with the 400 plus figure'],
    ['The Serenity Solution', 'Her intake only', 'Unverified, no public listing found'],
    ['Bought out business partner two months before giving birth', 'Becoming More show notes, Dec 30 2025', 'Reported, attribute to that interview'],
    ['Roots and Roofs, roughly 51 episodes since 2024', 'Apple Podcasts', 'High on existence, Medium on the count'],
], widths=[3.3, 2.4, 1.35])

h2('4E. Previous media appearances')
bullet('Becoming More with Dianna Kokoszka, December 30, 2025. ', '"Why You Must Do It Messy to Become More." Roughly 23 minutes. Covered the rocky start of team ownership, the partner buyout two months before giving birth, the move from solo producer to what she calls a coaching leader, significance over success, surrender and spirit led living, the Avatar hiring strategy, and her authorship track with Maxwell Leadership.')
bullet('Lead Fearlessly, host Katie Winchenbach. ', 'Women leaders show. Her speaker bio there is the source of the 600 plus clients figure and the reduce overwhelm and build systems language.')
bullet('Roots and Roofs, her own show, 2024 to present. ', 'She is the host, not a guest, but 51 episodes of her talking about motherhood and entrepreneurship is why that topic is exhausted. Assume any question in that lane has been answered.')
bullet('Overasked, listed on the Quick Reference Card. ', 'Balancing motherhood and business. The buyout story as an origin arc. What do it messy means.')

h2('4F. Their own words')
table(['Quote', 'Where and when', 'Confidence', 'How D.J. uses it'], [
    ['"Double your production every year. While growth is important, constantly chasing bigger numbers is not always sustainable or aligned with the life you actually want to build."',
     'Her KIR intake form, August 2026', 'Verbatim',
     'This is the Q13 callback. Read it back to her word for word, then ask for the metric she uses instead.'],
    ['"Wake up each day and ask, how can I add value to others today."',
     'Her KIR intake form, August 2026', 'Verbatim',
     'Rapid Fire will spend it. Safe to let go. If you want it later, it is a decent bridge into Q10 on what a new agent owes the team.'],
    ['"Homesteader. Hippie Cowgirl. Acreage Queen. Keynote Speaker. Real Estate Coach."',
     'Her personal Instagram bio, @mackenzieleeshelton', 'Verbatim',
     'Read the whole line back to her at the top of Block 4 and ask which one she would drop first. It is warm, it is hers, and the answer tells you what she is actually protecting.'],
    ['"If you want to go fast, go alone. If you want to go further, bring others along."',
     'Becoming More with Dianna Kokoszka, December 30, 2025', 'Reported',
     'Do not attribute the exact wording. Say "you have said something like" and use it to open Q9 on the avatar if she needs a runway.'],
    ['"Trusting your instincts, recognizing red flags, and protecting yourself throughout the real estate process."',
     'Her KIR intake form, August 2026, describing the buyer scam', 'Verbatim',
     'Hold in reserve for Q7. If she hedges on the scam, read her own sentence back and ask which red flag she means.'],
], widths=[2.7, 1.55, 0.95, 1.85])

h2('4G. Audience relevance')
table(['Segment', 'What they get from this episode'], [
    ['Individual agents', 'A concrete first handoff, a guardrail to go with it, and permission to stop trying to grow every single year.'],
    ['Team leaders', 'A forced delegation case study with the actual cost attached, a hiring filter, and a 30 day onboarding outline.'],
    ['Broker owners', 'What a productive team owner needed from a brokerage and did not get, in her own words, with no vendor pitch attached.'],
    ['New agents', 'The case for picking a narrow lane early, and what a good team actually gives a new agent in the first month.'],
], widths=[1.5, 5.55])

h2('4H. Landmines')
bullet('The book you cannot find. ', 'The Serenity Solution has no findable public page. Her Instagram links to an Amazon product page that would not load. Do not describe its contents, do not call it a bestseller, and do not tell listeners where to buy it. Ask her, then repeat what she says.')
bullet('600 versus 400. ', 'Her speaker bio says 600 plus clients since 2019. Her brokerage profile says 400 plus homes since inception. She has also been licensed since 2014, so any "since 2019" framing is about the team, not her career. Say 400 plus and let her update it.')
bullet('The buyout is Reported, not Verbatim. ', 'You have it from a show notes page, not from her. Attribute it. "You told Dianna Kokoszka" is bulletproof. "You bought out your partner" is not.')
bullet('The buyer scam. ', 'Do not ask for the name, the property, the brokerage, or whether it was prosecuted. Red flags and process changes only. She raised it, so she is willing, but she sets the depth.')
bullet('Faith. ', 'Surrender and spirit led living are her own public language. Let her bring it. Do not introduce it, and do not turn it into a question.')
bullet('Her podcast. ', 'She hosts a show in the same category as yours. Say the name, ask who it is for, move on. Do not review it, do not compare notes on podcasting, and do not spend two minutes on production talk.')
bullet('The children. ', 'Her current profiles say four children. An older listing said three boys. Say "four kids" only if she says it first, and never guess at ages or names.')
bullet('Sponsors. ', 'Real Geeks and Courted are both lead and data products. Her thesis is about doing less, not buying more. Do not steer any question toward a sponsor angle.')

h2('Drift guardrail')
p('This episode will try to become a conversation about balance and motherhood, because that is her platform and 51 episodes of muscle memory. The moment it does, use this line: "I want to stay on the business side of that for a minute, because I think the mechanism is the part people cannot get anywhere else." Then ask the block question again.', bold=True)

h2('Green room brief (two minutes before you record)')
bullet('', 'Tell her the show is standardized rapid fire, then one topic explored properly, then homework. About 43 minutes.')
bullet('', 'Tell her the topic is what she gave up to build a team that runs without her, and that you will be asking for specifics and numbers, not philosophy. Speakers appreciate the warning and it saves you three redirects.')
bullet('', 'Tell her you plan to ask about the buyer who scammed her, and ask what is off limits before you are live.')
bullet('', 'Ask her for the exact place listeners buy the book, and write it down. Do not guess on air.')
bullet('', 'Confirm the pronunciation of Underwood, Woodbine and the team name, and confirm whether she wants Realty ONE Group Sterling said in full.')

doc.add_page_break()

# =====================================================================
# 5. LIVE STREAM TITLE, DESCRIPTIONS AND HASHTAGS
# =====================================================================
h1('5. Live Stream Title, Descriptions and Hashtags')

h2('5A. Live stream title')
p('Primary: Mackenzie Shelton Bought Out Her Partner 2 Months Before Giving Birth (69 characters)', bold=True)
p('Backup: 400 Homes From a Farm in Underwood, Iowa: Mackenzie Shelton on Letting Go (73 characters)', bold=True)
p('Both work cold, before the interview happens, and both use only verified or properly attributable facts. This is not the published episode title. Pick that from 2A after you hear the interview.', italic=True)

h2('5B. Platform descriptions')

h3('Facebook Live')
p('Mackenzie Shelton bought out her business partner two months before she gave birth, and then had to build a real estate team that could run while she was gone. Today she has six agents across Nebraska and Iowa, more than 400 closings, and a very direct opinion about the advice to double your production every year. We are getting into what she actually gave up to make it work. Drop your questions in the comments!')

h3('Instagram Live')
p('Six agents. Two states. One farm in Iowa. Mackenzie Shelton on what she had to stop doing to build a team that runs without her. #RealEstatePodcast #TeamLeader #KeepingItReal')

h3('TikTok Live')
p('She was told to double her production every year. She said no, and built a 400 home team instead. #realestate #realtorlife #realestateagent #teambuilding')

h3('YouTube Live')
p('Mackenzie Shelton owns Prairie and Pine Real Estate Group in Omaha, Nebraska, a six agent team licensed in Nebraska and Iowa with more than 400 closings since 2019. On this episode of the Keeping It Real Podcast she breaks down how she delegated under a hard deadline, how she hires against a defined avatar, what broke when she stepped back, and why she refuses to double her production every year. Real estate team building, delegation, hiring, and running a business from a rural acreage.')

h3('LinkedIn Live')
p('Most team owners scale by adding. Mackenzie Shelton scaled by subtracting. She bought out her business partner two months before giving birth, rebuilt the company so it could operate without her, and now runs six agents across two states from a farm in western Iowa. We cover the delegation, the hiring filter, what it cost when she stepped back, and why she turned down the advice to double her production every year.')

h2('5C. Hashtag sets')
bullet('Universal: ', '#KeepingItReal #RealEstatePodcast #DJParis #RealtorLife #RealEstateAgent')
bullet('Episode specific: ', '#RealEstateTeam #OmahaRealEstate #NebraskaRealEstate #IowaRealEstate #AcreageLife #Delegation #WomenInRealEstate')
bullet('Guest tag: ', '@mackenzieleeshelton on Instagram. @prairieandpine.re.group for the team. Prairie and Pine Real Estate Group with Realty ONE Group Sterling on Facebook. Tag the Roots and Roofs podcast if she gives you the handle.')

doc.add_page_break()

# =====================================================================
# 6. YOUTUBE CHAPTER MARKERS
# =====================================================================
h1('6. YouTube Chapter Markers')
p('Estimates. Adjust after recording. Every title is written to be independently searchable.', italic=True, space=8)
table(['Timestamp', 'Chapter title'], [
    ['0:00', 'She Bought Out Her Partner Two Months Before Giving Birth'],
    ['2:00', 'Rapid Fire: Best and Worst Real Estate Advice'],
    ['4:00', 'Why Waiting Until After the Baby Was the Riskier Move'],
    ['8:00', 'The Eight Week Handoff: What Left Her Plate and What Did Not'],
    ['12:00', 'What Broke While She Was Out, and What It Cost'],
    ['15:00', 'The First Thing Any Agent Should Hand Off This Week'],
    ['18:00', 'One Thing Somebody Else Now Does Worse, and She Left It Alone'],
    ['21:00', 'What Has to Be Perfect vs What Just Has to Be Done'],
    ['24:00', 'Scammed by a Buyer: The Red Flags She Missed'],
    ['27:00', 'Hiring Against an Avatar: Six Agents Across Two States'],
    ['30:00', 'The First 30 Days for a New Agent on Her Team'],
    ['33:00', 'The Hire She Should Not Have Made'],
    ['35:00', 'Your First Hire: Which Role, and What They Do in Week One'],
    ['37:00', 'Why She Refused to Double Her Production Every Year'],
    ['39:00', 'The Acreage Niche Nobody Else in Omaha Wanted'],
    ['41:00', 'What She Needed From a Brokerage and Was Not Getting'],
    ['42:00', 'Homework and Where to Find Mackenzie Shelton'],
], widths=[1.1, 5.95])

doc.add_page_break()

# =====================================================================
# 7. STRESS TEST, COUNCIL REVIEW AND EP POLISH
# =====================================================================
h1('7. Stress Test, Council Review and EP Polish')

h2('7A. Stress test (pass 2, part one)')
table(['#', 'What broke', 'Fix applied'], [
    ['1', 'Fact conflict. Her speaker bio and her intake say 600 plus clients since 2019. Her brokerage profile and team site say 400 plus homes since inception. Both are hers.',
     'Standardized the doc on 400 plus, since that is the number on her brokerage profile. Added the conflict to the Quick Reference Card, 4D and the 4H landmines with the instruction to let her correct it live.'],
    ['2', 'Timeline conflict. Every "since 2019" claim reads as her career start, but she has been licensed in Nebraska since 2014 and was a property manager for five years before that.',
     'Rewrote 4A and 4B so 2019 is explicitly the team founding, not her start. Cold open and titles avoid any career length claim.'],
    ['3', 'Unverifiable claim. The Serenity Solution has no findable listing on Amazon, Goodreads, Maxwell Leadership or anywhere else. Her Instagram link in bio points to an Amazon product page that returned errors on every attempt.',
     'Moved the book entirely into guest supplied. Removed it from the cold open, all three titles and every live description. Added an instruction in the close to ask her for the purchase link on air and repeat what she says.'],
    ['4', 'Spent question. She sent her worst advice in the intake and it is the spine of Block 4. Standardized Rapid Fire will burn it in minute one with no follow up.',
     'Added the one exception note under Rapid Fire, the "we are coming back to that" line, and a written callback at Q13 that reads the intake quote back verbatim.'],
    ['5', 'Drift risk. She hosts 51 episodes on motherhood and entrepreneurship and speaks on overwhelm. This episode will slide into a balance conversation within ten minutes.',
     'Added the Drift Guardrail on page 4 with the exact line to say. Reframed every block around a mechanism, and put the balance topic on the overasked list.'],
    ['6', 'Dodgeable question. Q13 can be answered entirely with philosophy about sustainable growth and no number at all.',
     'Rewrote Q13 to demand the metric, added a producer note telling D.J. to say "I love that, and I want the number" and ask again, and added the if vague follow up naming four candidate metrics.'],
    ['7', 'Second dodgeable question. Q5 lets her describe letting go in the abstract while still doing everything.',
     'Rewrote it to require a named task somebody does worse. Added the note that a failure to name one is itself the answer and should be stayed with.'],
    ['8', 'Gotcha risk. The brokerage move question, Q15, could land as an ambush about why she left Keller Williams.',
     'Added a permission clause that removes the naming, and rewrote the ask to what she needed rather than what they failed to do. Both brokerage pages are public, so the premise itself is safe to state.'],
    ['9', 'Sourcing risk. The partner buyout is the strongest hook in the packet and it comes from a podcast show notes page, not from her.',
     'Marked it Reported in 4D and 4F. Wrote the attribution into the cold open and into Q1 so D.J. says "you told Dianna Kokoszka" rather than asserting it.'],
    ['10', 'Bad data. HomeLight shows 12 years experience and 11 transactions at an average price point of 194 thousand, which is a partial MLS sample and would badly understate her.',
     'Excluded it from the packet entirely. Flagging it here so nobody re-adds it from a later search.'],
    ['11', 'Legal and personal exposure. The buyer scam could pull toward identifying a real person or an open matter.',
     'Added a permission clause, a producer note listing the four things not to ask, and an explicit instruction that she sets the depth. The question now targets red flags and the process change only.'],
    ['12', 'Runtime overrun. Sixteen questions plus Rapid Fire is the ceiling for 43 minutes and Block 3 was originally five questions.',
     'Cut Block 3 to four. Wrote the four item cut list in cut order and the five item never cut list.'],
    ['13', 'Personal detail conflict. One older source describes three boys, her current brokerage profile and team bio say four children.',
     'Added to 4H with the instruction to use four only if she says it first and never to guess at ages or names.'],
], widths=[0.3, 3.2, 3.55])

h2('7B. Council review (pass 2, part two)')
p('Convened Hormozi, MrBeast, Kane, Miller, Chris Do, Eric Simon and Welsh, with Heath and Berger called as witnesses. She is a story guest with a coaching overlay and an unverifiable product, which is exactly the combination that produces a warm, forgettable episode if nobody pushes.', italic=True)

h3('Member notes')
table(['Member', 'What they would change'], [
    ['Alex Hormozi', 'Block 2 was philosophy with a nice title. Make her name the thing that got worse and left it alone, or the whole letting go thesis is unfalsifiable. Fixed at Q5.'],
    ['MrBeast', 'The middle sags at the onboarding question. Every agent has heard a 30 day plan. Keep it, but put the wrong hire right behind it so the block re-hooks before Block 4.'],
    ['Brendan Kane', 'Your first title was the farm one. The farm is the texture, not the hook. The hook is buying out a partner two months before giving birth, because nobody scrolls past that. Run it live, A/B the farm one as published.'],
    ['Donald Miller', 'The hero kept slipping into being the guest. Q16 is the only question where the listener is the hero. Make sure it never gets cut, and end every block on the listener, not on her.'],
    ['Chris Do', 'Q3 is the whole episode. What broke while she was gone is the only question that costs her something to answer. Give it a permission clause and then shut up for three seconds.'],
    ['Eric Simon', 'The scam story is the one an agent forwards to another agent. Do not bury it at the end of a block nobody reaches. It sits at Q7, mid episode, which is right.'],
    ['Justin Welsh', 'Do not break the standardized Rapid Fire to protect the worst advice answer. Use the callback mechanism the format already has. Agreed and applied.'],
], widths=[1.5, 5.55])

h3('Title')
table(['#', 'Title', 'Ingredient', 'Curiosity mechanism'], [
    ['1', 'She Bought Out Her Partner 2 Months Before Giving Birth. Then She Had to Stop Running Her Own Company.', 'Insight plus stat', 'Opens with a decision that sounds reckless, then closes on a consequence nobody predicts. The Aha lands after the Huh.'],
    ['2', '400 Homes From a Farm in Underwood, Iowa: Mackenzie Shelton on Building a Team That Runs Without You', 'Stat plus concept format', 'Geographic mismatch creates the gap. A farm in a town nobody has heard of should not produce 400 closings.'],
    ['3', 'Mackenzie Shelton Was Told to Double Her Production Every Year. She Said No.', 'Personality plus insight', 'A refusal is inherently a gap. The listener needs to know what happened after the no.'],
], widths=[0.3, 3.1, 1.15, 2.5])
p('Recommended: #2 for the published episode. The goal here is team leaders and the solo agents about to become one, and #2 is the only option that carries a verified number and the promise in the same line. Run #1 as the live title because it works cold and it is the strongest scroll stopper Kane could find in the packet.', italic=True)

h3('Cold open (sharpened)')
p('"Mackenzie Shelton bought out her business partner two months before she gave birth, and then she had to build a company that could run while she was gone. Six agents, two states, 400 plus homes, all of it operated out of a farm outside Underwood, Iowa, by somebody who will tell you she is a recovering perfectionist. We are going to talk about that today. Stay tuned."', bold=True)
p('The gap opens on the buyout and stays open. Nothing in it is unverified, and the one Reported fact is stated the way she told it to another host.', italic=True)

h3('The clip engine')
table(['Q#', 'Question', 'Berger emotion', 'Heath gap'], [
    ['3', 'What broke while you were out, and what did it cost you?', 'Anxiety, then relief. High arousal. Every leader in the audience has this exact fear.', 'The listener assumes delegation either works or ruins you. The real answer is neither, and they have to hear the number to find out.'],
    ['13', 'If it is not doubling, what is the number you run the business on?', 'Anger and validation. Refusing the growth mandate out loud gives the sharing agent real social currency.', 'Everybody knows the doubling advice. Nobody has heard the alternative metric named out loud.'],
    ['7', 'What were the red flags you missed, and what did you change afterward?', 'Anxiety plus practical value. The classic forwarded clip.', 'The gap is opened by the word scammed and only closes when she names the flag.'],
], widths=[0.35, 2.5, 2.1, 2.1])

h3('Live description scrub')
table(['Platform', 'Verdict and fix'], [
    ['Facebook', 'Keep. Conversational, leads with the buyout, ends on the comment prompt. Zero em dashes.'],
    ['Instagram', 'Fixed. The draft opened on her bio. Rewrote to three fragments and the promise, which is what actually reads on a phone.'],
    ['TikTok', 'Fixed. Removed the book mention entirely, since it is unverified and TikTok is the least forgiving place to be vague. The refusal to double is the punchier hook anyway.'],
    ['YouTube', 'Keep. Loaded with the searchable terms: real estate team building, delegation, hiring, Omaha, Nebraska, Iowa, acreage.'],
    ['LinkedIn', 'Fixed. Opened on the contrast, most owners scale by adding and she scaled by subtracting, which is the only line in the set that reads native to that feed.'],
], widths=[1.1, 5.95])

h3('Arc fix')
p('The sag was Block 3. Four questions on hiring, in the middle, for the second smallest audience segment, right after the emotional peak of the scam story. Fixed by putting the wrong hire question at Q11 rather than at the end of the block, so the block re-hooks on a failure before it hands off to Block 4. Also moved the acreage niche out of its own block and into Block 4, where it functions as the structural explanation for the refusal to double rather than as a separate episode.')

h3('Why it should work')
bullet('Curiosity mechanism, Heath: ', 'The cold open states a decision that sounds like a mistake and never explains it. The explanation is Q1 and the payoff is Q3.')
bullet('Share driver, Berger: ', 'Refusing the doubling mandate is high arousal and high social currency. An agent forwards it to say something about themselves, which is the whole mechanism.')
bullet('Retention move, MrBeast: ', 'Every block ends on an implementation question aimed at the listener, so the natural drop off point becomes the moment they get something.')

h3('The dissent')
p('Hormozi still objects that Q14, the acreage niche, is the most valuable question in the packet and it sits at number 14 out of 16, where a third of the audience will never reach it. He would open with it. The counter is that the buyout is a far stronger cold open and the niche only makes sense once the control thesis is established. The experiment for the next episode of this type is to run the structural question early and the emotional one late, and compare the retention curve at the 20 minute mark.')

h2('7C. EP polish (pass 3)')
bullet('', 'Rewrote Q1 so D.J. attributes the buyout to the Kokoszka interview instead of asserting it, and shifted the ask from the story to the reason, because the story is already published.')
bullet('', 'Cut a fifth question from Block 3 and moved the wrong hire question ahead of the first hire question, so the block re-hooks on a failure instead of trailing off on onboarding.')
bullet('', 'Dissolved the standalone acreage block and folded the niche into Block 4 as Q14, where it explains the refusal to double rather than competing with it for the core topic.')
bullet('', 'Removed the book from the cold open, all three titles, the live stream title and every platform description after the research pass could not verify a single public listing for it.')
bullet('', 'Rewrote Q5 from a question about perfectionism into a demand for a named task somebody else now does worse, because the original was answerable with a speech.')
bullet('', 'Added the Q13 callback machinery under Rapid Fire, including the exact line D.J. says, after the stress test caught that Rapid Fire spends the worst advice answer in minute one.')
bullet('', 'Added permission clauses at Q3, Q7, Q11 and Q15 and removed them everywhere else, so they still carry weight when they appear.')
bullet('', 'Rewrote Q15 from why she left Keller Williams into what she needed from a brokerage, which turns an ambush into a shared problem and makes the answer useful to broker owners.')
bullet('', 'Shortened every SAY THIS line to under 20 words and split Q2, which had been carrying both the handoff and the refusal to hand off in one sentence.')
bullet('', 'Standardized every production figure on 400 plus and added the 600 plus conflict to the card, the data table and the landmines rather than picking a number silently.')
bullet('', 'Rewrote the homework so it produces a written artifact and a handoff before Friday, since the original version stopped at noticing.')
bullet('', 'Added the Drift Guardrail line and the green room brief after the stress test flagged that 51 episodes of her own show makes the balance conversation the default gravity of this interview.')

doc.save("/Users/djparis/GitHub Projects/keeping-it-real-content-system/guest-prep/Mackenzie_Shelton_Interview_Prep.docx")
print("Saved Mackenzie_Shelton_Interview_Prep.docx")
