#!/usr/bin/env python
# Builds the KIRP interview prep .docx for Megan Walters, The Walters Team at eXp Realty (Columbia, MO)
# Built to prompts/06_interview_prep.md v5: research -> draft -> stress test + council -> EP polish
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)

for name, size, color in (('Heading 1', 18, 0x1F3864), ('Heading 2', 14, 0x2E5496), ('Heading 3', 11.5, 0x2E5496)):
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor((color >> 16) & 255, (color >> 8) & 255, color & 255)
    st.font.bold = True

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.8)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def p(text='', bold=False, italic=False, size=None, space=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space)
    r = par.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    return par


def rich(segments, style=None, space=6):
    par = doc.add_paragraph(style=style)
    par.paragraph_format.space_after = Pt(space)
    for seg in segments:
        text = seg[0]
        b = seg[1] if len(seg) > 1 else False
        i = seg[2] if len(seg) > 2 else False
        r = par.add_run(text)
        r.bold, r.italic = b, i
    return par


def bullet(text_bold, text_rest=''):
    par = doc.add_paragraph(style='List Bullet')
    par.paragraph_format.space_after = Pt(2)
    if text_bold:
        par.add_run(text_bold).bold = True
    if text_rest:
        par.add_run(text_rest)
    return par


def quoted(label, text):
    par = doc.add_paragraph(style='Intense Quote')
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.space_before = Pt(3)
    r = par.add_run(label + ' ')
    r.bold = True
    r.italic = True
    r2 = par.add_run(text)
    r2.italic = True
    return par


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            par = cells[i].paragraphs[0]
            run = par.add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def q(num, question, if_vague, reveals, serves, note=None, short=None, permission=None):
    rich([(str(num) + '. ', True), (question, True)], space=2)
    if short:
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(3)
        par.paragraph_format.left_indent = Inches(0.25)
        r = par.add_run('SAY THIS: ')
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        r2 = par.add_run('"' + short + '"')
        r2.bold = True
        r2.font.size = Pt(10.5)
    if permission:
        quoted('First, say this:', '"' + permission + '"')
    quoted('If vague, ask:', if_vague)
    quoted('Ideal answer reveals:', reveals)
    quoted('Serves:', serves)
    if note:
        rich([('PRODUCER NOTE: ', True), (note, False, True)], space=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def bridge(text):
    rich([('BRIDGE TO NEXT BLOCK (read as written): ', True), ('"' + text + '"', False, True)])


# =====================================================================
# TITLE
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('KEEPING IT REAL PODCAST')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Interview Prep: Megan Walters')
r.bold = True
r.font.size = Pt(20)
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('The Walters Team at eXp Realty  |  Columbia, Missouri  |  Target runtime 43 minutes')
r.italic = True
r.font.size = Pt(10)

# =====================================================================
# 1. QUICK REFERENCE CARD
# =====================================================================
h1('1. Quick Reference Card')
p('One page. Glance at this during the interview.', italic=True, space=8)

h3('Who she is')
bullet('Name: ', 'Megan Walters. Owner, The Walters Team at eXp Realty. Runs it with her husband Jesse.')
bullet('Based: ', 'Columbia, Missouri. Lifelong Columbia resident. Office at 717 Cherry Street Suite C.')
bullet('Markets: ', 'Columbia and mid-Missouri. Residential, land, new construction. Also an investor with a rental portfolio.')
bullet('Guest type: ', 'A, producing agent and team leader, with an educator overlay. Push for systems, daily habits, specific numbers.')

h3('Verified numbers (safe to say on air)')
bullet('2022: ', 'National Association of REALTORS 30 Under 30 honoree. She was 28. NAR listed her at House of Brokers Realty at the time.')
bullet('That year, individually: ', '66 sides and 17.6 million dollars in volume. This is the strongest verified number in the packet. Use it.')
bullet('2020: ', 'Columbia Board of REALTORS Salesperson of the Year. Also a COMO Magazine Women of Excellence finalist.')
bullet('2018 and 2019: ', 'Columbia Business Times Top of the Town, Top Salesperson.')
bullet('COMO Magazine 20 Under 40: ', 'Her team site says 2021, another listing says 2020. Say "a couple of years back" or ask her.')
bullet('Founded: ', 'Homes by Megan in 2017, rebranded to The Walters Team, moved to eXp. New office opened next to the Tiger Hotel in December 2022.')
bullet('Team (April 2023 profile): ', 'Megan, Jesse Walters, Autumn Lear (first agent hired), Brittney Stone (marketing manager).')
bullet('Local market right now: ', 'Columbia median sale price around 345 thousand, roughly two months of supply, median 16 days on market, mid 2026.')

h3('Guest supplied, do not state as verified')
bullet('', '1,000 plus career transactions. 150 plus team deals a year. 50 plus rental doors in five years. Top 1 percent nationwide. Seven figure business. No bought leads and no cold calling.')
bullet('', 'The 7 Figure Real Estate Academy and Day One Agent. Neither has a findable public page. Let her name and describe them.')
bullet('', 'Her own Instagram bio does say "Top 5 in Mid-MO" and "50+ doors investor," so the doors number is at least her own public claim.')

h3('Personal, publicly shared, use for warmth only')
bullet('', 'Psychology degree from the University of Missouri. Worked hospitality through high school and college, serving food and tending bar.')
bullet('', 'Failed the loan officer exam at Veterans United three times, each time by one point. Passed the real estate exam on the first try.')
bullet('', 'She and Jesse started Camacho Coffee in 2017. Read the landmine in 4H before you mention it.')
bullet('', 'Wife and mom. Instagram bio reads "Wife, Boy Mom, Real life + Real Estate."')

h3('Contact and social')
bullet('', 'Website waltersteamrealty.com. Email megan@thewalters.team. Phone 573.808.6457. Instagram @meganwalters (about 7,300 followers). YouTube The Walters Team. Facebook The Walters Team, eXp Realty. Link in bio beacons.page/meganwalters.')

h3('Connection to KIR')
bullet('', 'No prior KIR appearance found. First time on the show.')
bullet('', 'She is at eXp, so no Kale overlap and nothing to manage there. Sponsors Real Geeks and Courted are both lead and data products, and her whole thesis is that she does not buy leads. Do not let that turn into a defense of the sponsors.')

h3('THE CORE TOPIC')
p('She built a business she owns instead of one she rents. No bought leads, no cold calling, no desk hours, and then she turned the proceeds into doors she also owns. One topic, four angles: what replaced the desk, where the deals come from, who runs it besides her, and what she does with the money.', bold=True)

h3('Overasked questions to avoid')
bullet('', '"How did you get started at 22." Inside Real Estate already ran the coffee to closings origin arc. It is in your cold open instead.')
bullet('', '"How did you build confidence as a young agent." There is a published clip with that exact title.')
bullet('', '"What is it like working with your spouse." Every profile asks it and the answer is always charming and useless.')

h3('The "I have interviewed hundreds" moment')
p('Q5, and nowhere else: "I have interviewed hundreds of agents on this show and almost every one of them has bought leads at some point. You say you never have."', italic=True)

h3('Live stream title (paste into Restream before you hit record)')
p('Megan Walters Failed the Loan Officer Exam 3 Times. Then NAR Named Her 30 Under 30 (82 characters)', bold=True)
p('Backup: No Bought Leads, No Cold Calls: Megan Walters on Building in a Small Market (75 characters)', italic=True)
p('This is not the published episode title. Pick that after you hear the interview. Recommendation is in 2A.', italic=True)

h3('WATCH OUT FOR')
bullet('Rapid Fire will spend her two best answers. ', 'She sent both her best and worst advice in the intake. Worst advice, "you need to be at your desk 8 hours a day to succeed," is the entire spine of Block 1. Best advice, the one about letting people talk about you as you grow, is the emotional center of Q7. When she gives them in Rapid Fire, say "Love it, and we are coming back to that," then run the callbacks as written.')
bullet('Do not say the big numbers as fact. ', 'The 1,000 transactions, the 150 a year, the top 1 percent nationwide are all guest supplied. Say the 66 sides and the 17.6 million, then ask her to update it live.')
bullet('Camacho Coffee. ', 'A Columbia listing for it shows as closed. Do not use the present tense and do not congratulate her on it.')
bullet('Her open house story. ', 'It involves a stranger having a mental health crisis. If she tells it, let her tell it. Do not push for clinical detail and do not ask for the neighborhood.')

h3('THE TWO STANDING REMINDERS')
p('Ask the short version.', bold=True, space=2)
p('Count to three before you respond.', bold=True)

doc.add_page_break()

# =====================================================================
# 2. EPISODE FRAMEWORK
# =====================================================================
h1('2. Episode Framework')

h2('2A. Title Options')
table(['#', 'Title', 'Why it works'], [
    ['1', 'She Failed the Loan Officer Exam 3 Times. Then NAR Put Her in the 30 Under 30. (Megan Walters)',
     'Number, verified claim, name. Opens the curiosity gap with the failure and closes it with the award.'],
    ['2', '66 Sides at 28 With No Bought Leads: Megan Walters on Building in a Small Market',
     'Leads with the one hard number in the packet and the promise the largest audience segment came for.'],
    ['3', 'Megan Walters Was Told to Sit at a Desk 8 Hours a Day. She Did the Opposite.',
     'Her own words as the hook. Every agent who has heard that advice stops scrolling.'],
], widths=[0.35, 3.6, 3.1])
p('Recommended: #2. The goal for this episode is producing agents who act, and #2 is the only title that carries a verified number and the core promise in the same line. Run #1 as the live title because it works cold, before the interview has happened.', italic=True)

h2('2B. Cold Open Hook')
p('"Megan Walters failed the loan officer exam three times, each time by exactly one point, so she went and took the real estate exam instead and passed on her first try. Six years later the National Association of REALTORS named her one of its 30 Under 30, and she says she got there without buying a single lead. We are going to talk about that today. Stay tuned."', bold=True)

h2('2C. Episode Arc')
p('Core topic: She built a business she owns rather than rents, and then converted what it paid her into rental doors she also owns.', bold=True)
p('Why this topic: Her origin story is already on another podcast and her awards are on her website. What is not anywhere is the mechanism, where the deals actually come from when you have never bought a lead, who runs the business besides her, and what she does with the money. That is the part the largest slice of this audience cannot get anywhere else, and it works in any market of any size.')
p('Four angles, one topic. Block 1 kills the desk myth. Block 2 opens the engine. Block 3 shows the team that runs it. Block 4 shows where the money goes. Every block ends on something a listener can do this week.', italic=True)

doc.add_page_break()

# =====================================================================
# 3. INTERVIEW QUESTIONS
# =====================================================================
h1('3. Interview Questions')

h2('Rapid Fire (0:00 to 2:00, standardized, read as written)')
p('1. Best real estate advice you have ever received?', space=2)
p('2. Worst real estate advice you have ever received?', space=2)
p('3. One tool or app you cannot run your business without?', space=2)
p('4. What would surprise people most about your day to day?', space=2)
p('No follow ups. "Love it" and move on. One exception, below.', italic=True, space=4)
rich([('THE ONE EXCEPTION: ', True), ('She already sent you her best and worst advice, and both are load bearing later. After each, say: "Love it, and we are coming back to that." Then run the callback in Q1 (worst advice) and Q7 (best advice). Do not discuss either one here.', False, True)])

h2('BLOCK 1: The Desk Myth (2:00 to 12:00)')
p('Audience note: individual agents who still believe hours equal production. The largest segment, and the fastest place to earn their attention.', italic=True, space=2)
p('Arc: the number, the swap, the busywork, the next 90 minutes.', italic=True)

q(1,
  'Somebody early on told you the way to succeed in this business was to be at your desk eight hours a day. The year NAR named you 30 Under 30 you personally did 66 sides and 17.6 million dollars. When did you figure out the desk advice was wrong, and what did you do instead?',
  'Give me the year, and what you were doing with those eight hours instead.',
  'The moment she traded presence for process, with a date attached.',
  'Individual agents, new agents.',
  short='You did 66 sides at 28. When did you stop believing the desk advice?')

q(2,
  'Take me through yesterday. Not a typical day, yesterday. What time did you start, what were the first three things you actually touched, and what did you not touch at all?',
  'What time, and what was the actual task, not the category.',
  'The real operating rhythm, including what she has stopped doing.',
  'Individual agents, team leaders.',
  note='Do not accept "it depends" or a description of a typical day. Ask again for yesterday. This is the question agents will rewind.',
  short='Walk me through yesterday. First three things you touched, and what you ignored.')

q(3,
  'Every agent who builds systems builds a few that do nothing. What did you build, track, or pay for that felt productive and produced zero business?',
  'Name the tool or the habit, and what it cost you before you killed it.',
  'Credibility, plus permission for the listener to cancel something this week.',
  'Individual agents, team leaders.',
  short='What did you build that felt productive and produced nothing?')

q(4,
  'There is an agent listening who is four hours into a day at their desk with nothing to show for it. What are the three things you would have them do with the next 90 minutes instead?',
  'Give me the three, in order.',
  'Portable implementation with no tools and no budget.',
  'Individual agents, new agents.',
  short='Four hours at a desk, nothing to show. Next 90 minutes, what do they do?')

bridge('So the hours were never the thing, the system was. Which raises the obvious question. If you are not at the desk and you are not dialing, where is the business actually coming from.')

h2('BLOCK 2: Where the Deals Actually Come From (12:00 to 22:00)')
p('Audience note: the largest segment, and specifically the agent deciding this month whether to renew a lead contract.', italic=True, space=2)
p('Arc: the claim, the mechanism, the cost of being seen, the first 30 days.')

q(5,
  'I have interviewed hundreds of agents on this show, and almost every one of them has bought leads at some point. Most of them still do. You say you never have, and you do not cold call. So where does the business actually come from? Give me the rough percentages.',
  'Past clients, sphere, content, agent referrals, open houses, builders. What percent each?',
  'Whether this is a repeatable engine or just fifteen years of being known in a small town.',
  'Individual agents, team leaders, new agents.',
  note='This is the load bearing question of the episode. If she answers with a philosophy, ask again for the split. Do not move on without numbers. Say the "hundreds of agents" line here and nowhere else.',
  short='No bought leads, no cold calls. Where does the business come from? Percentages.')

q(6,
  'Your team describes what it does as innovating through media and content creation, and that is the part every agent copies badly. What do you actually post, how often, and how does a post become a signed listing agreement? Walk me through one that did.',
  'One specific post, what it said, and what happened after it.',
  'The content to contract path with a real example instead of an encouragement to be authentic.',
  'Individual agents.',
  note='If she gives you follower counts or view counts, redirect. You want the transaction, not the metric.',
  short='Walk me through one post that turned into a signed listing.')

q(7,
  'The best advice you ever got was about people talking about you as you grow, and letting them. In a town where everybody sees everything you post, that advice had to come from somewhere. What were they saying, and what did it cost you before you learned to let it go?',
  'Was it other agents, clients, or friends?',
  'The honest price of being visible in a small market, which is the fear stopping half the audience from posting anything.',
  'Individual agents, new agents.',
  permission='Tell me if this is too personal, but I want to ask about that advice you sent me.',
  note='This is the callback you promised in Rapid Fire. Ask it, then count to three and say nothing. This is the most likely clip in the episode.',
  short='What were people saying about you, and what did it cost you?')

q(8,
  'Say an agent just moved to a market of 100,000 people, they know nobody, and their lead contract ends this month. First 30 days, no bought leads. What do they do?',
  'Day one, day two. Be literal.',
  'A playbook that travels to any market, which is what makes this episode shareable outside Missouri.',
  'New agents, individual agents.',
  short='New market, no bought leads, no contacts. What are the first 30 days?')

bridge('That engine runs on you being known, which works right up until you are the only person who can do the work. You hired your first agent at 26.')

h2('BLOCK 3: The Team That Runs It Without Her (22:00 to 32:00)')
p('Audience note: team leaders, and the solo agent who is one good year away from their first hire.', italic=True, space=2)
p('Arc: the filter, the roles, the miss, the broker.')

q(9,
  'You have said you have to date somebody before you marry them when it comes to hiring. Autumn Lear was your first agent. What did dating her look like before you hired her, and how long did it take?',
  'How many months, and what specifically were you watching for?',
  'A hiring filter another team leader can copy on Monday.',
  'Team leaders, broker owners.',
  short='You date before you marry on hires. What did that look like with your first one?')

q(10,
  'A lead comes in at nine on a Tuesday morning. Who touches it first, what do they do with it, and at what point do you personally get involved?',
  'Name the seat, not the person.',
  'Whether this is an actual structure or her plus some helpers.',
  'Team leaders, individual agents.',
  short='Lead comes in Tuesday at nine. Who touches it, and when do you get involved?')

q(11,
  'Who did you hire that did not work out, and what did you miss during the dating phase?',
  'Was it skill, effort, or fit? And how long did you wait before you did something about it?',
  'The real cost of scaling, and the tell she now watches for.',
  'Team leaders, broker owners.',
  permission='You do not have to name names, and I am not asking you to.',
  short='Who did not work out, and what did you miss?')

q(12,
  'Flip this for the broker owners listening. You built your own brand inside somebody else\'s brokerage before you went to eXp. What did a broker actually do for you that mattered, and what did you have to build yourself because nobody was going to?',
  'One thing in each column. And what should a broker owner do this week for their top producer?',
  'Retention insight for broker owners and realistic expectations for every agent shopping brokerages.',
  'Broker owners, individual agents.',
  permission='I am not asking you to say anything unkind about anybody you have worked with.',
  short='What did your broker actually do for you, and what did you have to build alone?')

bridge('So that is the business. I want to spend the last stretch on what you do with what it pays you, because most agents never get to that part.')

h2('BLOCK 4: Turning Commission Into Doors (32:00 to 40:00)')
p('Audience note: any agent who has had a great year and has nothing to show for it but a nicer car.', italic=True, space=2)
p('Arc: the first door, the math now, the suspicion, the first check.')

q(13,
  'Your Instagram bio says 50 plus doors. Take me back to door number one. What was it, what did it cost, and where did the down payment come from?',
  'Was it a house, a duplex, what year? And did you represent yourself?',
  'The actual on ramp instead of the portfolio highlight reel.',
  'Individual agents, team leaders.',
  permission='Say as much or as little about the money as you want here.',
  short='Take me to door number one. What was it, and where did the money come from?')

q(14,
  'Columbia is sitting at about two months of supply and homes are going in roughly two weeks. In a market that tight, how does a rental deal still pencil for you, and where are you finding them if they are not on the MLS?',
  'Give me the numbers on the last one you bought.',
  'Whether the portfolio is a 2019 story or a 2026 one.',
  'Individual agents, all segments.',
  note='Verified as of mid 2026: median sale price around 345 thousand, roughly 2.2 months of supply, median 16 days on market. She will know the local number better than the aggregator does, so offer it and let her correct you.',
  short='Two months of supply, sixteen days on market. How does a rental still pencil?')

q(15,
  'Here is the pushback. There is an agent listening right now who represents buyers in a market where good rentals are scarce, and they are thinking, if Megan is buying doors, is she showing her clients the best deals or keeping them? Answer that agent directly. What is your rule?',
  'Written or unwritten? And has it ever cost you a deal you wanted?',
  'The ethical line, which is the thing that makes this entire block credible instead of a flex.',
  'All segments.',
  note='Deliver this flat, as the listener\'s question, not as an accusation. You are handing her the microphone to settle it, not putting her on trial.',
  short='An agent listening thinks you keep the best deals. What is your rule?')

q(16,
  'An agent closes twenty deals this year and ends up with nothing but a nicer car. What is the very first move you would have them make with the next commission check?',
  'How much of it, and into what?',
  'One transferable wealth habit the listener can start on their next closing.',
  'Individual agents, new agents.',
  short='Twenty deals, nothing to show for it. What do they do with the next check?')

h2('The Close (40:00 to 43:00)')

h3('Homework assignment (read verbatim)')
p('"Here is what I want you to do before the next episode. Open your last twenty closed transactions and write down where each one actually came from. By name, not by category. Then count how many came from somebody you already knew. That number tells you exactly what to fund next year and what to cancel. Not next month. This week."', bold=True)

h3('Guest close')
bullet('', '"Where can people find you and follow you?" Instagram @meganwalters is the main one. Website waltersteamrealty.com.')
bullet('', 'She teaches agents through The 7 Figure Real Estate Academy and Day One Agent. Neither has a page we could find, so let her say the names and describe them. Do not put numbers, prices, or results in her mouth.')
bullet('', 'Ask if there is anything free for listeners. If there is, get the exact link and say it twice.')

h3('Optional levity beat, only if the room needs it')
p('She once chased a cat around a subdivision for two hours believing she had let a seller\'s cat out, and the sellers did not own a cat. Two minutes, good clip, and it works as a re-hook if the energy dips between Block 3 and Block 4. Do not use it if you are at time.', italic=True)

h3('If you are running long, cut these first')
bullet('1. Q3, the busywork system. ', 'Q1 and Q2 already establish the desk thesis. This one deepens it, it does not carry it.')
bullet('2. Q12, the broker owner flip. ', 'Smallest audience segment, and it is the one question that pulls the conversation away from her business and toward brokerage shopping.')
bullet('3. Q9, the dating metaphor. ', 'Q10 gets you the team structure faster. Cut Q9 and open Block 3 on the Tuesday morning lead.')
bullet('4. Q14, whether the math still works. ', 'Block 4 survives on Q13 and Q15. Cut this before you cut either of those.')

h3('Never cut')
bullet('Q5. ', 'The lead source split. This is the episode.')
bullet('Q6. ', 'The post that became a listing. This is the proof for Q5.')
bullet('Q7. ', 'What people were saying about her. The human beat and the most likely clip.')
bullet('Q13. ', 'Door number one. The thing no other guest this quarter can offer.')
bullet('Q15. ', 'The conflict rule. Without it, Block 4 reads as a flex.')

doc.add_page_break()

# =====================================================================
# 4. RESEARCH BRIEF
# =====================================================================
h1('4. Research Brief')
p('Reference material. Read this the morning of, not during.', italic=True)

h2('4A. Background')
p('Columbia native, psychology degree from the University of Missouri. She worked hospitality through high school and college, serving food and tending bar, and a regular customer suggested real estate as a way to help people and pay off her student loans. She went at lending first, at Veterans United, and failed the loan officer exam three times, each time by a single point. She took the real estate exam instead and passed on the first try. She started at 22, launched Homes by Megan in 2017, later rebranded to The Walters Team and moved to eXp. She and her husband Jesse started Camacho Coffee the same year.')

h2('4B. Career Timeline (verified entries only)')
table(['Year', 'Role or milestone', 'Notable'], [
    ['2017', 'Founded Homes by Megan, Columbia MO', 'Same year she and Jesse launched Camacho Coffee, an air roasted bean company'],
    ['2018, 2019', 'Columbia Business Times Top of the Town, Top Salesperson', 'Two consecutive years'],
    ['2020', 'Columbia Board of REALTORS Salesperson of the Year', 'Also a COMO Magazine Women of Excellence finalist'],
    ['2021', 'COMO Magazine 20 Under 40', 'Lawrence Bulgin Award appears as 2020 on one listing and 2021 on another'],
    ['2022', 'NAR 30 Under 30 honoree, age 28, at House of Brokers Realty', '66 individual sides, 17.6 million dollars in individual volume'],
    ['Dec 2022', 'Opened an office next to the Tiger Hotel', 'Around the rebrand to The Walters Team'],
    ['Apr 2023', 'Team of four profiled by COMO Business Times', 'Megan, Jesse, Autumn Lear, Brittney Stone'],
    ['Current', 'The Walters Team at eXp Realty, 717 Cherry Street Suite C', 'Team site claims four ICON team awards and top 1 percent in mid-Missouri'],
], widths=[0.85, 3.1, 3.1])

h2('4C. What Makes Her Interesting for This Audience')
bullet('The small market proof: ', 'She hit 66 sides and 17.6 million in a college town, not a metro. Half this audience believes their market is the ceiling, and she is the counterexample with a number attached.')
bullet('The lead source claim: ', 'No bought leads and no cold calling, in a business where almost every guest has done both. If she can produce the percentage split, that single answer is worth the episode.')
bullet('The failure that starts it: ', 'Three failed loan officer exams, each by one point, before she ever sold a house. Nobody in this audience has a better reason to keep going.')
bullet('The doors: ', 'She converted agent income into a rental portfolio while still selling. Almost no producing agent guest can speak to both sides of that, and it is the question every top producer eventually asks.')
bullet('The visibility cost: ', 'She built on content in a town where everyone sees it, and the advice she quotes back is about people talking. That is the fear keeping most of this audience from posting anything at all.')

h2('4D. Key Data Points')
table(['Stat', 'Source', 'Confidence'], [
    ['66 sides, 17.6 million dollars individual volume', 'NAR 30 Under 30 profile, 2022', 'High'],
    ['NAR 30 Under 30, 2022, age 28, House of Brokers Realty', 'NAR 30 Under 30 profile', 'High'],
    ['CBOR Salesperson of the Year 2020', 'Team site, CBOR member page, Agent Pronto', 'High'],
    ['Founded Homes by Megan 2017, rebranded to The Walters Team', 'COMO Business Times, April 2023', 'High'],
    ['Failed loan officer exam three times, each by one point', 'COMO Business Times, April 2023', 'High, reported'],
    ['Team of four as of April 2023', 'COMO Business Times', 'Medium, likely out of date'],
    ['COMO 20 Under 40', 'Team site says 2021, another listing says 2020', 'Conflicting, do not state the year'],
    ['Columbia median sale price about 345 thousand, about 2.2 months supply, 16 days on market', 'Market aggregators, mid 2026', 'Medium, offer it and let her correct'],
    ['50 plus rental doors', 'Her own Instagram bio and her intake', 'Medium, self reported'],
    ['Top 5 in mid-Missouri', 'Her own Instagram bio', 'Medium, self reported'],
    ['Top 3 in volume and top 2 in sides by RealTrends', 'Team marketing copy, no year or market attached', 'Unverified, do not say it'],
    ['1,000 plus career transactions, 150 plus deals a year, top 1 percent nationwide', 'Guest bio only', 'Unverified, ask her live'],
    ['The 7 Figure Real Estate Academy, Day One Agent', 'Guest bio only, no public page found', 'Unverified, let her describe them'],
], widths=[3.0, 2.7, 1.4])

h2('4E. Previous Media Appearances')
bullet('Inside Real Estate, Episode 44. ', '"From Coffee to Closings: Megan Walters Rise in Real Estate." The origin arc, hospitality into real estate. There is a separate published clip titled "How Megan Walters Built Real Confidence as a Young Agent."')
bullet('COMO Business Times, April 2023. ', 'The team profile. Career timeline, the hiring philosophy, the office move, the growth stance.')
bullet('NAR 30 Under 30 profile, 2022. ', 'The production numbers and the hospitality quote.')
bullet('Overasked, avoid: ', 'the origin story, the young agent confidence question, and what it is like working with her husband. All three are already published and all three will eat ten minutes.')

h2('4F. Their Own Words')
table(['Quote', 'Where and when', 'Confidence', 'How you use it'], [
    ['"If they are talking about you as you grow, let them. It means they are noticing you and you are doing something right. Your grass is green, focus there."',
     'Her intake form to you, 2026', 'Verbatim',
     'Read it back at the top of Q7, then ask what it cost her. This is the emotional center of the episode.'],
    ['"You need to be at your desk 8 hours a day to succeed in real estate."',
     'Her intake, quoted as the worst advice she received', 'Verbatim',
     'Open Q1 with it. It is the thesis of Block 1 and she handed it to you.'],
    ['"I was so stubborn, I just went and took the real estate exam and passed on my first try."',
     'COMO Business Times, April 2023', 'Reported',
     'Paraphrase out loud, do not quote it word for word. Use it in the cold open or as the setup to Q1.'],
    ['"I am going to have to date you before I marry you."',
     'COMO Business Times, April 2023, on hiring', 'Reported',
     'Opens Q9. Say "you have said something like" and let her deliver her own version.'],
    ['"My goal is just to continue to grow, not to be the biggest team by any means, just really creating a family that wants to help each other succeed."',
     'COMO Business Times, April 2023', 'Reported',
     'Hold in reserve. If she describes aggressive growth now, read it back and ask her to reconcile it.'],
    ['"It teaches you to put your humility aside to serve another, even if it is not convenient for you at the time."',
     'NAR 30 Under 30 profile, on hospitality work', 'Reported',
     'Hold in reserve for Q4 or Q8, if the answers get abstract and you need to pull her back to the work.'],
], widths=[2.5, 1.6, 0.9, 2.1])

h2('4G. Audience Relevance')
table(['Segment', 'What they get'], [
    ['Individual agents', 'A lead source split from someone who has never bought a lead, and a content to contract path with a real example.'],
    ['Team leaders', 'A hiring filter, a lead routing structure by seat, and the hire that did not work out.'],
    ['Broker owners', 'What a top producer actually needed from her broker, and what she built herself when nobody did.'],
    ['New agents', 'A 30 day plan for a small market with no contacts and no budget, and the first thing to do with a commission check.'],
], widths=[1.4, 5.7])

h2('4H. Landmines')
bullet('Camacho Coffee. ', 'A Columbia listing for it currently shows as closed. We could not confirm whether the roasting business ended or just a location. Do not use the present tense, do not congratulate her on it, and if it comes up ask an open question and let her set the tense.')
bullet('The academy and Day One Agent. ', 'No findable public page for either. Do not quote enrollment, price, or student results. Let her name them in the close.')
bullet('The big production numbers. ', '1,000 plus career transactions, 150 plus a year, top 1 percent nationwide, seven figures. All guest supplied. The verified public number is 66 sides and 17.6 million in her 30 Under 30 year. Say the verified one and ask her to update it live.')
bullet('Brokerage history. ', 'NAR listed her at House of Brokers Realty in her honor year while she was already running her own brand. Do not assert she has always been at eXp and do not assert the year she moved.')
bullet('Award years. ', '20 Under 40 shows as 2021 on her team site and 2020 elsewhere. Lawrence Bulgin has the same problem. Say "a couple of years back."')
bullet('RealTrends ranking. ', '"Top 3 in volume and top 2 in sides" appears in her marketing with no year and no market attached. Do not state it.')
bullet('The open house story. ', 'A neighbor having a mental health crisis, which ended with her locking down a house full of people. If she raises it, let her tell it her way. Do not ask for the neighborhood, do not ask what was wrong with the woman, and do not play it for laughs.')

h2('Drift Guardrail')
p('The drift on this episode is the origin story. Two of her three prior appearances are the hospitality to real estate arc and she will land there without being asked. It is already spent in your cold open. If she starts telling it, the line is: "I put that in the open so people have it. What I want is the version of you that exists now."', italic=True)

h2('Green Room Brief (two minutes before you record)')
bullet('', 'Ask her to confirm four things on tape so you can use them: career transaction count and how she counts it, the team\'s deals last year, current door count, and whether she is still at eXp.')
bullet('', 'Tell her Rapid Fire is going to take her best and worst advice in the first two minutes, that you will not follow up there, and that you are coming back to both later.')
bullet('', 'Ask whether the open house story is fair game, and whether the cat story is.')
bullet('', 'Ask what she wants the academy and Day One Agent called on air, and whether there is a link for listeners.')
bullet('', 'Ask about Camacho Coffee off air so you know the tense before you are live.')

doc.add_page_break()

# =====================================================================
# 5. LIVE STREAM TITLE, DESCRIPTIONS, HASHTAGS
# =====================================================================
h1('5. Live Stream Title, Descriptions and Hashtags')

h2('5A. Live Stream Title')
p('Live stream title: Megan Walters Failed the Loan Officer Exam 3 Times. Then NAR Named Her 30 Under 30 (82 characters)', bold=True)
p('Backup: No Bought Leads, No Cold Calls: Megan Walters on Building in a Small Market (75 characters)', italic=True)
p('Both work cold, before the interview happens, and both use only verified facts. This does not have to match the published episode title. Pick that one after you hear the interview, and 2A has the recommendation.', italic=True)

h2('5B. Platform Descriptions')

h3('Facebook Live')
p('Megan Walters failed the loan officer exam three times before she ever sold a house. Today she runs one of the top teams in mid-Missouri, and she says she has done it without ever buying a lead or making a cold call. We are getting into where the business actually comes from, who runs it besides her, and what she does with the money. Drop your questions in the comments!')

h3('Instagram Live')
p('66 sides at 28. Zero bought leads. Megan Walters on how she really gets clients. #RealEstateAgent #RealtorLife #KeepingItReal')

h3('TikTok Live')
p('She failed the loan officer exam 3 times. Now she is a NAR 30 Under 30 with 50 plus rental doors. #realtor #realestate #realestateagent')

h3('YouTube Live')
p('Megan Walters, owner of The Walters Team at eXp Realty in Columbia, Missouri, and a National Association of REALTORS 30 Under 30 honoree, joins the Keeping It Real Podcast. We cover how she generates business without buying leads or cold calling, how she structured her team, and how she turned commission income into rental property. Real estate agent training, lead generation, team building and real estate investing, on the Keeping It Real Podcast with D.J. Paris.')

h3('LinkedIn Live')
p('Most agents assume production is a function of hours and lead spend. Megan Walters built a top producing team in a mid-Missouri market with neither. We break down her actual lead sources by percentage, how her team routes and closes them, and how she converted agent income into a rental portfolio.')

h2('5C. Hashtag Sets')
bullet('Universal: ', '#KeepingItReal #RealEstatePodcast #DJParis #RealtorLife #RealEstateAgent')
bullet('Episode specific: ', '#LeadGeneration #RealEstateTeam #ColumbiaMO #MidMissouri #RealEstateInvesting #30Under30 #SmallMarketAgent')
bullet('Guest tags: ', 'Instagram @meganwalters. Facebook The Walters Team, eXp Realty. YouTube The Walters Team. Website waltersteamrealty.com.')

doc.add_page_break()

# =====================================================================
# 6. YOUTUBE CHAPTER MARKERS
# =====================================================================
h1('6. YouTube Chapter Markers')
p('Estimates. Adjust after recording. Every title is written to be searchable on its own.', italic=True)
table(['Timestamp', 'Chapter title'], [
    ['0:00', 'She Failed the Loan Officer Exam 3 Times, Each Time By One Point'],
    ['[post-ads]', 'Rapid Fire: Best and Worst Real Estate Advice She Ever Got'],
    ['', '"Be at Your Desk 8 Hours a Day" and Why It Is Wrong'],
    ['', 'What a Top Producer Actually Did Yesterday, Hour By Hour'],
    ['', 'The System She Built That Produced Zero Business'],
    ['', 'Four Hours In With Nothing to Show: What to Do With the Next 90 Minutes'],
    ['', 'Never Bought a Lead: Where the Business Actually Comes From, By Percentage'],
    ['', 'The Social Post That Turned Into a Signed Listing'],
    ['', 'What People Said About Her as She Grew, and What It Cost'],
    ['', 'New Market, No Contacts, No Bought Leads: The First 30 Days'],
    ['', 'Date Them Before You Marry Them: How She Vetted Her First Hire'],
    ['', 'A Lead Comes In Tuesday at 9am: Who Touches It First'],
    ['', 'The Hire That Did Not Work Out and the Tell She Missed'],
    ['', 'What a Broker Actually Did For Her, and What She Had to Build Alone'],
    ['', 'Rental Door Number One: What It Cost and Where the Down Payment Came From'],
    ['', 'Two Months of Supply: How a Rental Still Pencils in 2026'],
    ['', 'Does She Keep the Best Deals From Her Own Buyers? Her Rule'],
    ['', '20 Deals and Nothing to Show For It: What to Do With the Next Check'],
    ['', 'Homework, and Where to Find Megan Walters'],
], widths=[1.1, 6.0])

doc.add_page_break()

# =====================================================================
# 7. STRESS TEST, COUNCIL, EP POLISH
# =====================================================================
h1('7. Stress Test, Council Review and EP Polish')

h2('7A. Stress Test (pass 2, part one)')
table(['#', 'What broke', 'Fix applied'], [
    ['1', 'The bio claims 1,000 plus career transactions. The only verified public production number is 66 sides in her 30 Under 30 year. The two cannot both anchor the episode.',
     'Built all packaging on the verified 66 sides and 17.6 million. Moved the career count into the green room brief as something she confirms on tape.'],
    ['2', 'The 7 Figure Real Estate Academy and Day One Agent have no findable public page. Nothing about either could be verified.',
     'Removed both from the interview body entirely. They appear only in the close, where she names and describes them herself.'],
    ['3', 'Camacho Coffee is in every profile of her, and a Columbia listing for it now shows as closed.',
     'Pulled the coffee business out of every question, added a landmine with an instruction to never use the present tense, and added an off air check to the green room brief.'],
    ['4', 'Rapid Fire spends her best and worst advice in minute one, and both are load bearing later.',
     'Wrote the deflection line in three places, built the worst advice callback into Q1 and the best advice callback into Q7.'],
    ['5', 'Q5 was dodgeable. She could answer "relationships and referrals" and never give a number.',
     'Rewrote it to demand percentages and added a producer note telling D.J. not to move on without the split.'],
    ['6', 'Q6 was dodgeable in the other direction. She could answer with follower counts.',
     'Rewrote it around one specific post that produced one specific signed listing, with a redirect note.'],
    ['7', 'Drift risk into the origin story. Two of her three prior appearances are that arc.',
     'Spent the origin story in the cold open so it is used before the interview starts, and wrote a guardrail line.'],
    ['8', 'The investor question read as an ambush in the draft.',
     'Reframed Q15 around a described listener asking the question, with a note to deliver it flat.'],
    ['9', 'Runtime. The draft ran five blocks and twenty questions against 43 minutes.',
     'Cut to four blocks and sixteen questions, and deleted a standalone small market block that was really the same argument as Block 2.'],
    ['10', 'Two award years conflict across sources, and the RealTrends ranking has no year or market attached.',
     'Kept all three out of the read. Added them to Landmines with the safe on air phrasing.'],
    ['11', 'Her unusual story involves a stranger in mental health crisis and could be played for laughs on air.',
     'Added an explicit landmine, and moved the cat story into the close as the designated levity beat instead.'],
], widths=[0.35, 3.3, 3.4])

h2('7B. Council Review (pass 2, part two)')

h3('Member notes')
table(['Member', 'What they would change'], [
    ['Alex Hormozi', 'Block 4 is the smallest audience and it is getting a quarter of the runtime. Earn it or shrink it. Q16 is the only thing in there a normal agent can act on.'],
    ['MrBeast', 'The sag is right at the handoff from her engine to her team. Nobody cares about her org chart yet. Re-hook with the fact that she hired at 26, not with a transition sentence.'],
    ['Brendan Kane', 'The failed exam is the hook, and it is a stat plus a personality ingredient in one line. Your first title buried it behind her name. Lead with the failure.'],
    ['Eric Simon, The Broke Agent', 'Q7 is the one an agent forwards to another agent. Every one of them has felt watched in their market. Do not soften it and do not rush past the silence.'],
    ['Chris Do', 'Q13 is where she either gets real about money or gives you a portfolio brag. The permission clause is what decides which one you get.'],
    ['Donald Miller', 'The hero has to be the agent listening, not Megan. Every block already ends on an implementation question. Keep it that way and the grunt test passes.'],
    ['Chip and Dan Heath (witness)', 'Called on the curiosity claim. The failed exam opens the gap and the 30 Under 30 closes it, in that order. Do not front load the award.'],
    ['Jonah Berger (witness)', 'Called on the share claim. High arousal is in the visibility question and the ethics question, not in the systems questions. Retold in one sentence, "she has never bought a lead" survives intact.'],
], widths=[1.7, 5.4])

h3('Title')
table(['#', 'Title', 'Ingredient', 'Curiosity mechanism'], [
    ['1', 'She Failed the Loan Officer Exam 3 Times. Then NAR Put Her in the 30 Under 30. (Megan Walters)', 'Stat plus personality', 'Gap opens on the failure, closes on the award, in that order'],
    ['2', '66 Sides at 28 With No Bought Leads: Megan Walters on Building in a Small Market', 'Stat plus insight', 'The number creates the question, "no bought leads" makes it urgent'],
    ['3', 'Megan Walters Was Told to Sit at a Desk 8 Hours a Day. She Did the Opposite.', 'Insight', 'Names advice the listener has personally received, then contradicts it'],
], widths=[0.3, 3.2, 1.3, 2.3])
p('Recommended: #2 to publish. The goal for this episode is producing agents who change something, and #2 is the only one carrying a verified number and the core promise in the same line. #1 runs live, because it works cold before the interview exists. #3 is the A/B for the Shorts cut.', italic=True)

h3('Cold open (sharpened)')
p('"Megan Walters failed the loan officer exam three times, each time by exactly one point, so she went and took the real estate exam instead and passed on her first try. Six years later the National Association of REALTORS named her one of its 30 Under 30, and she says she got there without buying a single lead. We are going to talk about that today. Stay tuned."', bold=True)

h3('The Clip Engine')
table(['Q#', 'Question', 'Berger emotion', 'Heath gap'], [
    ['7', 'What were people saying about you, and what did it cost you?', 'High arousal, anger and vindication. Social currency for any agent afraid to post.', 'She quoted the advice in her intake. She has never said publicly what caused it.'],
    ['15', 'An agent listening thinks you keep the best deals. What is your rule?', 'Anxiety and mild controversy, the highest sharing emotion in the episode.', 'Names the suspicion the listener already had, then closes it in her own words.'],
    ['13', 'Door number one. What was it, and where did the money come from?', 'Awe plus practical value, the combination that gets saved rather than just watched.', 'Everyone has seen the 50 doors. Nobody has seen the first one.'],
], widths=[0.35, 2.4, 2.3, 2.05])

h3('Live description scrub')
table(['Platform', 'Verdict and fix'], [
    ['Facebook', 'Keep. Leads on the failure, ends on the comment prompt, no jargon.'],
    ['Instagram', 'Keep. Three fragments and out. Any longer and it truncates before the hook lands.'],
    ['TikTok', 'Rewritten so it is not the Facebook line at a shorter length. Leads on the number and the doors, which is what performs there.'],
    ['YouTube', 'Keep. Carries the searchable terms, agent training, lead generation, team building, real estate investing, without reading like a keyword dump.'],
    ['LinkedIn', 'Rewritten to open on the assumption it breaks rather than on the guest. That is the only version that stops a business feed.'],
], widths=[1.1, 6.0])

h3('Arc fix')
p('The one place this sags is the handoff out of Block 2. She has just given a personal answer about being talked about, and the next block is org structure, which is a cold start. The bridge line does the re-hook by naming that she hired her first agent at 26, so the block opens on a decision rather than on a chart. If it still sags on tape, drop the cat story in there instead of in the close.')

h3('Why it should work')
bullet('Curiosity mechanism (Heath): ', 'The failed exam opens the gap before the award closes it, and the packaging never front loads the punchline.')
bullet('Share driver (Berger): ', 'Two high arousal beats, being talked about and the ethics question, both sitting inside an episode otherwise built on practical value.')
bullet('Retention move (MrBeast): ', 'Every block ends on a do it tomorrow question, so there is no natural place to stop listening.')

h3('The dissent')
p('Hormozi is still objecting that Block 4 serves the smallest slice of the audience for a quarter of the runtime. The experiment for the next episode of this type: cut the investing block to two questions, door number one and the ethics rule, and give the recovered time back to the lead source block. If the retention graph on this one drops after the 32 minute mark, he was right.')

h2('7C. EP Polish (pass 3)')
bullet('', 'Rebuilt every title, the cold open, the live stream title, and all five platform descriptions on the verified 66 sides and 17.6 million, and stripped the 1,000 transactions, the 150 a year, and the top 1 percent nationwide out of all packaging.')
bullet('', 'Cut the draft from five blocks and twenty questions to four blocks and sixteen, and deleted the standalone small market block, which was making the same argument as Block 2 with worse questions.')
bullet('', 'Moved the entire origin story into the cold open, because it is the one place it cannot eat interview time and it is already published on another show.')
bullet('', 'Rewrote Q5 to demand percentages and Q6 to demand one specific post and one specific listing, then added producer notes on both telling D.J. not to accept the substitute.')
bullet('', 'Reframed Q15 from a direct challenge into a described listener\'s pushback, which turns the one possible ambush in the episode into a shared problem she gets to settle.')
bullet('', 'Added the Rapid Fire deflection line in three places, the Quick Reference Card, the Rapid Fire section, and the green room brief, because both of her intake answers are load bearing and the format will spend them in minute one.')
bullet('', 'Wrote permission clauses on exactly four questions, Q7, Q11, Q12, and Q13, all of which touch money, a former employer, or something personal, and removed them everywhere else so the device does not become a tic.')
bullet('', 'Added the green room brief, because four separate load bearing facts are guest supplied and the only real fix is getting her to own them on tape before the interview starts.')
bullet('', 'Moved the cat story out of the body and into the close as an optional levity beat, and added an explicit instruction not to play the open house story for laughs.')
bullet('', 'Wrote all nineteen chapter markers as standalone searchable headlines rather than block labels.')
bullet('', 'Rewrote the TikTok and LinkedIn descriptions so they are not the Facebook one at a different length.')
bullet('', 'Put Q12 second in the cut order and ended Block 3 on it anyway, so the block still lands on an implementation question if it survives.')
bullet('', 'Added the drift guardrail line, since her published media is almost entirely the origin arc and she will return to it unprompted.')
bullet('', 'Added the Camacho Coffee landmine after finding a Columbia listing marked closed, because it appears in every profile of her and would otherwise get congratulated on air.')
bullet('', 'Checked every short version against the twenty word ceiling. Q1, Q4, and Q8 had to be rewritten to survive the compression.')
bullet('', 'Swept the document for em dashes, curly quotes, and AI-speak. Zero.')

doc.save("/Users/djparis/GitHub Projects/keeping-it-real-content-system/guest-prep/Megan_Walters_Interview_Prep.docx")
print('saved')
