#!/usr/bin/env python
# Builds the KIRP interview prep .docx for Chris Wands, The Wands Team at Douglas Elliman (Miami)
# Built to prompts/06_interview_prep.md v5: research -> draft -> stress test + council -> EP polish
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)

for name, size, color in (('Heading 1', 18, 0x1F3864), ('Heading 2', 14, 0x2E5496), ('Heading 3', 11.5, 0x2E5496)):
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor((color >> 16) & 255, (color >> 8) & 255, color & 255)
    st.font.bold = True

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.8)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def p(text='', bold=False, italic=False, size=None, space=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space)
    r = par.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    return par


def rich(segments, style=None, space=6):
    par = doc.add_paragraph(style=style)
    par.paragraph_format.space_after = Pt(space)
    for seg in segments:
        text = seg[0]
        b = seg[1] if len(seg) > 1 else False
        i = seg[2] if len(seg) > 2 else False
        r = par.add_run(text)
        r.bold, r.italic = b, i
    return par


def bullet(text_bold, text_rest=''):
    par = doc.add_paragraph(style='List Bullet')
    par.paragraph_format.space_after = Pt(2)
    if text_bold:
        par.add_run(text_bold).bold = True
    if text_rest:
        par.add_run(text_rest)
    return par


def quoted(label, text):
    par = doc.add_paragraph(style='Intense Quote')
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.space_before = Pt(3)
    r = par.add_run(label + ' ')
    r.bold = True
    r.italic = True
    r2 = par.add_run(text)
    r2.italic = True
    return par


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            par = cells[i].paragraphs[0]
            run = par.add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def q(num, question, if_vague, reveals, serves, note=None, short=None, permission=None):
    rich([(str(num) + '. ', True), (question, True)], space=2)
    if short:
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(3)
        par.paragraph_format.left_indent = Inches(0.25)
        r = par.add_run('SAY THIS: ')
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        r2 = par.add_run('"' + short + '"')
        r2.bold = True
        r2.font.size = Pt(10.5)
    if permission:
        quoted('First, say this:', '"' + permission + '"')
    quoted('If vague, ask:', if_vague)
    quoted('Ideal answer reveals:', reveals)
    quoted('Serves:', serves)
    if note:
        rich([('PRODUCER NOTE: ', True), (note, False, True)], space=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def bridge(text):
    rich([('BRIDGE TO NEXT BLOCK (read as written): ', True), ('"' + text + '"', False, True)])


# =====================================================================
# TITLE
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('KEEPING IT REAL PODCAST')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Interview Prep: Chris Wands')
r.bold = True
r.font.size = Pt(20)
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Founder, The Wands Team at Douglas Elliman  |  Miami Beach, Florida  |  Target runtime 45 minutes')
r.italic = True
r.font.size = Pt(10)

# =====================================================================
# PAGE 1: QUICK REFERENCE CARD
# =====================================================================
h1('1. Quick Reference Card')
p('One page. Glance at this during the interview.', italic=True, space=8)

h3('Who he is')
bullet('Name: ', 'Christopher M. Wands, goes by Chris. Executive Director of Luxury Sales, Douglas Elliman. Founder of The Wands Team.')
bullet('Based: ', 'Miami Beach, Florida. Originally Port Washington, New York.')
bullet('Markets: ', 'Miami Beach, Fisher Island, Coconut Grove, Brickell, Downtown, Coral Gables, Boca Raton, Delray Beach, Fort Lauderdale. Team also works a New York City referral corridor.')
bullet('Guest type: ', 'A. Producing team leader. Push for systems, daily habits, specific numbers.')

h3('Verified numbers (safe to say on air)')
bullet('2016: ', 'Entered the business with Douglas Elliman. Won Douglas Elliman South Florida Rookie of the Year in his first year.')
bullet('2022: ', 'Founded The Wands Team. Nine people as of the May 2025 Inman Teams Spotlight.')
bullet('Elliman national team rankings: ', 'No. 4 by rental transactions, No. 7 by GCI, No. 8 by volume, No. 10 by transactions. No. 11 medium team nationwide by GCI. Pinnacle Award, top three percent of the company.')
bullet('RealTrends Verified: ', 'No. 7 team in Miami by volume, No. 67 in Florida. Average sale price $1.87 million.')
bullet('New development: ', 'Sales executive on Five Park Miami Beach. Team also works The Perigon, The Standard Residences, Opus Coconut Grove, Park Grove.')

h3('Personal, publicly shared, use for warmth only')
bullet('', 'Fourth generation real estate. Great grandparents and grandparents were developers and contractors.')
bullet('', 'Moved to Miami in 2012 for the University of Miami and never left.')
bullet('', 'Competitive sailboat racer as a kid. Traveled constantly for it.')
bullet('', 'First sales job out of school was at a water sports company, where he became the top salesperson almost immediately.')

h3('Contact and social')
bullet('', 'Website movetofl.com. Phone 561.420.9610. Instagram @chris_wands and @thewandsteam. LinkedIn Christopher M. Wands. YouTube The Wands Team.')

h3('Connection to KIR')
bullet('', 'No prior Keeping It Real appearance found. No podcast appearances found anywhere, which means this is likely his first real long form interview. That is an advantage. The stories are not rehearsed.')

h3('THE CORE TOPIC')
p('How the deals nobody else wanted, rentals and portal leads and small transactions, became the acquisition engine underneath a luxury business.', bold=True)

h3('Overasked questions to avoid')
bullet('', 'What is happening in the South Florida market right now. He has answered this in citybiz and Inman and it does not travel to a Chicago listener.')
bullet('', 'Why is everyone moving from New York to Florida. Every Miami guest gets this. It is a headline, not a system.')
bullet('', 'How did you get started in real estate. Voyage MIA already ran the full origin story.')

h3('The "I have interviewed hundreds" moment')
p('Use it once, at Q6, exactly like this: "I have interviewed hundreds of agents and almost every one of them tells me rentals are a loss leader they tolerate. You ranked nationally in them."', italic=True)

h3('Live stream title (paste into Restream before you hit record)')
rich([('Primary: ', True), ('Chris Wands: The Miami Luxury Team That Ranks #4 in Rentals, and Why That\'s the Point', False, True), (' (85 characters)', False, True)])
rich([('Backup: ', True), ('Chris Wands on the Miami Deals Every Luxury Agent Turns Down', False, True), (' (60 characters)', False, True)])
p('This does not have to match the published episode title. Pick that after you hear the interview.', italic=True, space=8)

h3('WATCH OUT FOR')
bullet('Rapid Fire will spend his best material. ', 'He sent his best and worst advice in the intake. Worst was "stay in your lane," and that answer is load bearing for Block 2. When he says it, say "Love it, and we are coming back to that." Then run the callback at Q7.')
bullet('The billionaire lead story is load bearing too. ', 'He sent it as his most unusual experience. Do not let it come out as a throwaway anecdote in minute two. It is Q3 and it is the proof for the whole core topic.')
bullet('Three guest supplied numbers. Do not state any of them as fact. ', 'The $55 million Perigon sale, the 35 properties in year one, and the $1 billion in career sales. None of the three appear in any independent source. Ask him to confirm each on air.')
bullet('Two conflicting volume numbers exist. ', 'Inman reported 48.6 sides and $100.3 million. RealTrends Verified shows $60.3 million and 32.3 sides. Different periods, probably, but do not pick one. Ask him what a normal year looks like now.')
bullet('Drift risk: the Miami market lecture. ', 'Three of his five suggested topics are market commentary. If the conversation slides there, use the guardrail at the bottom of Section 4.')

h3('THE TWO STANDING REMINDERS')
p('Ask the short version. Count to three before you respond.', bold=True)

doc.add_page_break()

# =====================================================================
# PAGE 2: EPISODE FRAMEWORK
# =====================================================================
h1('2. Episode Framework')

h2('2A. Title Options')
table(['#', 'Title', 'Why It Works'],
      [['1', 'The #4 Rental Team at Douglas Elliman Is a Luxury Team. Chris Wands on Why Small Deals Win',
        'Number plus a contradiction. A luxury team ranking nationally in rentals is the whole episode in one line.'],
       ['2', 'They Handed Him the Scraps. Now He Runs One of Elliman\'s Top Teams. (Chris Wands)',
        'His own words, a status reversal, and his name. Curiosity gap opens before it closes.'],
       ['3', 'The Miami Luxury Broker Who Tells Clients to Rent (Chris Wands, The Wands Team)',
        'Specific claim that costs the guest something. Name and company for search.']],
      widths=[0.4, 3.4, 3.2])
p('Recommended: #1. It carries the only verified number that is also a contradiction.', italic=True)

h2('2B. Cold Open Hook')
p('Read before the ads:', italic=True, space=3)
p('"Inside one of the largest luxury brokerages in the country, there is a Miami team that ranks number four nationally in rental transactions, and its founder says that ranking is the reason he has a luxury business at all. We are going to talk about that today. Stay tuned."', bold=True)

h2('2C. Episode Arc')
rich([('Core topic: ', True), ('How the deals nobody else wanted, rentals and portal leads and small transactions, became the acquisition engine underneath a luxury business.', False, True)])
rich([('Why this topic: ', True), ('It is the one thing about him that is both verified in the data and completely unasked everywhere else, and unlike a South Florida market update it works identically for an agent in Chicago, Cleveland, or Boise.', False, True)])
p('Structure: Rapid Fire, then four topic blocks of roughly ten minutes each, then the close. Sixteen numbered questions total. That is the ceiling for 45 minutes with real answers.', space=8)

table(['Segment', 'Time', 'Purpose'],
      [['Rapid Fire', '0:00 to 2:00', 'Four standardized questions. No follow ups. Deflection line ready on question two.'],
       ['Block 1: The Scraps Years', '2:00 to 12:00', 'What he actually did with the deals nobody wanted.'],
       ['Block 2: The Rental Engine', '12:00 to 22:00', 'The national rental ranking, and the system that converts renters into buyers.'],
       ['Block 3: Telling a Buyer Not to Buy', '22:00 to 33:00', 'The advisory line, priced against the new cost of condo ownership.'],
       ['Block 4: Scaling Without Breaking It', '33:00 to 42:00', 'Solo to nine. What he hands a new agent and what they have to earn.'],
       ['The Close', '42:00 to 45:00', 'Homework, where to find him.']],
      widths=[2.2, 1.3, 3.5])

doc.add_page_break()

# =====================================================================
# SECTION 3: INTERVIEW QUESTIONS
# =====================================================================
h1('3. Interview Questions')

h2('Rapid Fire (0:00 to 2:00, standardized, read as written)')
p('Four questions. One to two sentence answers. No follow ups. You say "Love it" and move on.', italic=True, space=6)
bullet('1. ', 'Best real estate advice you have ever received?')
bullet('2. ', 'Worst real estate advice you have ever received?')
bullet('3. ', 'One tool or app you cannot run your business without?')
bullet('4. ', 'What would surprise people most about your day to day?')
p('')
rich([('DEFLECTION LINE, question 2 only: ', True), ('He is going to say "stay in your lane." That is the spine of Block 2. Say "Love it, and we are coming back to that," then move to question 3. Do not follow up. The callback is Q7.', False, True)])
rich([('What he already sent you: ', True), ('Best was bet on yourself, trust your gut, and do not focus on the money, focus on helping people first. Worst was stay in your lane. You already have both answers, so listen for what he adds live rather than for the answer itself.', False, True)])

h2('BLOCK 1: The Scraps Years (2:00 to 12:00)')
p('Audience note: individual agents and new agents. This is the block where a two year agent decides whether to keep listening.', italic=True)
p('Arc: the scrap, the method, the proof, the rule.', italic=True, space=8)

q(1,
  'When you started at Elliman you have said that people would hand you their scraps. I do not want the philosophy, I want the actual scrap. What was the first one somebody handed you, and what did you do with it that the person who handed it off would not have done?',
  'What was the price point, and how long did it take you?',
  'The specific behavior, not the attitude. Speed of response, showing up in person, doing the paperwork somebody else did not want to do.',
  'Individual agents, new agents',
  short='You said people handed you their scraps. What was the first one?')

q(2,
  'Walk me through what actually happens on a twenty five hundred dollar a month rental in your world. Start to finish. Who does the showings, how many hours does it eat, and what goes into the file after it closes?',
  'How many hours, and what exactly goes in the file after closing?',
  'The mechanism. If the answer is only about service, push for the hours and the data capture. The file is the whole point.',
  'Individual agents, team leaders',
  note='He can answer this one entirely on feel. Do not accept feel. You want a number of hours and a description of what gets recorded.',
  short='Walk me through a $2,500 rental start to finish. Who does what, how many hours?')

q(3,
  'You once got a lead off Realtor.com that turned out to be a billionaire. Before you knew who he was, when it was just another portal lead in your inbox, what did you do with it that most agents do not do with portal leads?',
  'How fast do you answer a portal lead, and who on the team answers it?',
  'The response protocol. He told us the sale did not close, which makes this better, because the lesson is the process and not the win.',
  'Individual agents, new agents',
  note='He sent this as his most unusual experience so he will want to tell it as a story. Let him tell it, then pull him back to the protocol. The protocol is the takeaway.',
  short='A Realtor.com lead turned out to be a billionaire. What did you do with it first?')

q(4,
  'Somebody is listening right now who just got handed a two thousand dollar rental by their team lead and feels a little insulted by it. What is the exact thing you want them to do with it this week?',
  'What do they say to that client, and what do they write down?',
  'A do it tomorrow instruction. One action, not a mindset.',
  'New agents, individual agents',
  short='An agent just got handed a $2,000 rental. What do they do with it this week?')

bridge('So the small deal is not the deal, it is the entry point, and you built a whole ranking out of that idea, which is where I want to go next.')

h2('BLOCK 2: The Rental Engine (12:00 to 22:00)')
p('Audience note: individual agents first, team leaders second. This is the contrarian core of the episode.', italic=True)
p('Arc: the contradiction, the number, the callback, the minimum system.', italic=True, space=8)

q(5,
  'Inside a luxury brokerage, your team ranks number four nationally by rental transactions. Most luxury agents I talk to treat rentals as something you do for a friend and then never again. You built a national ranking out of them. Why?',
  'Was that a decision or an accident? Walk me through the moment you decided to lean in.',
  'Whether the rental volume is strategy or leftover. Either answer is good radio, but you want to know which.',
  'Individual agents, team leaders',
  short='Your team is #4 at Elliman in rentals. Why chase rentals at all?')

q(6,
  'I have interviewed hundreds of agents and almost every one of them tells me rentals are a loss leader they tolerate. You ranked nationally in them. Of the renters your team placed in the last two years, how many have since bought a home through you, and how do you know that number?',
  'Give me a count and tell me where it lives. Is it a field in your CRM or is it in your head?',
  'A conversion rate and a tracking system. This is the single most valuable thirty seconds in the episode if he has the number.',
  'Individual agents, team leaders',
  note='This is your one use of "I have interviewed hundreds" for the episode. It is also the most dodgeable question in the prep. He will want to give you one great story about a renter who bought a penthouse. Do not accept the story in place of the number. If he does not have the number, that is also an honest answer and worth saying out loud.',
  short='How many of your renters have bought through you, and how do you track it?')

q(7,
  'You told us the worst advice you ever got was stay in your lane. Ranking nationally in rentals while running a luxury team is pretty much the definition of not staying in your lane. Where did that advice actually cost you, and where was the person who gave it to you right?',
  'What is one thing you tried that you should not have?',
  'The tension. The second half of the question is what keeps this from being a victory lap, and it is where he becomes credible.',
  'All segments',
  note='This is the callback you set up in Rapid Fire. Do not skip the second half of the question.',
  short='Worst advice was stay in your lane. Where did it cost you, and where was it right?')

q(8,
  'Say a solo agent wants to run rentals as a pipeline instead of a chore. What is the minimum system? What do they put in the database, how often do they follow up, and what month do they call?',
  'What is the actual field, and what does the call sound like?',
  'A copyable system with a cadence attached. This is the block payoff.',
  'Individual agents, new agents',
  short='Minimum system to turn rentals into a pipeline. What field, what cadence?')

bridge('That is the case for chasing the small deal, so now I want to ask you about the other side of it, which is the deal you talk somebody out of.')

h2('BLOCK 3: Telling a Buyer Not to Buy (22:00 to 33:00)')
p('Audience note: individual agents in every market. The Florida specifics are the setup, the advisory skill is the takeaway.', italic=True)
p('Arc: what changed, the client he told to rent, the objection, the script.', italic=True, space=8)

q(9,
  'Since the milestone inspections and the reserve studies came due, buyers down there are inheriting special assessments that run anywhere from ten thousand to over a hundred thousand dollars a unit. What are you doing on a condo now that you were not doing three years ago?',
  'Name the documents. Who pulls them and at what point in the process?',
  'A due diligence checklist an agent in any state can copy for HOA and condo deals.',
  'Individual agents, new agents',
  note='The assessment range is verified from Florida condo law coverage, not from him. State it as what is happening in the market, not as his number.',
  short='What do you check on a condo now that you did not check three years ago?')

q(10,
  'You have said you wanted to advise, not just transact. So tell me about the last client you told to rent instead of buy. What were the numbers that made that the right call?',
  'What was the price to rent math, and what was their time horizon?',
  'A real threshold. The moment where he named a number that cost him a commission.',
  'Individual agents, team leaders',
  permission='Keep it anonymous if you need to, but I want a real one, not a hypothetical.',
  short='Last client you told to rent instead of buy. What were the numbers?')

q(11,
  'There is an agent listening right now, three years in, one kid at home, and she hears a guy with a team of nine in Miami say tell your client to rent, and she thinks that is easy advice when you have that behind you. She cannot afford to talk somebody out of a deal this month. What do you say to her?',
  'What does she actually do differently on Monday?',
  'Whether his advice survives contact with an agent who is not him. If he cannot answer this, the whole block was luxury talk.',
  'Individual agents, new agents',
  short='An agent three years in cannot afford to talk clients out of deals. Answer her.')

q(12,
  'Give me the rent versus buy conversation as a script. What are the three questions an agent asks a buyer to figure out which side of the line that buyer is on?',
  'Say them the way you would say them to a client, not the way you would write them down.',
  'Three portable questions. This is the most clippable thirty seconds in the episode.',
  'All segments',
  short='Three questions an agent asks to figure out rent or buy. Go.')

bridge('That is how you handle the client, and I want to finish on how you handle the people who work for you, because you went from doing this alone to running nine.')

h2('BLOCK 4: Scaling Without Breaking It (33:00 to 42:00)')
p('Audience note: team leaders and broker owners, with the last question aimed back at solo agents.', italic=True)
p('Arc: the readiness signal, week one, the personal cost, the number before the hire.', italic=True, space=8)

q(13,
  'You have said something like start small and only scale when the infrastructure is ready to support it. You went from solo to a team of nine. What was the specific thing that told you the infrastructure was ready, and what broke first when it was not?',
  'What broke? Name it.',
  'A readiness signal a team leader can check against their own business, plus the failure that taught it.',
  'Team leaders, broker owners',
  note='This is a paraphrase of an Inman quote. Say "you have said something like," do not read it as a quote.',
  short='What told you the infrastructure was ready to hire? What broke first?')

q(14,
  'When somebody new joins The Wands Team, what do you actually hand them in week one? Leads, a market, a script, a rental desk. Be specific about what they get and what they have to earn.',
  'What does a new person on your team have to produce before they get anything else?',
  'The onboarding package. This is the recruiting question dressed as an operations question, which is exactly how our broker owners want it.',
  'Team leaders, broker owners, new agents',
  short='What does a new agent on your team get in week one?')

q(15,
  'You have said you wish you had understood earlier how much time this career demands. You were first in the office and last to leave for years. What did that actually cost you?',
  'Would you run those years the same way again?',
  'The honest beat. This is the moment that earns the follow.',
  'All segments',
  permission='Tell me if this is too personal, but.',
  short='What did first in, last out actually cost you?')

q(16,
  'Last one on the team side. Somebody listening has four agents and is about to hire a fifth. What is the one number they should look at before they do it?',
  'Where do they find that number?',
  'One metric, named out loud. Do not let him give three.',
  'Team leaders, broker owners',
  short='Four agents, about to hire a fifth. What number do they check first?')

h2('The Close (42:00 to 45:00)')

h3('Homework assignment (read verbatim)')
p('"Here is what I want you to do before the next episode. Open your database and pull every rental and every small deal you have closed in the last three years. Pick five of those people and call them this week. Ask one question: are you still renting. That is it. Not next month. This week."', bold=True)
p('Under thirty minutes. Costs nothing. Directly on the core topic, and the listener knows exactly what done looks like.', italic=True)

h3('Guest close')
bullet('', '"Where can people find you, follow you, or work with you?"')
bullet('', 'He will point to movetofl.com and Instagram. The team handles New York to South Florida referrals, so if a listener has a client heading south, that is the natural handoff. Let him say it, do not say it for him.')

h3('Optional levity beat, only if you are running short')
p('He sent a story about a buyer who asked to sleep over in a house before purchasing it. It is a good closer if you have ninety seconds. Do not tell it for him and do not force it if you are at time.', italic=True)

h3('If you are running long, cut these first')
table(['Cut order', 'Question', 'Why it is expendable'],
      [['1', 'Q14, week one onboarding', 'Team leaders get most of this from Q13 and Q16. It is the least portable question in the doc.'],
       ['2', 'Q9, condo due diligence', 'The most Florida specific question in the episode. Q10 carries the same block without it.'],
       ['3', 'Q3, the billionaire portal lead', 'Painful to cut because it is a great story, but Q1 and Q2 already establish the same principle.'],
       ['4', 'Q5, why rentals at all', 'Q6 asks the harder version of the same thing. If time is tight, skip to the number.']],
      widths=[0.8, 2.4, 3.8])

h3('Never cut')
bullet('Q6. ', 'The renter to buyer conversion number. This is the episode.')
bullet('Q1. ', 'The first scrap. It is the opening image the whole hour hangs on.')
bullet('Q10. ', 'The client he told to rent. The credibility question.')
bullet('Q12. ', 'The three question script. The clip.')
bullet('Q15. ', 'What it cost him. The human beat.')

doc.add_page_break()

# =====================================================================
# SECTION 4: RESEARCH BRIEF
# =====================================================================
h1('4. Research Brief')
p('Reference material. Read this the morning of, not during the interview.', italic=True, space=8)

h2('4A. Background')
p('Wands grew up in Port Washington, New York, in a family that had been in real estate for three generations before him, developers and contractors on both his great grandparents and grandparents sides. He moved to Miami in 2012 for the University of Miami and stayed. His first job out of school was selling at a water sports company, where he became the top salesperson quickly. He joined Douglas Elliman in 2016, won South Florida Rookie of the Year that first year, and by 2018 was a company Top Producer. He founded The Wands Team in 2022 and it now ranks among Elliman teams nationally.')

h2('4B. Career Timeline (verified entries only)')
table(['Year', 'Role / Company', 'Notable'],
      [['2012', 'Moved to Miami, University of Miami', 'Left Port Washington, New York'],
       ['After graduating', 'Sales, water sports company', 'Became top salesperson quickly'],
       ['2016', 'Joined Douglas Elliman', 'Douglas Elliman South Florida Rookie of the Year'],
       ['2018', 'Douglas Elliman Top Producer', 'Voyage MIA reports over $35 million sold that year'],
       ['2022', 'Founded The Wands Team', 'Built to serve an expanding client base'],
       ['May 2025', 'Inman Teams Spotlight', 'Team of 9. No. 11 medium team nationwide by GCI. Pinnacle Award'],
       ['April 2026', 'citybiz Q&A published', 'Market outlook interview']],
      widths=[1.1, 2.5, 3.4])

h2('4C. What Makes Him Interesting for This Audience')
bullet('The rental ranking: ', 'A luxury team inside Douglas Elliman ranks No. 4 in the company nationally by rental transactions. That contradiction is the episode, and it is a strategy any agent in any market can copy tomorrow.')
bullet('The average sale price: ', 'RealTrends Verified puts his average sale at $1.87 million. That is a real number a listener can relate to, not a $55 million headline, and it means his business is built on volume and repeat, not on one trophy deal a year.')
bullet('He asked for the small deal topic himself: ', 'Not overlooking small deals is on his own list of things he wants to discuss. He is bringing the material, which almost never happens.')
bullet('Fourth generation, but he started with nothing handed to him: ', 'The family legacy and the scraps story sit right next to each other, which is a tension worth naming.')
bullet('No podcast history: ', 'Nothing found across any show. The answers will not be rehearsed, which is rare for a guest at this production level.')

h2('4D. Key Data Points')
table(['Stat', 'Source', 'Confidence'],
      [['Douglas Elliman South Florida Rookie of the Year, 2016', 'Elliman bio, Voyage MIA, citybiz', 'High'],
       ['Founded The Wands Team in 2022', 'movetofl.com, Inman', 'High'],
       ['Team of 9', 'Inman Teams Spotlight, May 2025', 'High'],
       ['No. 4 team by rental transactions at Elliman', 'Inman, movetofl.com', 'High'],
       ['No. 7 by GCI, No. 8 by volume, No. 10 by transactions', 'Inman, movetofl.com', 'High'],
       ['No. 11 medium team nationwide by GCI, Pinnacle Award', 'Inman Teams Spotlight, May 2025', 'High'],
       ['RealTrends Verified: $60.28M volume, 32.3 sides, $1.87M average', 'RealTrends team profile', 'High'],
       ['48.6 sides and $100,282,407 volume', 'Inman Teams Spotlight, May 2025', 'Medium, conflicts with RealTrends, different period'],
       ['Sales executive on Five Park Miami Beach', 'LinkedIn, Elliman project coverage', 'Medium'],
       ['Over $35 million sold in 2018', 'Voyage MIA', 'Medium'],
       ['$1 billion in career sales', 'His own bio and website, which says sales and pendings', 'Unverified, guest supplied'],
       ['35 properties sold in his first year', 'Guest intake only', 'Unverified, guest supplied'],
       ['Record setting $55 million sale at The Perigon', 'Guest intake only', 'Unverified, guest supplied'],
       ['Miami condo median days to contract 85 in June 2026, up from 68', 'Market reporting', 'Medium, market context not his number'],
       ['Florida condo special assessments running $10K to $100K+ per unit', 'Florida condo law and market coverage', 'Medium, market context not his number']],
      widths=[3.0, 2.6, 1.4])

h2('4E. Previous Media Appearances')
table(['Outlet', 'Date', 'Topics covered'],
      [['Inman, Teams Spotlight', 'May 15, 2025', 'Team structure, rankings, advice to new team leaders, what he would tell his younger self'],
       ['citybiz, Q&A', 'April 21, 2026', 'South Florida luxury market normalization, buyer selectivity, what he is watching'],
       ['Voyage MIA', 'Undated', 'Full origin story, family background, early struggle, philosophy'],
       ['ARTRPRNR Magazine', 'Undated', 'Profile, could not access the full text'],
       ['Keeping It Real', 'None found', 'First appearance']],
      widths=[1.6, 1.3, 4.1])
p('No podcast appearances found on any platform. Treat every story as first telling.', italic=True)

h2('4F. Their Own Words')
p('Read Verbatim lines as quotes. Paraphrase everything marked Reported out loud, using "you have said something like."', italic=True, space=6)
table(['Quote', 'Where and when', 'Confidence', 'How you use it'],
      [['"Always bet on yourself and trust your gut, and do not focus solely or primarily on the money, focus on helping people first."', 'His KIR intake form', 'Verbatim', 'Rapid Fire will surface it. Leave it there.'],
       ['"Stay in your lane."', 'His KIR intake form, worst advice', 'Verbatim', 'Read it straight back to him at Q7. It is the spine of Block 2.'],
       ['"People would hand me their scraps."', 'Voyage MIA', 'Reported', 'Opens Q1. Say "you have said people handed you their scraps," do not attribute exact wording.'],
       ['"Real estate is a lifestyle, not a 9 to 5 job, and it is definitely not for the faint of heart."', 'Inman, May 2025', 'Reported', 'Hold in reserve for Q15 if he hedges on what the years cost him.'],
       ['"Start small and focus on building a strong foundation. Scale only when the infrastructure is ready to support it."', 'Inman, May 2025', 'Reported', 'Paraphrased into Q13. Ask him what ready actually looked like.'],
       ['"I wish I had truly understood the time and dedication this career demands. Success does not happen overnight."', 'Inman, May 2025', 'Reported', 'The setup for Q15. Paraphrase it, then ask what it cost.'],
       ['"This business is not about quick wins, it is a long game. Prioritize building genuine relationships over chasing commissions."', 'Inman, May 2025', 'Reported', 'Reserve. Use it if he starts selling the audience instead of teaching them.']],
      widths=[2.6, 1.3, 0.9, 2.2])

h2('4G. Audience Relevance')
table(['Segment', 'What they get from this episode'],
      [['Individual agents', 'A pipeline built out of the deals they are currently resenting, with a database field and a call cadence attached.'],
       ['Team leaders', 'The readiness signal for the next hire, and what a small deal desk does for new agent ramp time.'],
       ['Broker owners', 'A recruiting angle, which is that new agents survive longer when the brokerage gives them a small deal channel instead of telling them to farm luxury.'],
       ['New agents', 'Permission to take the two thousand dollar rental, plus the exact thing to do with it.']],
      widths=[1.6, 5.4])

h2('4H. Landmines')
bullet('Do not state the $55 million Perigon sale as fact. ', 'It appears nowhere outside his intake. Verified public Perigon and Five Park numbers are much lower, an $18 million Perigon contract that topped Miami-Dade weekly signings in May 2025 and a $17.5 million Five Park penthouse. If you want it, ask: "You mentioned a record at The Perigon. Walk me through that deal."')
bullet('Do not state 35 properties in year one as fact. ', 'Voyage MIA reports over $35 million sold in 2018, which is a different claim about a different year, and it is easy to conflate the two. Ask him to confirm the year one number if you want it on tape.')
bullet('Do not pick a volume number. ', 'Inman reported 48.6 sides and $100.3 million. RealTrends Verified shows $60.3 million and 32.3 sides. Both are probably right for different periods. If you need it live, ask "what does a normal year look like for the team now."')
bullet('$1 billion is career sales and pendings on his own site. ', 'Say "north of a billion in career sales and pendings, is that right," and let him own it.')
bullet('There are other Wands agents at Douglas Elliman. ', 'A Robert Wands and a Dawn Wands both appear in Elliman listings. He has said he is fourth generation, so a family connection is likely, but nothing confirms it. Do not assert it. If you are curious, ask.')
bullet('Sponsor awareness. ', 'Real Geeks and Courted are the show sponsors. If he names a lead source in Q3 or Q8, do not steer it toward or away from either one.')

h2('Drift Guardrail')
p('Three of his five suggested topics are South Florida market commentary. If the conversation slides into a market update, use this line: "Before we get too deep into Miami, I want to pull this back to the agent in Cleveland listening, because the thing you did with the small deals works anywhere." Then go to the next numbered question.', italic=True)

h2('Green Room Brief (two minutes before you record)')
bullet('', 'Tell him the episode is about small deals and the rental engine, not a Miami market update. He suggested the small deal topic himself so he will be glad, but he has almost certainly prepared market commentary.')
bullet('', 'Warn him about Rapid Fire. Four fast questions, one to two sentences each, and you will not follow up. Tell him you will circle back to the worst advice answer later so he does not spend it.')
bullet('', 'Ask him before the tape rolls whether he has the renter to buyer conversion number handy. If he does not, he can look it up in ninety seconds and the whole episode gets better.')
bullet('', 'Confirm the three guest supplied numbers off air: the Perigon sale, 35 properties in year one, and the billion. Then you can say them on air with him instead of about him.')

doc.add_page_break()

# =====================================================================
# SECTION 5: LIVE STREAM
# =====================================================================
h1('5. Live Stream Title, Descriptions and Hashtags')

h2('5A. Live Stream Title')
rich([('Primary: ', True), ('Chris Wands: The Miami Luxury Team That Ranks #4 in Rentals, and Why That\'s the Point', True), (' (85 characters)', False, True)])
rich([('Backup: ', True), ('Chris Wands on the Miami Deals Every Luxury Agent Turns Down', True), (' (60 characters)', False, True)])
p('This is set before the interview happens, so it only uses what research already verified. It does not have to match the published episode title. Pick that after you hear the tape.', italic=True)

h2('5B. Platform Descriptions')

h3('Facebook Live')
p('Chris Wands runs one of Douglas Elliman\'s top teams in Miami, and his team ranks number four in the entire company by rental transactions. Most luxury agents will not touch a rental. He built a business on them. Today he walks through exactly how the small deals turn into the big ones. Drop your questions in the comments!')

h3('Instagram Live')
p('A luxury team that ranks #4 in rentals. Chris Wands explains why the deals other agents turn down are the whole business. #RealEstatePodcast #RealtorLife #MiamiRealEstate')

h3('TikTok Live')
p('He got handed everybody else\'s scraps. Now he runs one of Elliman\'s top teams. #realtor #realestate #miami #realestateagent #realestatetips')

h3('YouTube Live')
p('Chris Wands, founder of The Wands Team at Douglas Elliman in Miami Beach, joins Keeping It Real Podcast to break down how rentals and small transactions became the pipeline underneath a luxury business. We cover the renter to buyer conversion system, when to tell a client to rent instead of buy, and how he scaled from solo agent to a team of nine. Real estate agent training, team building, and lead conversion with host D.J. Paris.')

h3('LinkedIn Live')
p('The Wands Team at Douglas Elliman ranks number four in the company nationally by rental transactions, inside one of the most competitive luxury markets in the country. Founder Chris Wands joins us to explain why that ranking is a deliberate acquisition strategy and not an accident. We get into the conversion system, the advisory line on rent versus buy, and what he checks before adding the next agent.')

h2('5C. Hashtag Sets')
bullet('Universal: ', '#KeepingItReal #RealEstatePodcast #DJParis #RealtorLife #RealEstateAgent')
bullet('Episode specific: ', '#MiamiRealEstate #LuxuryRealEstate #DouglasElliman #RealEstateTeam #RentalsToSales #MiamiBeach #RealEstateLeads')
bullet('Guest tags: ', '@chris_wands and @thewandsteam on Instagram, Christopher M. Wands on LinkedIn, The Wands Team on YouTube, MOVETOFL0RIDA on Facebook')

doc.add_page_break()

# =====================================================================
# SECTION 6: CHAPTERS
# =====================================================================
h1('6. YouTube Chapter Markers')
p('Estimates. Adjust after recording. Each title is written to be independently searchable.', italic=True, space=6)
table(['Timestamp', 'Chapter Title'],
      [['0:00', 'The Luxury Team That Ranks #4 in Rentals'],
       ['2:00', 'Rapid Fire: Best and Worst Real Estate Advice'],
       ['4:00', 'They Handed Him Their Scraps: Chris Wands on His First Deals at Elliman'],
       ['7:00', 'What a $2,500 Rental Actually Costs an Agent in Hours'],
       ['9:30', 'The Realtor.com Lead That Turned Out to Be a Billionaire'],
       ['11:00', 'What to Do With the $2,000 Rental Nobody Wants'],
       ['12:30', 'Why a Luxury Team Chases Rental Transactions'],
       ['15:30', 'How Many Renters Actually Buy: The Conversion Number'],
       ['18:30', 'Stay In Your Lane Was the Worst Advice He Ever Got'],
       ['20:30', 'The Minimum System to Turn Rentals Into a Pipeline'],
       ['22:30', 'What Agents Check on Condos Now That Assessments Have Changed'],
       ['26:00', 'The Client He Told to Rent Instead of Buy'],
       ['29:00', 'Answering the Agent Who Cannot Afford to Lose a Deal'],
       ['31:30', 'Three Questions to Decide Rent Versus Buy'],
       ['33:30', 'Solo to Nine: How He Knew the Team Was Ready'],
       ['36:00', 'What a New Agent Gets in Week One on The Wands Team'],
       ['38:30', 'What First In and Last Out Actually Cost Him'],
       ['40:30', 'The One Number to Check Before You Hire Your Fifth Agent'],
       ['42:00', 'Homework and Where to Find Chris Wands']],
      widths=[1.2, 5.8])

doc.add_page_break()

# =====================================================================
# SECTION 7
# =====================================================================
h1('7. Stress Test, Council Review and EP Polish')

h2('7A. Stress Test (pass 2, part one)')
table(['#', 'What broke', 'Fix applied'],
      [['1', 'The draft opened Block 1 by asserting he sold 35 properties in his first year. That number exists only in his intake, and a separate source reports over $35 million sold in 2018, which is easy to conflate.', 'Pulled the number out of every question, the cold open, and all four titles. Moved it to Landmines as guest supplied and added a green room instruction to confirm it off air.'],
       ['2', 'The $55 million Perigon sale was in the first title option and the cold open. Nothing independent confirms it. Public Perigon and Five Park records top out around $17.5 to $18 million.', 'Removed it from packaging entirely. Rebuilt all three titles and the live stream title on the verified No. 4 rental ranking. Added the exact ask if D.J. wants it on tape.'],
       ['3', 'Two conflicting volume figures. Inman reported 48.6 sides and $100.3 million, RealTrends Verified shows $60.3 million and 32.3 sides.', 'Named both in 4D with the conflict marked, and wrote the safe on air phrasing into Landmines so D.J. asks rather than asserts.'],
       ['4', 'Rapid Fire question two spends the worst advice answer, and stay in your lane is the spine of Block 2.', 'Added the deflection line to Rapid Fire and the Quick Reference Card, and built the callback into Q7.'],
       ['5', 'The billionaire portal lead was going to come out as a throwaway in Rapid Fire question four, or worse, get told as a name drop with no lesson.', 'Made it Q3 with a producer note that pulls him back to the response protocol after the story lands.'],
       ['6', 'Q6 was dodgeable. He can answer any renter conversion question with one great anecdote about a renter who bought a penthouse.', 'Rewrote it to demand a count and a tracking location, and added the producer note telling D.J. not to accept the story instead.'],
       ['7', 'Q2 had the same problem. Walk me through a rental invites a service answer.', 'Rewrote it to ask for hours and for what goes in the file, and added the note.'],
       ['8', 'The draft had five blocks and twenty one questions for a 45 minute runtime.', 'Cut to four blocks and sixteen questions, dropped the standalone South Florida market block entirely, and folded the current moment material into Block 3 where it has a job.'],
       ['9', 'Drift risk. Three of his five suggested topics are Miami market commentary and he will arrive prepared to deliver them.', 'Wrote the guardrail line, put it in Section 4, and told him in the green room brief what the episode is actually about.'],
       ['10', 'Q11 read as an ambush in the draft, because it put the skepticism in D.J.\'s mouth.', 'Reframed it around a described listener, an agent three years in with a kid at home, so it lands as a shared problem instead of a challenge.'],
       ['11', 'Q15 touches what his career cost him personally and had no softening.', 'Added the permission clause as words D.J. says.'],
       ['12', 'The rent versus buy block was all Florida. A Chicago listener gets nothing from milestone inspections.', 'Kept Q9 as the setup but ended the block on Q12, a portable three question script, and marked Q9 as the second thing to cut.']],
      widths=[0.4, 3.3, 3.3])

h2('7B. Council Review (pass 2, part two)')

h3('Member notes')
table(['Member', 'What they would change'],
      [['Alex Hormozi', 'Q6 is the only question in here that pays the listener a number. Protect it and cut anything that competes with it for airtime.'],
       ['MrBeast', 'The sag is minute twelve to fifteen, right after the origin block. Open Block 2 on the contradiction, not on a warm up. Fixed.'],
       ['Brendan Kane', 'The first title was built on a $55 million number you cannot verify. The rental ranking is the better hook anyway, because it is a contradiction and not a brag. Ingredient is the stat.'],
       ['Gary Vaynerchuk', 'The LinkedIn description was the Facebook one with longer words. Rewrote it around the strategy, which is what that audience actually clicks.'],
       ['Donald Miller', 'The hero was drifting toward Chris. Q11 puts the listening agent back in the center. Keep it, it is the grunt test for the whole episode.'],
       ['Eric Simon', 'Every agent has been handed a garbage rental and felt insulted by it. Q4 is the question that gets this sent from one agent to another.'],
       ['Justin Welsh', 'Do not break the standardized Rapid Fire to protect the stay in your lane answer. Use the deflection line instead. Fixed that way.'],
       ['Jon Youshaei', 'Chapter markers were labeled by block name. Rewrote all nineteen as standalone searchable headlines.'],
       ['Chip and Dan Heath', 'Ruling on the curiosity claim: the title opens the gap correctly, a luxury team ranking in rentals is a Huh before it is an Aha. It fails only if a listener does not know Elliman is a luxury brokerage, so the cold open says so explicitly.'],
       ['Jonah Berger', 'Ruling on the share claim: the emotion is high arousal, it is indignation on behalf of the agent who was handed the scraps. Social currency is real, because retelling this makes the sharer look strategic. The one sentence version survives, which is that the small deals are the pipeline.']],
      widths=[1.5, 5.5])

h3('Title')
table(['#', 'Title', 'Ingredient', 'Curiosity mechanism'],
      [['1', 'The #4 Rental Team at Douglas Elliman Is a Luxury Team. Chris Wands on Why Small Deals Win', 'Stat', 'Contradiction. Luxury and rentals do not belong in the same sentence, so the reader needs the resolution.'],
       ['2', 'They Handed Him the Scraps. Now He Runs One of Elliman\'s Top Teams. (Chris Wands)', 'Insight', 'Status reversal. The gap is the how, and the title refuses to close it.'],
       ['3', 'The Miami Luxury Broker Who Tells Clients to Rent (Chris Wands, The Wands Team)', 'Concept format', 'Self interest violation. Why would he talk himself out of a commission.']],
      widths=[0.4, 3.0, 1.0, 2.6])
p('Recommended: #1. The episode goal here is clips and reach, and it is the only option carrying a verified number that is also the contradiction.', italic=True)

h3('Cold open (sharpened)')
p('"Inside one of the largest luxury brokerages in the country, there is a Miami team that ranks number four nationally in rental transactions, and its founder says that ranking is the reason he has a luxury business at all. We are going to talk about that today. Stay tuned."', bold=True)

h3('The Clip Engine')
table(['Q#', 'Question', 'Berger emotion', 'Heath gap'],
      [['Q6', 'How many of your renters have bought through you, and how do you track it?', 'Excitement and awe if the number is real', 'Opens on a claim everyone believes, which is that rentals do not pay, then closes it with a count.'],
       ['Q12', 'Three questions an agent asks to figure out rent or buy.', 'Social currency, the sharer looks like an advisor', 'Promises a script and delivers it inside thirty seconds.'],
       ['Q4', 'An agent just got handed a $2,000 rental. What do they do with it this week?', 'Indignation flipped into usefulness', 'Names a feeling the listener has had, then resolves it with an instruction.']],
      widths=[0.5, 3.0, 1.7, 1.8])

h3('Live description scrub')
table(['Platform', 'Verdict and fix'],
      [['Facebook', 'Keep. Conversational, leads on the contradiction, ends with the comment prompt.'],
       ['Instagram', 'Fixed. The draft ran four sentences. Cut to two plus tags.'],
       ['TikTok', 'Fixed. Rewrote as the scraps line, which is the only version that works as a spoken hook.'],
       ['YouTube', 'Keep. Carries his name, the company, the city, and the searchable terms.'],
       ['LinkedIn', 'Fixed. Rewritten around the strategy rather than the personality.']],
      widths=[1.2, 5.8])

h3('Arc fix')
p('The most likely sag was minute twelve, where the origin story ends and the episode has to restart on a new idea. The fix was the bridge line out of Block 1, which states the through line out loud, plus opening Block 2 on the contradiction question rather than on setup. The second risk was Block 3 blurring into a Florida market segment, which the Q12 script question and the drift guardrail both defend against.')

h3('Why it should work')
bullet('Curiosity mechanism, Heath: ', 'A luxury team ranking nationally in rentals is a Huh that only the episode can turn into an Aha.')
bullet('Share driver, Berger: ', 'Every agent has been handed a deal beneath them. The takeaway survives one sentence of retelling, which is that the small deals are the pipeline.')
bullet('Retention move, MrBeast: ', 'The conversion number at Q6 sits at minute fifteen, right where listeners drop, and the whole first block is built to make them want it.')

h3('The dissent')
p('Hormozi is still objecting that the episode is worth nothing if Q6 comes back without a number, and he is right. The experiment is the green room ask, which is telling the guest before the tape rolls exactly which number you are going to want. If it works here, put a one line number request in every guest confirmation email going forward.')

h2('7C. EP Polish (pass 3)')
bullet('', 'Stripped the $55 million Perigon sale, the 35 properties, and the $1 billion out of all three titles, the cold open, the live stream title, and every platform description, then rebuilt the packaging on the verified No. 4 rental ranking.')
bullet('', 'Cut the draft from five blocks and twenty one questions to four blocks and sixteen, and deleted the standalone South Florida market block, which was the one block that could have been its own episode.')
bullet('', 'Rewrote Q2 and Q6 to demand hours, counts, and a tracking location instead of numbers he could substitute a story for, and added producer notes telling D.J. not to accept the substitute.')
bullet('', 'Reframed Q11 around a described listener rather than around D.J.\'s skepticism, which turns a possible ambush into a shared problem.')
bullet('', 'Moved the billionaire portal lead from a likely Rapid Fire throwaway into Q3, where it functions as proof for the core topic instead of as a name drop.')
bullet('', 'Added the Rapid Fire deflection line in three places, the Quick Reference Card, the Rapid Fire section, and the green room brief, because the worst advice answer is load bearing for Block 2 and the format will spend it in minute one.')
bullet('', 'Wrote permission clauses on exactly two questions, Q10 and Q15, and removed them everywhere else so the device does not become a tic.')
bullet('', 'Added the green room brief, because four separate facts in this episode are guest supplied and the only real fix is to get him to own them before the tape rolls.')
bullet('', 'Rewrote all nineteen chapter markers as standalone searchable headlines rather than block labels.')
bullet('', 'Rewrote the LinkedIn and TikTok descriptions so they are not the Facebook one at a different length.')
bullet('', 'Moved Q9 to second in the cut order and ended Block 3 on the portable three question script, so the block survives if the Florida specifics have to go.')
bullet('', 'Added the drift guardrail line, since three of his five suggested topics are market commentary and he will arrive prepared to deliver them.')
bullet('', 'Added the other Wands agents at Elliman to Landmines, because a fourth generation claim plus two same name agents in the company is exactly the kind of thing that gets asserted wrong on air.')
bullet('', 'Checked every short version against the twenty word ceiling. Q1, Q6, and Q11 had to be rewritten to survive the compression.')
bullet('', 'Swept the document for em dashes, curly quotes, and AI-speak. Zero.')

doc.save("/Users/djparis/GitHub Projects/keeping-it-real-content-system/guest-prep/Chris_Wands_Interview_Prep.docx")
print('saved')
