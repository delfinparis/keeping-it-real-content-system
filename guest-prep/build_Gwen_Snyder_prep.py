#!/usr/bin/env python
# Builds the KIRP interview prep .docx for Gwen (Gwenivere) Snyder
# v2 - after stress test, second pass, EP polish, council re-run
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)

for name, size, color in (('Heading 1', 18, 0x1F3864), ('Heading 2', 14, 0x2E5496), ('Heading 3', 11.5, 0x2E5496)):
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor((color >> 16) & 255, (color >> 8) & 255, color & 255)
    st.font.bold = True

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.8)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def p(text='', bold=False, italic=False, size=None, space=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space)
    r = par.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    return par


def rich(segments, style=None, space=6):
    par = doc.add_paragraph(style=style)
    par.paragraph_format.space_after = Pt(space)
    for seg in segments:
        text = seg[0]
        b = seg[1] if len(seg) > 1 else False
        i = seg[2] if len(seg) > 2 else False
        r = par.add_run(text)
        r.bold, r.italic = b, i
    return par


def bullet(text_bold, text_rest=''):
    par = doc.add_paragraph(style='List Bullet')
    par.paragraph_format.space_after = Pt(2)
    if text_bold:
        par.add_run(text_bold).bold = True
    if text_rest:
        par.add_run(text_rest)
    return par


def quoted(label, text):
    par = doc.add_paragraph(style='Intense Quote')
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.space_before = Pt(3)
    r = par.add_run(label + ' ')
    r.bold = True
    r.italic = True
    r2 = par.add_run(text)
    r2.italic = True
    return par


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            par = cells[i].paragraphs[0]
            run = par.add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def q(num, question, if_vague, reveals, serves):
    rich([(f'{num}. ', True), (question, True)], space=2)
    quoted('If vague, ask:', if_vague)
    quoted('Ideal answer reveals:', reveals)
    quoted('Serves:', serves)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


# =====================================================================
# TITLE
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('KEEPING IT REAL PODCAST')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Interview Prep: Gwen Snyder')
r.bold = True
r.font.size = Pt(24)
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('LIV Sotheby\'s International Realty, Greenwood Village CO  |  Founder, Referable  |  Prepared August 4, 2026  |  v2')
r.italic = True
r.font.size = Pt(10)

# =====================================================================
# PAGE 0 — WHAT CHANGED
# =====================================================================
h1('WHAT CHANGED IN v2')
p('Four passes ran on the first draft: stress test, rebuild, EP polish, council. If you already read v1, these are the only things you need to re-learn.', italic=True, size=9.5)
bullet('Her origin story has three incompatible published versions. ', 'v1 built the opening question on the Beverly Hills version. That question is gone. See Landmine 1.')
bullet('The dungeon story has an identifiability problem. ', 'v1 only flagged explicit tagging. There are now hard rules. See Landmine 4.')
bullet('Rapid fire question 4 will burn the dungeon clip if you do not pre-brief her. ', 'Fix is a green-room sentence. See the Rapid Fire producer note.')
bullet('Act 1 went from nine questions to seven. ', 'The perspective flip became a follow-up instead of its own slot, and the personal question moved to the end of Act 2.')
bullet('The title recommendation changed. ', 'The council resolved its own dissent. Different titles now ship to YouTube and to the podcast feed.')
bullet('The cold open no longer states a number of years. ', 'v1 contradicted its own landmine.')

doc.add_page_break()

# =====================================================================
# PAGE 1 — QUICK REFERENCE CARD
# =====================================================================
h1('PAGE 1: QUICK REFERENCE CARD')
p('Glance at this during the interview. Everything else is morning-of reading.', italic=True, size=9.5)

bullet('Name: ', 'Gwenivere Snyder, goes by Gwen. Pronouns not publicly stated; she/her used throughout her own press. Confirm "Gwen or Gwenivere" and the pronunciation in the green room before you record the intro.')
bullet('Title: ', 'Global Real Estate Advisor, LIV Sotheby\'s International Realty. Office: 8000 E Belleview Ave, Suite 200, Greenwood Village, CO 80111.')
bullet('Also runs: ', 'Referable, her coaching program for mid-career agents. Launched 2022. Private 1:1, 8-12 agent group cohorts, and half or full day workshops for brokerages and conferences.')
bullet('Licensed: ', 'Colorado, 2004.')
bullet('Markets: ', 'Denver metro luxury (Cherry Hills Village, Greenwood Village, Castle Pines, Highlands Ranch, Lone Tree, Parker, Broomfield) plus the mountain markets she added in 2018 (Evergreen, Conifer).')
bullet('Key stats: ', '$500M+ career volume, crossed in 2025. 97% list-to-sale ratio. #1 on her LIV Sotheby\'s team. Top 1% nationwide per RealTrends. 55+ professional athletes, coaches and managers as clients.')
bullet('Awards: ', '5280 Magazine Top Real Estate Producer 2026 (Double Black Diamond). Denver Real Producers 2025. RealTrends Verified and The Thousand, 2025. Mile High Leaders Top 100 Broker, Denver Metro. Denver Metro REALTORS Excellence Awards. REALM member.')
bullet('Stages: ', 'Denver Build 26 (spring 2026), Colorado Association of Realtors Annual Conference (fall 2025), Real Producers Summit (spring 2025).')
bullet('Philanthropy: ', 'Donates 10% of every closed deal to causes for women and children. Founding President and Executive Board member of the Women\'s Guild for A Precious Child.')
bullet('Find her: ', 'gwensnyderrealestate.com | IG @gweniveresnyder (37K) | Facebook @GwenivereLuxuryRE | X @GwenLuxuryRE | LinkedIn /in/gweniveresnyder | 303.718.1085')
bullet('KIR connection: ', 'FIRST TIME GUEST. No prior KIR appearance found. See the sponsor note.')
bullet('Guest type: ', 'A with a D overlay. Producing agent first, coach second. Interview her as a practitioner. Product cap: no more than 20% of questions about Referable, and the episode has to pay off for an agent who never buys a minute of her coaching.')
bullet('Episode length: ', 'NOT CONFIRMED. Everything here is built to 40 minutes. If it is 30, cut Act 1 Q3 and Act 2 Q2. If it is 60, Act 1 Q2 and Q5 are the two that reward more room.')

rich([('THE CORE TOPIC: ', True),
      ('How a 100% referral business actually gets built and maintained at half a billion dollars in volume, told as a mechanism an agent can copy this month, not as a personality trait you either have or you don\'t.',)])

rich([('THE ONE THING THAT MAKES THIS AN EPISODE AND NOT A LECTURE: ', True),
      ('"Referable" sounds like a compliment, not a plan. Your whole job is converting every warm word she uses into a step, a number, a script, or a calendar entry. If she says "I just take great care of people," you do not have an answer yet. Ask again, differently.', False, False)])

rich([('OVERASKED, DO NOT ASK: ', True)])
bullet('', '"Tell us about the car accident at 17." Her standard resilience answer, told in Bold Journey (Dec 2023) and CanvasRebel (Aug 2023). If she goes there herself, let her, then move. Do not open the door.')
bullet('', '"How did you get into real estate?" See Landmine 1. This question is a trap on this guest for reasons that have nothing to do with it being overasked.')
bullet('', '"What is it like working with professional athletes?" Every local profile runs this. It is a stat, not a story. If you want it, ask what athletes taught her about privacy.')
bullet('', '"What advice do you have for new agents?" Her published answer is already three words: grit, follow-up, consistency. Act 2 covers this properly.')

rich([('THE "I\'VE INTERVIEWED HUNDREDS" MOMENT (use once, Act 1 Q5): ', True),
      ('"I\'ve interviewed hundreds of agents on this show, and the most common thing I hear about referrals is that they take too long, so they buy leads to cover the gap. You\'re saying the leads are the slow part. Make that case."', False, False)])

doc.add_page_break()

h2('LANDMINES (read all five before you record)')

rich([('1. HER ORIGIN STORY HAS THREE INCOMPATIBLE PUBLISHED VERSIONS. ', True),
      ('Her 2021 ShoutOut Colorado profile says she modeled in Europe, worked in private aviation and entertainment in LA, sold high-end residential in Beverly Hills, then relocated to Colorado. Her Aug 2023 CanvasRebel interview says she started about twenty years ago at Sotheby\'s in Beverly Hills while pregnant with twins, pushed into it by a mentor named Mimi Starrett. Her current website says she was licensed in Colorado in 2004 and "immediately focused on the luxury segment of the Denver metro market," under a headline that reads "Twenty years, one market." The Beverly Hills chapter is not on her site anymore.', False, False)])
rich([('   WHAT YOU DO: ', True),
      ('Nothing. Do not referee her resume on air, do not read the Beverly Hills story in your intro, and do not pin a year to a place. She may have repositioned deliberately, and it is not the story. Act 1 Q1 has been rewritten to get everything you need from the beginning of her career without asserting where or when it started. If she tells the Beverly Hills version herself, take it at face value and keep going.', False, False)])

rich([('2. THE YEAR COUNT. ', True),
      ('Her intake bio says 23 years. Her own site says "20+ Years in Market" and "Twenty years, one market," off a 2004 license, which is 22. Three numbers, all hers. Say "licensed in 2004" or "more than two decades" and you cannot be wrong. Do not say a specific number of years, including in the cold open.', False, False)])

rich([('3. THE VOLUME NUMBER. ', True),
      ('Published career volume across sources: $150M (2021), $200M (Aug 2023), $325M (Dec 2023), $400M+ (LinkedIn), $500M+ (her site and Instagram now, milestone dated 2025). Read $500M+ and attribute it lightly ("she\'s crossed half a billion"). Never stack two figures in one sentence and never do the subtraction out loud. It is her published number and that is where your interest in it ends.', False, False)])

rich([('4. THE DUNGEON STORY IS AN IDENTIFIABILITY PROBLEM, NOT JUST A TASTE PROBLEM. ', True),
      ('She is describing real, living, named-in-her-own-mind people, saying they own a "very well known fashion brand," and alleging what she took to be a commercial operation. On a distributed show, if the brand becomes guessable, that is a real person being accused of something. HARD RULES: she does not name the brand, you do not guess, you do not narrow it (no city beyond "California," no decade, no "was it the one that..."), and you do not ask what happened to them. If she starts narrowing, your line is "don\'t tell me who, I don\'t want to know." That line is also funny, which is why it works. Keep the $7M Dr. Seuss house as the swap if she is uncomfortable, and tag the episode explicit if she goes past the written version.', False, False)])

rich([('5. SPONSOR EXPOSURE. ', True),
      ('Her entire thesis is no cold calls and no bought leads. Real Geeks sells lead generation. Courted sells agent data that feeds prospecting. Ask her the question and let her answer it fully, but do not co-sign it, do not add your own shot at the category, and do not let a clip go out that reads as the show trashing a sponsor\'s product category. The frame that keeps everyone whole: leads are a faucet, referrals are a well, most agents are running the faucet because they never dug the well. There is a scripted bridge on Act 1 Q5.', False, False)])

rich([('TONE READ, DO THIS IN THE FIRST TWO MINUTES: ', True),
      ('Her intake reads funny, unfiltered, self-deprecating, eats-pavement-in-front-of-clients. Her website reads discreet luxury concierge, "preparation and discretion," not a joke on the whole site. Those are two different guests and you will not know which one showed up until the rapid fire. If you get the intake version, lean into the stories and let her run. If you get the website version, drop the comedy setups entirely, do not fish for the funny, and go harder on mechanism and numbers, because the systems episode works with either guest and the comedy episode only works with one.', False, False)])

doc.add_page_break()

# =====================================================================
# PAGE 2 — EPISODE FRAMEWORK
# =====================================================================
h1('PAGE 2: EPISODE FRAMEWORK')

h2('2A. Title Options')
p('The council resolved these into a split ship. See Section 7.', italic=True, size=9.5)
table(['#', 'Title', 'Ships To', 'Why It Works'],
      [['1', 'She Has Never Made a Cold Call. Gwen Snyder on Why Chasing Leads Is the Slow Way to Build a Business.',
        'Podcast feed, social cuts',
        'Leads with the refusal, which every agent identifies with instantly and which costs nothing to see yourself in. "Slow way" inverts what they were trained to believe, so they click to argue.'],
       ['2', '$500 Million Sold, Zero Cold Calls: Gwen Snyder on the Referral Engine That Replaced Prospecting',
        'YouTube',
        'The number does search and authority work where authority is the currency. Two devices in one line, and the back half is A/B-able without touching the front.'],
       ['3', 'The 97% List-to-Sale Agent: Gwen Snyder on Becoming the Agent People Cannot Stop Naming',
        'Hold',
        'Her most under-quoted stat, but 97% needs a beat of explaining and a title that needs explaining has already lost. Keep for a clip caption.']],
      widths=[0.3, 2.9, 1.1, 2.6])

h2('2B. Cold Open Hook')
p('D.J. reads this before the ads:', italic=True)
rich([('"My guest today has sold more than half a billion dollars of Colorado real estate and has never made a single cold call. Not one. Not ever, and not by accident. We\'re going to talk about that today. Stay tuned."', False, True)])
p('Note the year count is gone (Landmine 2) and "not by accident" replaced it. It does the same work, because it turns a refusal into a decision, and a decision has a reason, and the reason is the episode.', size=9.5, italic=True)

h2('2C. Episode Arc')
rich([('Core Topic: ', True), ('The referral engine as an actual system with steps, timing and a failure mode, built by an agent who has run nothing else for two decades.', False, False)])
rich([('Why this topic: ', True), ('Every prior interview she has given is the resilience story or the origin story. Nobody has made her show her work. It is also the highest-value topic for the largest audience segment, because "referrals" is the thing every agent says they want and nobody can describe a process for.', False, False)])
rich([('What we are NOT doing: ', True), ('The car crash, the origin story, athlete name-drops, and a tour of the Referable curriculum. Her career arc is Section 4 reading, not an agenda. One sentence and back.', False, False)])

doc.add_page_break()

# =====================================================================
# SECTION 3 — INTERVIEW QUESTIONS
# =====================================================================
h1('SECTION 3: INTERVIEW QUESTIONS')

h2('3-Act Structure (40 minute target)')
table(['Act', 'Time', 'What Happens'],
      [['RAPID FIRE', '0:00-2:00', 'The 4 standardized questions. She pre-answered two in her intake. No follow-ups. Pre-brief required, see below.'],
       ['ACT 1: DEEP DIVE', '2:00-25:00', 'Seven questions, one topic: how the referral engine actually works. Roughly 3 minutes each, which is a real conversation.'],
       ['ACT 2: THE PLAYBOOK', '25:00-36:00', 'Four tactical questions, then one human landing.'],
       ['THE CLOSE', '36:00-40:00', 'Homework, where to find her, Referable.']],
      widths=[1.5, 1.3, 4.1])

h3('RAPID FIRE (standardized, read as written, no follow-ups)')
p('1. Best real estate advice you\'ve ever received?')
p('2. Worst real estate advice you\'ve ever received?')
p('3. One tool or app you can\'t run your business without?')
p('4. What would surprise people most about your day-to-day?')

rich([('PRE-BRIEF, SAY THIS IN THE GREEN ROOM (this is the fix for a real problem): ', True),
      ('"One thing before we start. In the rapid fire I\'m going to ask what would surprise people about your day-to-day. Don\'t give me the door story there. I want to set that one up properly later so it lands." Without this, question 4 pulls the dungeon story straight into a two-minute segment where the format forbids you from following up, and the best clip in the episode dies in a zone where you are supposed to say "love it" and move on.', False, False)])

rich([('PRODUCER NOTE: ', True),
      ('She pre-answered 1 and 2. Best: "If you\'re not early, you\'re late," framed as a mindset and not a clock. Worst: "Cold call your way to more business," which she has never done and never will. That worst-advice answer is your cold open and your Act 1 opener. Let it land, say "Love it," move. You are planting it on purpose to pick it up ninety seconds later with the full runway. Question 3 is the one place a coach can start listing a CRM stack. Past two sentences, "Love it" over the top and go.', False, False)])

doc.add_page_break()

h2('ACT 1: THE DEEP DIVE, The Referral Engine (23 min)')
p('Audience note: individual producing agents, the largest segment. Every one of them has been told to prospect and most of them hate it. She is the proof of the alternative, but only if she shows the mechanism instead of the philosophy.', italic=True)

q(1,
  'Ninety seconds ago you told me the worst advice you ever got was to cold call your way to a business, and that you\'ve never done it and never will. So take me back to the very beginning. Brand new license, no past clients, nobody owes you anything. If you\'re not calling strangers, where does the first deal actually come from?',
  'Name one actual person and how that one deal happened. Not the category, the person.',
  'That referral-only is something you bootstrap, not something you inherit. Every listener with an empty database needs to hear that year one had a mechanism.',
  'Individual agents, new agents.')

rich([('WHY IT IS WORDED THIS LOOSELY: ', True), ('See Landmine 1. It gets you the founding mechanism without asserting a year, a city, or a brokerage, all of which conflict across her own published sources. Let her choose which version of the beginning she tells.', False, False)])
doc.add_paragraph().paragraph_format.space_after = Pt(0)

q(2,
  '"Referable" sounds like a personality trait, and personality traits aren\'t transferable. So make it a process for me. A file closes on a Friday. What actually happens in the next ninety days that turns that one client into three more?',
  'Day one, day thirty, day ninety. Is it in a calendar, a CRM, or your head? Do you do it or does someone on your team?',
  'The actual cadence with timing and triggers. This is the most stealable minute in the episode. Do not accept a sentiment. Ask a third time if you have to.',
  'Individual agents, team leaders.')

q(3,
  'Your list-to-sale ratio is 97%. Most agents hear that and think great negotiator. I think it means you\'re telling sellers no. What\'s the exact sentence you use when a seller wants a number you won\'t take?',
  'Give me the words you say out loud, not the philosophy of pricing.',
  'That the referral engine runs on refusals as much as service. A walked listing is a reputation deposit. And listeners get a script they can use Tuesday.',
  'Individual agents, team leaders.')

rich([('IF SHE DOESN\'T TRACK IT: ', True), ('v1 asked "how many listings do you turn down in a year," which invites "I don\'t count." That is cut. If she still deflects, go to: "Okay, forget the number. When was the last time you walked away from one, and what was the conversation?" A story beats a statistic here anyway.', False, False)])
doc.add_paragraph().paragraph_format.space_after = Pt(0)

q(4,
  'Here\'s what worries me about a referral-only business. It doesn\'t crash, it just gets quiet. There\'s no bad month you can point at, the phone is only a little slower and then a little slower. Has that happened to you?',
  'What was the leading indicator you missed? What do you watch now so it can\'t sneak up on you again?',
  'The failure mode and the early warning system. This is the most credible thing she can say, and it is the moment the skeptics decide to trust her.',
  'Individual agents, team leaders.')

q(5,
  'I\'ve interviewed hundreds of agents on this show, and the most common thing I hear about referrals is that they take too long, so they buy leads to cover the gap. You\'re saying the leads are the slow part. Make that case.',
  'Put numbers on it. Conversion, cost, time to close, what a referral is worth over five years versus a purchased lead.',
  'The economic argument, not the moral one. If all you get is "leads feel gross," you do not have the answer yet.',
  'All segments. This is the clip.')

rich([('SPONSOR GUARDRAIL, SCRIPTED: ', True), ('Ask it, let her answer it, do not agree out loud, do not pile on. Your bridge out, word for word: "So for the agent who\'s already got a lead source and isn\'t going to turn it off, what\'s the referral layer they build underneath it?" That converts an either/or into a both/and, gets you a better answer, and protects the show. See Landmine 5.', False, False)])
doc.add_paragraph().paragraph_format.space_after = Pt(0)

q(6,
  'You\'ve walked buyers through thousands of homes. My producer tells me a very well known fashion family once opened a locked door in front of you and it was a full dungeon, and all you said was "great storage." I need you to tell that story.',
  'Nothing. Get out of the way and let her tell it. One reaction when she lands it, not five.',
  'The funniest sixty seconds of the episode and the only clip that travels outside real estate. It is also the pattern interrupt that re-hooks the middle of the show.',
  'All segments.')

rich([('THEN, AS A SEPARATE FOLLOW-UP (do not bundle these, the funny answer will eat the analytical one): ', True),
      ('"Okay, but here\'s the real question underneath that. Twenty years of seeing what\'s behind people\'s locked doors. What has that taught you about who\'s actually making the decision in a room?"', False, True)])
rich([('AND REMEMBER: ', True), ('Landmine 4. No brand, no city, no decade, no guessing. If she starts narrowing it, "don\'t tell me who, I don\'t want to know" is your exit and it gets a laugh.', False, False)])
doc.add_paragraph().paragraph_format.space_after = Pt(0)

q(7,
  'Something changed this year that I don\'t think agents have connected to your world yet. Clear Cooperation is basically dead, and Sotheby\'s own 2026 report says 81% of their agents hear privacy as the client\'s top concern. If the best inventory increasingly never hits the public market, your network isn\'t branding anymore. It\'s access. Am I overstating that?',
  'Give me a real one. A property that traded inside a network and never went public, and what got you in the room.',
  'That "referable" stopped being a soft skill and became an inventory question. This reframes the whole episode for anyone who thought the topic was feel-good.',
  'All segments. This is your industry-news beat.')

rich([('PERSPECTIVE FLIP, ASK IT HERE AS A FOLLOW-UP: ', True),
      ('"And if I\'m a broker-owner watching the best inventory move inside private networks, what does that change about who I hire?" That is your required flip for the broker-owner and team-leader segments, and it lands harder attached to this question than it did as its own slot in v1.', False, True)])
rich([('SOURCING NOTE: ', True), ('The Clear Cooperation collapse is well documented as of May 2026 and the 81% figure is from Sotheby\'s International Realty\'s own 2026 Luxury Outlook, which is safe ground with a Sotheby\'s agent. Do not go further and assert specific Zillow or Realtor.com listing policies on air. You have not read those policies and you do not need them to make the point.', False, False)])

doc.add_page_break()

# =====================================================================
# ACT 2
# =====================================================================
h2('ACT 2: THE PLAYBOOK (11 min)')
p('Audience note: for everyone, including the agent with six closings a year and no database. Translate everything she said into Monday morning.', italic=True)

q(1,
  'Someone listening does six deals a year, has maybe forty past clients in a spreadsheet, and has never asked any of them for anything. What\'s the minimum viable version of what you do? Not the whole system. The first thing, the one they start Monday.',
  'How many people, how often, and what do they actually say? If it\'s a text, give me the text.',
  'The floor, not the ceiling. If the answer needs a CRM, a stager and an assistant, it does not serve the audience and you have to push again.',
  'Individual agents, new agents.')

rich([('LISTEN CLOSELY HERE, THE HOMEWORK DEPENDS ON IT: ', True), ('The written homework in the close assumes her answer is personal, specific, and has no ask attached. If she describes something different, rewrite the homework live off her answer instead of reading the card. Her system beats my guess at her system.', False, False)])
doc.add_paragraph().paragraph_format.space_after = Pt(0)

q(2,
  'You said the best advice you ever got was "if you\'re not early, you\'re late," and that it\'s a mindset, not a clock. Fine, but give me the clock version too. One place in a deal that\'s already under contract where being early changed the outcome.',
  'What did you do, how many days out, and what would the average agent have done instead?',
  'A concrete behavior tied to her own best-advice answer, which closes the loop from rapid fire. Being early is the actual referral mechanism hiding inside the philosophy.',
  'Individual agents, team leaders.')

q(3,
  'What do agents get wrong when they decide to build a referral business? Where do they blow it in month one?',
  'What does the bad version sound like? Give me the wrong text message.',
  'The anti-pattern. Listeners remember the mistake longer than the method, and it stops them from doing the cringe version of what she just described.',
  'Individual agents, new agents.')

q(4,
  'You wake up tomorrow in Denver with everything you know and zero contacts. Nobody has your number, nobody owes you a favor, and you still refuse to cold call. What are the first ninety days?',
  'Week one specifically. What\'s on the calendar Monday?',
  'Whether the system is portable or whether it only works because she is two decades in. This is where she proves the thesis or admits the head start. Either answer is good tape.',
  'All segments.')

h3('THE HUMAN LANDING (last question before the close)')
q(5,
  'Last thing. You\'ve been open about the fact that there was real weight behind the scenes while all of this was happening. Most agents in this business act like that part doesn\'t exist, like the smile in the listing photo is the whole story. What does a listing appointment look like on a day when everything at home is falling apart?',
  'What do you do in the car before you walk in? Has a client ever known?',
  'The honest beat. It is why she is worth an hour instead of a blog post, and it is the moment people decide to follow her.',
  'All segments.')

rich([('WHY IT MOVED: ', True), ('In v1 this was the last question of Act 1, which meant going from "everything at home is falling apart" straight into "what does a listener do Monday morning." That seam was bad. Here it lands the episode and hands you a natural walk into the homework.', False, False)])

rich([('CONSENT NOTE, DO THIS BEFORE YOU RECORD: ', True), ('Her bio says she is open about having carried real personal hardship, and she has publicly referenced family health challenges without naming specifics. Ask in the green room: "Is there anything on the personal side you want me to steer toward, and anything you want me to stay away from?" Then ask exactly what she green-lights and nothing else. Never surface an illness or a family member she has not named publicly. If she declines the area entirely, cut this and give the time to Act 1 Q4, which pulls the same vulnerability from the business side.', False, False)])

h2('THE CLOSE')
rich([('HOMEWORK (read verbatim, unless Act 2 Q1 told you different): ', True)])
rich([('"Here\'s what I want you to do before the next episode. Open your database and pull five past clients you haven\'t talked to in over a year. Not a mass email. Not a market update. Five individual texts, and each one has to include one specific thing you remember about them, their kid, their dog, the awful wallpaper in the basement. No ask. Nothing about real estate. Five texts, before Friday."', False, True)])
p('Under 30 minutes, free, and everyone knows what done looks like. It is also the literal first rep of what she describes in Act 2 Q1, which is why you have to actually listen to that answer before you read this.', italic=True, size=9.5)

rich([('GUEST CLOSE: ', True)])
bullet('', '"Where can people find you and follow you?" (gwensnyderrealestate.com, Instagram @gweniveresnyder)')
bullet('', 'Then once, cleanly: "You built a program around this called Referable. Who is it actually for, and who is it not for?" The "not for" half keeps it from being a commercial, and she will answer it honestly.')
bullet('', 'Ask if she has anything free for listeners. If she does, get the exact URL on air and put it in the show notes.')

doc.add_page_break()

# =====================================================================
# SECTION 4 — RESEARCH BRIEF
# =====================================================================
h1('SECTION 4: RESEARCH BRIEF')
p('Morning-of reading. Not for use during the interview.', italic=True, size=9.5)

h2('4A. Background')
p('Colorado native. Licensed in Colorado in 2004. Her earlier published interviews describe a first career modeling in Europe, then private aviation and entertainment in Los Angeles, then high-end residential real estate in Beverly Hills, entered while pregnant with twins at the encouragement of a mentor named Mimi Starrett, before relocating to Colorado. Her current website tells it differently, as a single Colorado market entered in 2004 and never left. Both versions are hers, published four years apart. See Landmine 1 and do not adjudicate it on air. At 17 she survived a car wreck that severed her femur and required a long rehabilitation, which is the story she reaches for whenever an interviewer asks about resilience. What every source agrees on, going back to 2021: she has never run a prospecting-based business. Referral and repeat, only, the entire way.')

h2('4B. Career Timeline (verified entries only)')
table(['Year', 'Role / Company', 'Notable'],
      [['2004', 'Licensed, Colorado', 'Her site says she immediately focused on Denver metro luxury. Her earlier interviews place the start of her career in Beverly Hills. Both are hers.'],
       ['2010', 'Colorado', 'First recognized as a top producer (her own About page).'],
       ['2018', 'Colorado mountain markets', 'Expanded the practice into Evergreen and Conifer.'],
       ['2021', 'Christie\'s International Real Estate, Denver', 'Listed there March 2021. ~$150M career volume at the time. $53M sold that year.'],
       ['2022', 'LIV Sotheby\'s International Realty, Greenwood Village', 'Joined. Launched Referable the same year.'],
       ['2025', 'LIV Sotheby\'s / Referable', 'Crossed $500M career volume. RealTrends Verified and The Thousand. Denver Real Producers feature. Spoke at Real Producers Summit.'],
       ['2026', 'LIV Sotheby\'s / Referable', '5280 Top Producer, Double Black Diamond. Spoke at Denver Build 26.']],
      widths=[0.7, 2.4, 3.8])

h2('4C. What Makes Her Interesting to This Audience')
bullet('The refusal is absolute: ', 'She has never cold called and says she never will. Most "referral based" agents quietly run paid leads too. She does not, and she has two decades of production to argue with.')
bullet('97% list-to-sale: ', 'Almost unremarked on her own site, and the most interesting number she has. It implies pricing discipline and a willingness to walk that nobody has ever made her explain.')
bullet('She has started over: ', 'Whichever origin version is right, she rebuilt in Denver after a market change and changed brokerages twice in the last five years. The system traveled. That is the spine of Act 2 Q4.')
bullet('She gives away 10% of every closed deal: ', 'Not a foundation, not a year-end gift. Ten percent of every check. Unusual and specific, worth one sentence on air, not five.')
bullet('The dungeon story: ', 'On the record, clean as she tells it, and genuinely funny. The most clippable minute in the intake, with the caveats in Landmine 4.')

h2('4D. Key Data Points')
table(['Stat', 'Source', 'Confidence'],
      [['$500M+ career volume, crossed 2025', 'Her site and Instagram', 'High as a published claim. Climbed fast across sources, see Landmine 3.'],
       ['97% list-to-sale ratio', 'Her About and coaching pages', 'Medium. Self-reported, no third party.'],
       ['Licensed in Colorado 2004', 'Her About page timeline', 'High'],
       ['"23 years" in the business', 'Guest-supplied bio', 'CONFLICTS with her own site ("20+", "twenty years") and the 2004 date. Do not say a number.'],
       ['#1 on her LIV Sotheby\'s team', 'Her About page', 'Medium. Self-reported.'],
       ['$53M sold in 2021', 'Syndicated agent profiles', 'Medium'],
       ['Top 1% nationwide, RealTrends', 'Multiple independent profiles', 'High'],
       ['Top .33% nationwide', 'Her Instagram bio only', 'Unverified. Do not read on air.'],
       ['55+ pro athletes, coaches, managers', 'Multiple profiles, self-reported', 'Medium'],
       ['10% of every closed deal donated', 'CanvasRebel, Bold Journey, profiles', 'High. Consistent across years.'],
       ['5280 Top Producer 2026, Double Black Diamond', 'Her press page and the 5280 directory', 'High'],
       ['Denver Build 26, CAR conference, RP Summit', 'Her press page', 'Medium. Her site is the only source.'],
       ['Clear Cooperation effectively dead, 2026', 'Inman, May 15 2026', 'High'],
       ['81% of Sotheby\'s agents cite privacy as top client concern', 'Sotheby\'s International Realty 2026 Luxury Outlook', 'High'],
       ['Referable launched 2022', 'Her About and coaching pages', 'High']],
      widths=[2.6, 2.6, 1.7])

h2('4E. Previous Media Appearances')
bullet('Bold Journey Magazine, Dec 2023: ', 'Resilience. The car accident, "grit, follow up and consistency," mentorship, the 10% giving. Overasked territory.')
bullet('CanvasRebel Magazine, Aug 2023: ', 'Risk and origin. The accident in her father\'s words, the Beverly Hills start while pregnant with twins, Mimi Starrett, the 100% referral model. Overasked, and see Landmine 1.')
bullet('ShoutOut Colorado, Mar 2021: ', 'Short Q&A. Europe modeling, private aviation, Christie\'s, concierge service. Where the pre-real-estate career is documented.')
bullet('Live podcast recording with Jay Manansala: ', 'Billed as the lessons and strategies that got her through the toughest seasons of her career. The closest thing to a business-mechanics interview she has done. Assume "how did you survive the hard years" is spent.')
bullet('5280, Denver Real Producers, RealTrends: ', 'Rankings and directory features, not interviews.')
bullet('Nothing anywhere: ', 'A step by step teardown of how the referral engine actually runs. That is the gap this episode fills, and it is a wide one.')

h2('4F. Audience Relevance')
table(['Segment', 'What They Get'],
      [['Individual agents', 'A referral process with timing and triggers instead of a slogan, a pricing script, and a Monday-morning first rep.'],
       ['Team leaders', 'The failure mode of a referral business and its early warning signs, and where "being early" sits inside a deal already under contract.'],
       ['Broker-owners', 'What an office would have to change to produce referable agents instead of dialer operators, framed against private-listing access.'],
       ['New agents', 'The bootstrap. How the first deals happen when you have no database and refuse to prospect.']],
      widths=[1.5, 5.4])

doc.add_page_break()

# =====================================================================
# SECTION 5 — RESTREAM LIVE DESCRIPTIONS
# =====================================================================
h1('SECTION 5: RESTREAM LIVE DESCRIPTIONS & HASHTAGS')

h3('Facebook Live')
p('Gwen Snyder has sold over $500 million in Colorado luxury real estate and has never made a single cold call. Today she is walking us through the exact referral system she built instead, including the part most agents get wrong in the first month. Drop your questions in the comments!')

h3('Instagram Live')
p('$500M sold. Zero cold calls. Gwen Snyder on how to actually build a referral business. #RealEstateAgent #ReferralBusiness #LuxuryRealEstate #KeepingItReal #RealtorLife')

h3('TikTok Live')
p('She sold half a billion in real estate and never once cold called. Here is what she did instead. #realtor #realestateagent #referrals #luxuryrealestate #realestatetips')

h3('YouTube Live')
p('Gwen Snyder of LIV Sotheby\'s International Realty joins the Keeping It Real Podcast to break down the referral engine behind more than $500 million in Colorado luxury real estate sales. We cover why she has never cold called, what a 97% list to sale ratio actually requires, and how the collapse of Clear Cooperation is changing what a referral network is worth. Hosted by D.J. Paris.')

h3('LinkedIn Live')
p('Most agents are told to prospect. Gwen Snyder built a $500 million luxury practice on referrals and repeat clients without ever making a cold call. In this conversation she breaks down the system behind it, the failure mode nobody warns you about, and what the end of Clear Cooperation means for agents whose networks are now their inventory.')

h3('Hashtag Sets')
bullet('Universal: ', '#KeepingItReal #RealEstatePodcast #DJParis #RealtorLife #RealEstateAgent')
bullet('Episode-specific: ', '#ReferralBusiness #LuxuryRealEstate #DenverRealEstate #ColoradoRealEstate #RealEstateCoaching #TopProducer #SothebysRealty')
bullet('Guest tags: ', 'IG @gweniveresnyder | Facebook @GwenivereLuxuryRE | X @GwenLuxuryRE | LinkedIn /in/gweniveresnyder')

# =====================================================================
# SECTION 6 — YOUTUBE CHAPTERS
# =====================================================================
h1('SECTION 6: YOUTUBE CHAPTER MARKERS')
p('Estimates. Adjust after recording. Each title has to be worth clicking on its own.', italic=True, size=9.5)
table(['Timestamp', 'Chapter Title'],
      [['0:00', '$500 Million Sold and Never One Cold Call'],
       ['2:00', 'Rapid Fire: Best and Worst Real Estate Advice'],
       ['4:00', 'The First Deal With Zero Referrals: How to Start From Nothing'],
       ['8:00', 'The 90 Days After Closing That Turn One Client Into Three'],
       ['12:00', 'A 97% List to Sale Ratio Means Saying No: The Pricing Script'],
       ['15:00', 'The Silent Failure Mode of a Referral Business'],
       ['18:00', 'Referrals vs Paid Leads: The Actual Math'],
       ['21:00', 'The Locked Door, and What It Taught Her About Reading a Room'],
       ['24:00', 'Private Listings Are Back, So Your Network Is Now Your Inventory'],
       ['28:00', 'The Playbook: Your First Referral System in Under an Hour'],
       ['32:00', 'Starting Over in Denver With Zero Contacts: The First 90 Days'],
       ['35:00', 'What a Listing Appointment Looks Like on the Worst Day'],
       ['38:00', 'Homework and Where to Find Gwen Snyder']],
      widths=[1.2, 5.7])

doc.add_page_break()

# =====================================================================
# SECTION 7 — COUNCIL REVIEW
# =====================================================================
h1('SECTION 7: COUNCIL REVIEW (second pass)')
p('Re-run on the revised packet. Board convened: Kane, Miller, Hormozi, MrBeast, Byron Lazine, Eric Simon, Chris Do, Youshaei, with Heath and Berger called as witnesses. The first council left a dissent unresolved. This pass resolves it.', italic=True, size=9.5)

h2('Title: the dissent from v1 is now the plan')
p('Eric Simon\'s objection in the first pass was that "$500 million" reads as a luxury flex and half the audience cannot see themselves in it, while "she has never made a cold call" is every agent\'s fantasy and costs nothing to identify with. Chris Do backed him. The stress test then handed Simon a second argument he did not have the first time: her volume figure has moved a lot across four years of published sources, which is fine as her own claim but is a weak thing to hang the whole package on. Kane\'s counter still stands for one surface only, which is that on YouTube the number does search and authority work that the refusal cannot.')
rich([('Resolution, ship both: ', True), ('Title 1 (the refusal) on the podcast feed and every social cut. Title 2 (the number) on YouTube only. This is not a compromise, it is correct platform targeting, and it is the experiment v1 proposed run as policy instead of as a test.', False, False)])
table(['#', 'Title', 'Ingredient', 'Curiosity Mechanism'],
      [['1', 'She Has Never Made a Cold Call. Gwen Snyder on Why Chasing Leads Is the Slow Way to Build a Business.',
        'Insight plus personality', 'Violates a trained rule. "Slow way" is the inversion that makes an agent click to argue. Nothing to resolve it but pressing play.'],
       ['2', '$500 Million Sold, Zero Cold Calls: Gwen Snyder on the Referral Engine That Replaced Prospecting',
        'Stat plus insight', 'Gap opens on "zero." The number sets stakes, the zero makes the hole, and the method is withheld.'],
       ['3', 'The 97% List-to-Sale Agent: Gwen Snyder on Becoming the Agent People Cannot Stop Naming',
        'Stat', 'Held. Needs a beat of explanation, which a title cannot afford. Reuse as a clip caption on the Act 1 Q3 cut.']],
      widths=[0.3, 3.2, 1.0, 2.4])

h2('Cold-Open Hook (final)')
rich([('"My guest today has sold more than half a billion dollars of Colorado real estate and has never made a single cold call. Not one. Not ever, and not by accident. We\'re going to talk about that today. Stay tuned."', False, True)])
p('Heath: "not by accident" is doing the work that "in twenty-two years, on purpose" did in v1, without stating a number the packet elsewhere forbids. A refusal on its own reads as luck or a rich sphere and the gap closes itself. A refusal that was chosen has a reason behind it, and the reason is the episode. Kane adds that three fragments in a row is a rhythm device and D.J. should read them as three separate sentences, not one breath.', size=9.5)

h2('The Clip Engine')
table(['Q#', 'Question', 'Berger Emotion', 'Heath Gap'],
      [['Act 1, Q5', 'Referrals vs paid leads, numbers demanded',
        'High arousal. Anger and validation. Half the audience is paying for the thing she says is slower.',
        'Contradicts what the listener was trained to believe. Social currency: sharing it signals you are the agent who does not need to buy leads.'],
       ['Act 1, Q6', 'The locked door and "great storage"',
        'Amusement, the most shared emotion available, and it needs zero industry knowledge to enjoy.',
        'Pure narrative gap. The door is locked, then it opens. The only clip in the episode that travels outside real estate.'],
       ['Act 1, Q4', 'The referral business that goes quiet instead of crashing',
        'Anxiety. High arousal and badly under-used in this category. Names a fear agents have never articulated.',
        'Passes Berger\'s valuable-virality test: an agent retells it in one sentence and the point survives intact.'],
       ['Act 2, Q5', 'The listing appointment on the worst day',
        'Low arousal on its own, so it will not travel far. Include it anyway.',
        'Chris Do: this is not a reach clip, it is a follow clip. It converts the people the other three brought in.']],
      widths=[0.9, 1.9, 2.0, 2.1])

h2('Live-Description Scrub')
table(['Platform', 'Verdict and Fix'],
      [['Facebook', 'Keep. The question prompt is native and "the part most agents get wrong" is a second gap after the hook.'],
       ['Instagram', 'Keep. Fragments, not sentences. GaryVee: the only one that reads like it was written for its feed instead of pasted into it.'],
       ['TikTok', 'Keep. "Half a billion" outperforms "$500M" both spoken and scanned. Lowercase tags are correct here.'],
       ['YouTube', 'Updated. Now says "the collapse of Clear Cooperation" instead of the vaguer phrasing. Youshaei: this description is the only one doing real search work, so it should carry the newsiest language.'],
       ['LinkedIn', 'Keep. Byron Lazine signs off because it is the only description anchored to an actual industry headline, which is what earns a professional feed.']],
      widths=[1.1, 5.8])

h2('Arc: what the rebuild fixed and what it did not')
p('MrBeast\'s v1 note was a sag at minutes 15 to 21 where three analytical questions stacked with no story between them. The rebuild moved the locked-door story into that gap permanently rather than leaving it as a resequencing suggestion, so the fix is now structural instead of a note D.J. has to remember while recording. Welsh withdraws his v1 objection on the grounds that this is now the format rather than an exception to it.')
p('What did not get fixed: Act 1 still runs seven questions in twenty-three minutes, and MrBeast holds that the real retention risk on a systems episode is not the question count, it is that questions 2 and 3 both require her to produce a specific mechanism on demand. If she cannot, the middle goes soft no matter how it is sequenced. The mitigation is the pushback scripting on both questions, not the running order.', size=9.5)

h2('Why It Should Work')
bullet('Curiosity mechanism (Heath): ', 'The title and the cold open both state an absence and never the method. Nothing in either assumes the listener knows who Gwen Snyder is, which clears the Curse of Knowledge check that most guest-name titles fail.')
bullet('Share driver (Berger): ', 'Two independent engines with no overlap. The locked-door clip travels on amusement to anybody. The leads-versus-referrals clip travels on social currency strictly inside the industry. One recording, two audiences.')
bullet('Retention move (MrBeast): ', 'The worst-advice answer is planted at minute one and paid off at minute two, which teaches the listener early that this show closes its loops. The locked-door story then re-hooks the exact minute the analysis gets dense.')
bullet('Value density (Hormozi): ', 'Act 1 Q2 and Act 2 Q1 are the payload and they are deliberately not clips. Q2 is why an agent finishes the episode. The locked door is why they start it. Do not confuse the two jobs.')

h2('The Dissent')
p('Hormozi objects to the locked-door story getting three minutes of a systems episode. His argument: value density is the only thing that earns a listener\'s next episode, the story teaches nothing an agent can use, and every minute of it is a minute not spent making her produce a mechanism. Miller takes the other side, that the story is the only thing in forty minutes that makes her a person rather than a set of production numbers, and that nobody takes advice from a resume.')
rich([('Resolved for now: ', True), ('The story stays, capped at one beat, positioned as a pattern interrupt and not as the payload, with the analytical follow-up attached so it pays rent. That is why Q6 is split into a story and a follow-up instead of one compound question.', False, False)])
rich([('The experiment: ', True), ('Cut a version of the YouTube upload without the locked-door segment and compare 30-day average view duration against the full version. If retention holds without it, Hormozi is right and future systems episodes should stop buying entertainment beats with runtime. If it drops, the pattern interrupt is load-bearing and it should become standard on every deep-dive episode.', False, False)])

# =====================================================================
doc.save('/Users/djparis/GitHub Projects/keeping-it-real-content-system/guest-prep/Gwen_Snyder_Interview_Prep.docx')
print('Saved v2: Gwen_Snyder_Interview_Prep.docx')
